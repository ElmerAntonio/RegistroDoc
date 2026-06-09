"""
RegistroDoc Pro — Dashboard Principal
Diseño basado en la imagen de referencia:
  - Panel principal con métricas, gráfica de línea y barra
  - Marca de agua sutil de Panamá
  - Ventana fija 1280x720 (no redimensionable)
  - Paleta: fondo #0A1628, acentos cian #00DDEB / verde #00FF88
"""

import customtkinter as ctk
import tkinter as tk
import datetime
import time

# Matplotlib para las gráficas
import matplotlib
matplotlib.use("TkAgg")
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

from theme import C, FONT_TITLE, FONT_BODY, ACENTO

try:
    from utils.frases_educacion import frase_del_dia
    FRASES_OK = True
except ImportError:
    FRASES_OK = False

try:
    from tareas import obtener_pendientes, obtener_clase_actual, obtener_proxima_clase
    TAREAS_OK = True
except ImportError:
    TAREAS_OK = False


# ─── CLASE PRINCIPAL DEL DASHBOARD ─────────────────────────────────────────
class DashboardFrame(ctk.CTkFrame):
    """
    Panel dashboard completo.
    Parámetros:
        master      : contenedor padre
        engine      : DataEngine con los métodos del motor de datos
        app_principal: referencia a RegistroDocApp para navegar
    """

    def __init__(self, master, engine, app_principal=None, **kwargs):
        super().__init__(master, fg_color=C["fondo"], **kwargs)
        self.engine        = engine
        self.app           = app_principal
        self._acento       = ACENTO

        # To support dynamic individual search
        self._current_student = None

        self._student_cache = None

        # Stats una sola vez
        self._stats = self._cargar_stats()

        self._construir()

    def actualizar_vista(self):
        self._stats = self._cargar_stats()
        if hasattr(self, "_panel_container") and self._panel_container.winfo_exists():
            self._panel_container.destroy()
        self._panel_principal(self)

    # ══════════════════════════════════════════════════════════════════════
    #  DATOS
    # ══════════════════════════════════════════════════════════════════════
    def _cargar_stats(self) -> dict:
        try:
            return self.engine.get_dashboard_stats()
        except Exception:
            modal = getattr(self.engine, "modalidad", "premedia")
            fallback_g = ["7°", "8°", "9°"] if modal == "premedia" else ["1°"]
            return {
                "total": 92, "riesgo": 7, "honor": "SANTOS FIDEL",
                "honor_prom": "4.9", "asistencia": "98%",
                "grados": ["Sin datos"],
            }

    # ══════════════════════════════════════════════════════════════════════
    #  CONSTRUCCIÓN PRINCIPAL
    # ══════════════════════════════════════════════════════════════════════
    def _construir(self):
        self.pack_propagate(False)
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

        # Panel principal
        self._panel_principal(self)


    # ══════════════════════════════════════════════════════════════════════
    #  PANEL PRINCIPAL
    # ══════════════════════════════════════════════════════════════════════
    def _panel_principal(self, parent):
        panel = ctk.CTkScrollableFrame(parent, fg_color=C["fondo"],
                                       corner_radius=0,
                                       scrollbar_fg_color=C["fondo"],
                                       scrollbar_button_color=C["borde"])
        panel.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        panel.columnconfigure(0, weight=1)
        self._panel_container = panel

        # Marca de agua Panamá (canvas con texto muy tenue)
        self._marca_agua(panel)

        # Título y fecha
        self._titulo_seccion(panel)

        # 🕐 Clase actual según horario
        self._widget_clase_actual(panel)

        # Tarjetas métricas
        self._tarjetas_metricas(panel)

        # Frase motivacional del día
        self._frase_motivacional(panel)

        # 📋 Tareas pendientes programadas
        self._widget_tareas_pendientes(panel)

        # Panel de Hábitos
        self._panel_habitos(panel)

        # Asistente y Alertas del Docente
        self._asistente_docente(panel)

        # Gráficas
        self._graficas(panel)

        # Grados activos y accesos rápidos
        self._footer_panel(panel)

    # ── Marca de agua ──────────────────────────────────────────────────
    def _marca_agua(self, parent):
        """Canvas con texto PANAMÁ muy tenue como marca de agua."""
        cv = tk.Canvas(parent, bg=C["fondo"], height=0,
                       highlightthickness=0, bd=0)
        cv.pack(fill="x")
        # Se dibuja sobre toda la pantalla con place desde el panel principal
        # (técnica: se pone un label gigante transparente atrás)
        fondo_lbl = ctk.CTkLabel(
            parent,
            text="P  A  N  A  M  Á",
            font=ctk.CTkFont("Segoe UI", 72, "bold"),
            text_color="#0D2040",        # casi invisible sobre fondo
            fg_color="transparent")
        fondo_lbl.pack(pady=(0, 0))
        # No ocupa espacio visible — es un efecto decorativo

    # ── Frase Motivacional ─────────────────────────────────────────────
    def _frase_motivacional(self, parent):
        """Muestra una frase/versículo del día — offline."""
        if not FRASES_OK:
            return
        try:
            texto, referencia = frase_del_dia()
        except Exception:
            return

        frame = ctk.CTkFrame(parent, fg_color="#0C2137", corner_radius=10,
                             border_width=1, border_color="#1E3A5F")
        frame.pack(fill="x", padx=24, pady=(8, 4))

        inner = ctk.CTkFrame(frame, fg_color="transparent")
        inner.pack(fill="x", padx=16, pady=10)

        ctk.CTkLabel(inner, text="✨",
                     font=ctk.CTkFont("Segoe UI", 20)).pack(side="left", padx=(0, 10))

        txt_frame = ctk.CTkFrame(inner, fg_color="transparent")
        txt_frame.pack(side="left", fill="x", expand=True)

        ctk.CTkLabel(txt_frame, text=f"«{texto}»",
                     font=ctk.CTkFont("Segoe UI", 11, slant="italic"),
                     text_color="#94A3B8", wraplength=700,
                     anchor="w", justify="left").pack(anchor="w")

        ctk.CTkLabel(txt_frame, text=f"— {referencia}",
                     font=ctk.CTkFont("Segoe UI", 10, "bold"),
                     text_color="#FFD700", anchor="w").pack(anchor="w", pady=(2, 0))

    # ── Clase Actual según Horario ─────────────────────────────────────
    def _widget_clase_actual(self, parent):
        """Muestra qué clase tiene el docente en este momento."""
        if not TAREAS_OK:
            return
        try:
            horario = self.engine.obtener_horario()
        except Exception:
            return

        materia, rango, dia = obtener_clase_actual(horario)

        frame = ctk.CTkFrame(parent, fg_color="#0F2744", corner_radius=10,
                             border_width=1, border_color="#1E3A5F")
        frame.pack(fill="x", padx=24, pady=(8, 4))

        inner = ctk.CTkFrame(frame, fg_color="transparent")
        inner.pack(fill="x", padx=16, pady=10)

        ahora = datetime.datetime.now()
        dias_es = {0: "Lunes", 1: "Martes", 2: "Miércoles", 3: "Jueves",
                   4: "Viernes", 5: "Sábado", 6: "Domingo"}
        dia_nombre = dias_es.get(ahora.weekday(), "")
        hora_str = ahora.strftime("%I:%M %p")

        if materia:
            # Clase en progreso
            ctk.CTkLabel(inner, text="🟢",
                         font=ctk.CTkFont("Segoe UI", 18)).pack(side="left", padx=(0, 8))

            txt_f = ctk.CTkFrame(inner, fg_color="transparent")
            txt_f.pack(side="left", fill="x", expand=True)

            ctk.CTkLabel(txt_f, text=f"Clase Actual: {materia}",
                         font=ctk.CTkFont("Segoe UI", 14, "bold"),
                         text_color="#00FF88", anchor="w").pack(anchor="w")

            ctk.CTkLabel(txt_f, text=f"{dia_nombre} — {rango} — {hora_str}",
                         font=ctk.CTkFont("Segoe UI", 11),
                         text_color="#94A3B8", anchor="w").pack(anchor="w")
        else:
            # No hay clase ahora — mostrar la próxima
            prox_mat, prox_rango = obtener_proxima_clase(horario)

            ctk.CTkLabel(inner, text="🕐",
                         font=ctk.CTkFont("Segoe UI", 18)).pack(side="left", padx=(0, 8))

            txt_f = ctk.CTkFrame(inner, fg_color="transparent")
            txt_f.pack(side="left", fill="x", expand=True)

            if prox_mat:
                ctk.CTkLabel(txt_f, text=f"Próxima clase: {prox_mat}",
                             font=ctk.CTkFont("Segoe UI", 13, "bold"),
                             text_color="#38BDF8", anchor="w").pack(anchor="w")
                ctk.CTkLabel(txt_f, text=f"{dia_nombre} — {prox_rango} — Ahora: {hora_str}",
                             font=ctk.CTkFont("Segoe UI", 11),
                             text_color="#94A3B8", anchor="w").pack(anchor="w")
            else:
                msg = "No hay más clases hoy" if ahora.weekday() < 5 else "Fin de semana — ¡descanse!"
                ctk.CTkLabel(txt_f, text=msg,
                             font=ctk.CTkFont("Segoe UI", 13, "bold"),
                             text_color="#94A3B8", anchor="w").pack(anchor="w")
                ctk.CTkLabel(txt_f, text=f"{dia_nombre} — {hora_str}",
                             font=ctk.CTkFont("Segoe UI", 11),
                             text_color="#64748B", anchor="w").pack(anchor="w")

    # ── Tareas Pendientes ──────────────────────────────────────────────
    def _widget_tareas_pendientes(self, parent):
        """Muestra las tareas/evaluaciones programadas pendientes."""
        if not TAREAS_OK:
            return

        pendientes = obtener_pendientes()
        if not pendientes:
            return

        frame = ctk.CTkFrame(parent, fg_color="#0F2744", corner_radius=10,
                             border_width=1, border_color="#1E3A5F")
        frame.pack(fill="x", padx=24, pady=(8, 4))

        # Header
        hdr = ctk.CTkFrame(frame, fg_color="transparent")
        hdr.pack(fill="x", padx=16, pady=(10, 4))

        ctk.CTkLabel(hdr, text=f"📋 Tareas Pendientes ({len(pendientes)})",
                     font=ctk.CTkFont("Segoe UI", 14, "bold"),
                     text_color="#00DDEB", anchor="w").pack(side="left")

        # Lista de tareas (máximo 5 visibles) — clickables
        for t in pendientes[:5]:
            urgencia = t.get("_urgencia", "normal")
            dias = t.get("_dias_restantes", 999)

            colores = {
                "vencida": ("#EF4444", "⛔ VENCIDA"),
                "hoy": ("#F59E0B", "⚠️ HOY"),
                "urgente": ("#FB923C", f"🔔 En {dias} días"),
                "normal": ("#3B82F6", f"📅 En {dias} días"),
            }
            color, badge_txt = colores.get(urgencia, ("#3B82F6", ""))

            info = f"{badge_txt}  {t['titulo']} — {t.get('grado', '')} — {t.get('fecha_limite', '')}"

            def _ir_notas(app=self.app):
                if app:
                    try: app.mostrar_notas()
                    except: pass

            ctk.CTkButton(frame, text=info,
                          font=ctk.CTkFont("Segoe UI", 11),
                          fg_color="#0A1628", hover_color=C["hover"],
                          text_color=color, anchor="w", height=30,
                          command=_ir_notas).pack(fill="x", padx=16, pady=2)

        # Espaciado final
        ctk.CTkFrame(frame, fg_color="transparent", height=8).pack()

    # ── Título ─────────────────────────────────────────────────────────
    def _titulo_seccion(self, parent):
        f = ctk.CTkFrame(parent, fg_color="transparent")
        f.pack(fill="x", padx=24, pady=(10, 4))
        f.columnconfigure(1, weight=1)

        ctk.CTkLabel(f, text="Panel de Control Principal",
                     font=ctk.CTkFont("Segoe UI", 22, "bold"),
                     text_color=C["texto"]).grid(row=0, column=0, sticky="w")

        # Barra de búsqueda integrada al Dashboard
        search_f = ctk.CTkFrame(f, fg_color=C["input"], corner_radius=20, border_width=1, border_color=C["borde"], height=34)
        search_f.grid(row=0, column=1, sticky="ew", padx=(40, 40))
        search_f.grid_propagate(False)
        search_f.columnconfigure(1, weight=1)

        ctk.CTkLabel(search_f, text="🔍", font=ctk.CTkFont(size=14), text_color=C["texto_sec"], width=30).grid(row=0, column=0, padx=(10,0))

        self._search_var = tk.StringVar()
        self._search_entry = ctk.CTkEntry(
            search_f, textvariable=self._search_var,
            fg_color="transparent", border_width=0,
            placeholder_text="Buscar alumno para ver su rendimiento...",
            font=ctk.CTkFont("Segoe UI", 13),
            text_color=C["texto"])
        self._search_entry.grid(row=0, column=1, sticky="ew", padx=(4,10), pady=4)

        # We bind the Return key so we can perform the search
        self._search_entry.bind("<Return>", self._ejecutar_busqueda)

        fecha = datetime.datetime.now().strftime("%A %d de %B, %Y").capitalize()
        ctk.CTkLabel(f, text=fecha,
                     font=ctk.CTkFont("Segoe UI", 11),
                     text_color=C["texto_sec"]).grid(row=0, column=2, sticky="e")

        # Toast message frame for not found errors
        self._toast_lbl = ctk.CTkLabel(f, text="", text_color=C["rojo"], font=ctk.CTkFont("Segoe UI", 12, "bold"))
        self._toast_lbl.grid(row=1, column=1, pady=(2, 0))

    def _ejecutar_busqueda(self, event=None):
        texto = self._search_var.get().strip().lower()
        if not texto:
            self._current_student = None
            self._actualizar_dashboard_contextual()
            return

        encontrado = None
        try:
            if self._student_cache is None:
                start_cache = time.perf_counter()
                self._student_cache = []
                for g in self.engine.obtener_grados_activos():
                    for est in self.engine.obtener_estudiantes_completos(g):
                        self._student_cache.append({"grado": g, "nombre": est["nombre"]})
                print(f"[*] Built student cache in {time.perf_counter() - start_cache:.4f}s")

            start_search = time.perf_counter()
            for cached_est in self._student_cache:
                if texto in cached_est["nombre"].lower():
                    encontrado = cached_est
                    break
            print(f"[*] Search completed in {time.perf_counter() - start_search:.6f}s")

        except Exception as e:
            # Fallback mock for UI test
            print(f"[!] Search error: {e}")

        if encontrado:
            self._current_student = encontrado
            self._toast_lbl.configure(text="")
            self._actualizar_dashboard_contextual()
        else:
            self._toast_lbl.configure(text="❌ Estudiante no registrado")
            self._search_var.set("")
            self._current_student = None
            self._actualizar_dashboard_contextual()
            # Ocultar toast despues de 3 segundos
            def clear_toast():
                if self.winfo_exists() and hasattr(self, "_toast_lbl") and self._toast_lbl.winfo_exists():
                    self._toast_lbl.configure(text="")
            self.after(3000, clear_toast)

    def _actualizar_dashboard_contextual(self):
        # Refresh the cards and graph based on _current_student state
        if hasattr(self, "_cards_frame") and self._cards_frame.winfo_exists():
            self._cards_frame.destroy()
            self._tarjetas_metricas(self._panel_container)

        if hasattr(self, "_habitos_panel") and self._habitos_panel.winfo_exists():
            self._habitos_panel.destroy()
        self._panel_habitos(self._panel_container)

        if hasattr(self, "_graph_frame") and self._graph_frame.winfo_exists():
            self._graph_frame.destroy()
            self._graficas(self._panel_container)

    # ── Tarjetas métricas ────────────────────────────────────────────────
    def _tarjetas_metricas(self, parent):
        self._cards_frame = ctk.CTkFrame(parent, fg_color="transparent")
        # Ensure it always stays above footer, below title. Since pack is used, we have to insert it at correct index if recreating
        # We will pack it directly, but since we are recreating it during dynamic search,
        # a better approach would be to have dedicated container frames.
        # But we'll just reconstruct the entire view or place it carefully.
        # Actually it's easier to just destroy and pack everything below the title again.

        # Let's clean the container elements below the title before rendering
        for widget in parent.winfo_children():
            if widget not in (self._cards_frame, parent.winfo_children()[0], parent.winfo_children()[1]):
                if hasattr(self, "_cards_frame") and widget == self._cards_frame: continue
                # Do not destroy water mark or title section

        self._cards_frame.pack(fill="x", padx=20, pady=(6, 10))
        self._cards_frame.columnconfigure((0, 1, 2, 3), weight=1, uniform="card")

        if self._current_student:
            # Contextual mode for individual student
            nombre_est = self._current_student["nombre"]
            grado_est = self._current_student["grado"]
            
            # Find student index/ID
            id_est = None
            try:
                estudiantes = self.engine.obtener_estudiantes_completos(grado_est)
                for idx, est in enumerate(estudiantes):
                    if est["nombre"] == nombre_est:
                        id_est = str(idx + 1)
                        break
            except Exception:
                pass
                
            # Averages & Risk
            prom_val = None
            try:
                proms = self.engine.obtener_promedios_reales(grado_est, None, "Anual")
                if not proms:
                    proms = self.engine.obtener_promedios_reales(grado_est, None, "Trimestre 1")
                if proms and nombre_est in proms:
                    prom_val = proms[nombre_est]
            except Exception:
                pass
                
            if prom_val is not None:
                risk_status = "En Riesgo" if prom_val < 3.0 else "Estable"
                risk_sub = f"Promedio: {prom_val:.2f}"
                risk_color = C["rojo"] if prom_val < 3.0 else C["verde"]
            else:
                risk_status = "Sin notas"
                risk_sub = "No registrado"
                risk_color = C["texto_sec"]

            # Attendance calculations for individual
            total_dias = 0
            ausencias = 0
            tardanzas = 0
            excusas = 0
            if id_est:
                try:
                    for trim in ["Trimestre 1", "Trimestre 2", "Trimestre 3"]:
                        res = self.engine.obtener_estadisticas_asistencia(grado_est, trim, id_est)
                        total_dias += res.get("total_dias", 0)
                        ausencias += res.get("ausencias", 0)
                        tardanzas += res.get("tardanzas", 0)
                        excusas += res.get("excusas", 0)
                except Exception:
                    pass
            
            asist_pct = "0%"
            asist_sub = "Sin datos"
            if total_dias > 0:
                presentes = total_dias - ausencias
                pct = (presentes / total_dias) * 100
                asist_pct = f"{pct:.0f}%"
                asist_sub = f"{ausencias} faltas / {excusas} excusas"

            datos = [
                ("👤",  "Estudiante",           grado_est, nombre_est[:20],            self._acento),
                ("⚠️",  "Estatus de Riesgo",    risk_status, risk_sub,                 risk_color),
                ("🏆",  "Cuadro de Honor",      "—",    "No aplica",                C["texto_sec"]),
                ("📊",  "Asistencia Indiv.",    asist_pct, asist_sub,                  self._acento),
            ]
        else:
            # General mode
            s = self._stats
            total = str(s.get("total", 0))
            riesgo = str(s.get("riesgo", 0))
            honor = str(s.get("honor", "N/A"))
            asistencia = str(s.get("asistencia", "0%"))

            if honor == "N/A" or not honor or honor.strip() == "":
                honor_val = "—"
                honor_sub = "Sin registros"
            else:
                honor_val = "1"
                honor_sub = honor

            datos = [
                ("👥",  "Total Alumnos",       total,  "Matriculados",             self._acento),
                ("⚠️",  "Alumnos en Riesgo",   riesgo, "Nota promedio < 3.0",      C["rojo"]),
                ("🏆",  "Cuadro de Honor",     honor_val, honor_sub[:18],          C["amarillo"]),
                ("📊",  "Asistencia",          asistencia, "Promedio general",     C["verde"]),
            ]

        for col, (ico, titulo, valor, sub, color) in enumerate(datos):
            self._tarjeta(self._cards_frame, col, ico, titulo, valor, sub, color)

    def _tarjeta(self, parent, col, ico, titulo, valor, sub, color):
        card = ctk.CTkFrame(parent, fg_color=C["card"],
                            border_width=1, border_color=color,
                            corner_radius=12)
        card.grid(row=0, column=col, padx=6, pady=4, sticky="nsew")

        # Icono
        ctk.CTkLabel(card, text=ico, font=ctk.CTkFont(size=20),
                     fg_color="transparent",
                     text_color=color).pack(anchor="w", padx=16, pady=(14, 0))

        # Título (con wraplength para que no desborde)
        ctk.CTkLabel(card, text=titulo,
                     font=ctk.CTkFont("Segoe UI", 12),
                     text_color=C["texto_sec"],
                     wraplength=140,
                     justify="left").pack(anchor="w", padx=16, pady=(2, 0))

        # Valor grande
        ctk.CTkLabel(card, text=valor,
                     font=ctk.CTkFont("Segoe UI", 32, "bold"),
                     text_color=color,
                     wraplength=160,
                     justify="left").pack(anchor="w", padx=16, pady=(0, 2))

        # Subtítulo
        ctk.CTkLabel(card, text=sub,
                     font=ctk.CTkFont("Segoe UI", 11),
                     text_color=C["texto_sec"],
                     wraplength=150,
                     justify="left").pack(anchor="w", padx=16, pady=(0, 14))

    def _panel_habitos(self, parent):
        if self._current_student:
            return

        hab_stats = self._stats.get("habitos", {"S": 0, "R": 0, "X": 0})
        s_val = hab_stats.get("S", 0)
        r_val = hab_stats.get("R", 0)
        x_val = hab_stats.get("X", 0)
        total_vals = s_val + r_val + x_val

        self._habitos_panel = ctk.CTkFrame(parent, fg_color=C["card"], border_width=1, border_color=C["borde"], corner_radius=12)
        self._habitos_panel.pack(fill="x", padx=20, pady=(6, 10))

        # Title
        ctk.CTkLabel(self._habitos_panel, text="🧠 Indicador General de Hábitos y Actitudes (Grupo Consejero 8° B)",
                     font=ctk.CTkFont("Segoe UI", 14, "bold"),
                     text_color=C["cian"]).pack(anchor="w", padx=20, pady=(15, 10))

        # Container columns
        cols_frame = ctk.CTkFrame(self._habitos_panel, fg_color="transparent")
        cols_frame.pack(fill="x", padx=20, pady=(0, 20))
        cols_frame.columnconfigure((0, 1, 2), weight=1, uniform="habito_col")

        metrics = [
            ("Satisfactorio (S)", s_val, C["verde"]),
            ("Regular (R)", r_val, C["amarillo"]),
            ("No Satisface (X)", x_val, C["rojo"])
        ]

        for idx, (label, count, color) in enumerate(metrics):
            col_frame = ctk.CTkFrame(cols_frame, fg_color="transparent")
            col_frame.grid(row=0, column=idx, padx=10, pady=5, sticky="nsew")

            pct = (count / total_vals) * 100 if total_vals > 0 else 0
            
            ctk.CTkLabel(col_frame, text=label, font=ctk.CTkFont("Segoe UI", 12, "bold"), text_color=C["texto"]).pack(anchor="w")
            ctk.CTkLabel(col_frame, text=f"{count} evaluaciones ({pct:.1f}%)", font=ctk.CTkFont("Segoe UI", 11), text_color=C["texto_sec"]).pack(anchor="w", pady=(0, 6))

            pb = ctk.CTkProgressBar(col_frame, progress_color=color, fg_color=C["borde"], height=8)
            pb.pack(fill="x")
            pb.set(pct / 100.0)

    # ── Gráficas ─────────────────────────────────────────────────────────
    def _graficas(self, parent):
        self._graph_frame = ctk.CTkFrame(parent, fg_color="transparent")
        self._graph_frame.pack(fill="x", padx=20, pady=6)
        self._graph_frame.columnconfigure(0, weight=6)
        self._graph_frame.columnconfigure(1, weight=4)

        self._grafica_linea(self._graph_frame)
        self._grafica_barras(self._graph_frame)

    def _grafica_linea(self, parent):
        """Gráfica de escalada — evolución de notas T1→T2→T3."""
        card = ctk.CTkFrame(parent, fg_color=C["card"],
                            border_width=1, border_color=C["borde"],
                            corner_radius=12)
        card.grid(row=0, column=0, padx=(0, 8), sticky="nsew")

        ctk.CTkLabel(card, text="Tendencia de Rendimiento Académico  (X/Y)",
                     font=ctk.CTkFont("Segoe UI", 13, "bold"),
                     text_color=C["texto"]).pack(anchor="w", padx=16, pady=(12, 0))
        ctk.CTkLabel(card, text="Escala 1–5 por trimestre",
                     font=ctk.CTkFont("Segoe UI", 10),
                     text_color=C["texto_sec"]).pack(anchor="w", padx=16)

        # Matplotlib
        fig = Figure(figsize=(5.4, 2.8), dpi=96,
                     facecolor=C["card"])
        ax = fig.add_subplot(111)
        ax.set_facecolor(C["card"])
        fig.subplots_adjust(left=0.1, right=0.97, top=0.92, bottom=0.16)

        grados = self.engine.obtener_grados_activos()
        if not grados:
            grados = ["7°", "8°", "9°"] if self.engine.modalidad == "premedia" else ["1°"]

        x_labels = ["T1", "T2", "T3"]
        x = [1, 2, 3]
        
        y = []
        for t_name in ["Trimestre 1", "Trimestre 2", "Trimestre 3"]:
            notas_t = []
            for g in grados:
                proms = self.engine.obtener_promedios_reales(g, None, t_name)
                if proms:
                    notas_t.extend([v for v in proms.values() if v is not None])
            if notas_t:
                y.append(round(sum(notas_t)/len(notas_t), 2))
            else:
                y.append(1.0)

        ax.plot(x, y, color=self._acento, linewidth=2.2, label="Promedio Grupal",
                zorder=3, solid_capstyle="round")
        ax.fill_between(x, y, alpha=0.12, color=self._acento)
        ax.scatter(x, y, color=self._acento, s=55, zorder=5,
                   edgecolors=C["fondo"], linewidths=1.5)

        if getattr(self, "_current_student", None):
            student_name = self._current_student["nombre"]
            y_indiv = []
            for t_name in ["Trimestre 1", "Trimestre 2", "Trimestre 3"]:
                student_notes = []
                for g in grados:
                    proms = self.engine.obtener_promedios_reales(g, None, t_name)
                    if student_name in proms:
                        student_notes.append(proms[student_name])
                if student_notes:
                    y_indiv.append(round(sum(student_notes)/len(student_notes), 2))
                else:
                    y_indiv.append(1.0)
            ax.plot(x, y_indiv, color=C["amarillo"], linewidth=2.2, label=student_name[:12],
                    zorder=4, linestyle="--")
            ax.scatter(x, y_indiv, color=C["amarillo"], s=60, zorder=6,
                       edgecolors=C["fondo"], linewidths=1.5)
            ax.legend(loc="upper left", facecolor=C["fondo"], edgecolor=C["borde"], labelcolor=C["texto"], fontsize=8)

        # Ejes
        ax.set_xlim(0.7, 3.3)
        ax.set_ylim(1, 5.2)
        ax.set_xlabel("Trimestres", color=C["texto_sec"], fontsize=9)
        ax.set_ylabel("Notas 1–5", color=C["texto_sec"], fontsize=9)
        ax.tick_params(colors=C["texto_sec"], labelsize=8)
        for spine in ax.spines.values():
            spine.set_edgecolor(C["borde"])
        ax.grid(True, color=C["borde"], linewidth=0.5, alpha=0.6)
        ax.set_xticks(x)
        ax.set_xticklabels(x_labels)

        canvas = FigureCanvasTkAgg(fig, master=card)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="x", padx=10, pady=(4, 12))

    def _grafica_barras(self, parent):
        """Barras de asistencia / rendimiento por grupo."""
        card = ctk.CTkFrame(parent, fg_color=C["card"],
                            border_width=1, border_color=C["borde"],
                            corner_radius=12)
        card.grid(row=0, column=1, sticky="nsew")

        ctk.CTkLabel(card, text="Rendimiento por Grupo",
                     font=ctk.CTkFont("Segoe UI", 13, "bold"),
                     text_color=C["texto"]).pack(anchor="w", padx=16, pady=(12, 0))
        ctk.CTkLabel(card, text="Promedio final por sección",
                     font=ctk.CTkFont("Segoe UI", 10),
                     text_color=C["texto_sec"]).pack(anchor="w", padx=16)

        fig2 = Figure(figsize=(3.2, 2.8), dpi=96,
                      facecolor=C["card"])
        ax2 = fig2.add_subplot(111)
        ax2.set_facecolor(C["card"])
        fig2.subplots_adjust(left=0.14, right=0.97, top=0.9, bottom=0.18)

        grados = self.engine.obtener_grados_activos()
        if not grados:
            grados = ["7°", "8°", "9°"] if self.engine.modalidad == "premedia" else ["1°"]

        valores = []
        for g in grados:
            proms = self.engine.obtener_promedios_reales(g, None, "Anual")
            if not proms:
                proms = self.engine.obtener_promedios_reales(g, None, "Trimestre 1")
            if proms:
                avg = sum(proms.values()) / len(proms)
                valores.append(round(avg, 2))
            else:
                valores.append(1.0)

        colores  = [self._acento] * len(grados)

        ax2.bar(grados, valores, color=colores, width=0.6,
                edgecolor=C["fondo"], linewidth=0.8)

        ax2.set_ylim(1, 5.3)
        ax2.tick_params(colors=C["texto_sec"], labelsize=8)
        for spine in ax2.spines.values():
            spine.set_edgecolor(C["borde"])
        ax2.grid(True, axis="y", color=C["borde"],
                 linewidth=0.5, alpha=0.5)
        ax2.set_ylabel("Promedio", color=C["texto_sec"], fontsize=8)

        canvas2 = FigureCanvasTkAgg(fig2, master=card)
        canvas2.draw()
        canvas2.get_tk_widget().pack(fill="x", padx=10, pady=(4, 12))

    # ── Footer: grados activos + accesos rápidos ─────────────────────────
    def _footer_panel(self, parent):
        self._footer_frame = ctk.CTkFrame(parent, fg_color="transparent")
        ff = self._footer_frame
        ff.pack(fill="x", padx=20, pady=(4, 20))
        ff.columnconfigure((0, 1, 2), weight=1, uniform="bot")

        # Grados activos
        gc = ctk.CTkFrame(ff, fg_color=C["card"],
                          border_width=1, border_color=C["borde"],
                          corner_radius=12)
        gc.grid(row=0, column=0, padx=(0, 8), sticky="nsew", pady=4)

        ctk.CTkLabel(gc, text="Grados Activos",
                     font=ctk.CTkFont("Segoe UI", 14, "bold"),
                     text_color=C["texto"]).pack(anchor="w", padx=16, pady=(14, 8))

        try:
            grados = self.engine.obtener_grados_activos()
        except Exception:
            grados = ["Sin datos"]

        gf = ctk.CTkFrame(gc, fg_color="transparent")
        gf.pack(padx=16, pady=(0, 14))
        for g in grados:
            pill = ctk.CTkLabel(
                gf, text=f"  {g}  ",
                fg_color=C["input"],
                corner_radius=16,
                font=ctk.CTkFont("Segoe UI", 14, "bold"),
                text_color=self._acento,
                padx=12, pady=6)
            pill.pack(side="left", padx=4)

        # Tareas Próximas (reemplaza el horario duplicado)
        hc = ctk.CTkFrame(ff, fg_color=C["card"],
                          border_width=1, border_color=C["borde"],
                          corner_radius=12)
        hc.grid(row=0, column=1, padx=(0, 8), sticky="nsew", pady=4)

        ctk.CTkLabel(hc, text="📌 Tareas Próximas",
                     font=ctk.CTkFont("Segoe UI", 14, "bold"),
                     text_color=C["texto"]).pack(anchor="w", padx=16, pady=(14, 8))

        hf = ctk.CTkFrame(hc, fg_color="transparent")
        hf.pack(fill="x", padx=16, pady=(0, 14))

        # Mostrar próximas tareas si el módulo está disponible
        try:
            from tareas import obtener_pendientes as _pend
            pendientes = _pend()
            if pendientes:
                for t in pendientes[:3]:
                    urgencia = t.get("_urgencia", "normal")
                    colores_u = {"vencida": "#EF4444", "hoy": "#F59E0B",
                                 "urgente": "#FB923C", "normal": "#3B82F6"}
                    color = colores_u.get(urgencia, "#3B82F6")

                    # Botón clickable que lleva a Notas
                    def _ir_notas_tarea(g=t.get("grado"), m=t.get("materia"), app=self.app):
                        if app:
                            try:
                                app.mostrar_notas()
                                notas_frame = app._frames.get("NotasFrame")
                                if notas_frame:
                                    if g:
                                        notas_frame.combo_grado.set(g)
                                        notas_frame.al_cambiar_grado(g)
                                    if m and m != "General":
                                        opts = notas_frame.combo_materia.cget("values")
                                        if m in opts:
                                            notas_frame.combo_materia.set(m)
                                        notas_frame.cargar_descripciones()
                                    notas_frame.tabs.set("Tareas")
                            except Exception:
                                pass

                    btn = ctk.CTkButton(hf, text=f"📝 {t['titulo'][:25]}",
                                        font=ctk.CTkFont("Segoe UI", 11),
                                        fg_color=C["card_alt"], hover_color=C["hover"],
                                        text_color=color, anchor="w", height=28,
                                        command=_ir_notas_tarea)
                    btn.pack(fill="x", pady=2)
            else:
                ctk.CTkLabel(hf, text="✅ Sin tareas pendientes",
                             font=ctk.CTkFont("Segoe UI", 12),
                             text_color=C["verde"]).pack(anchor="w")
        except Exception:
            ctk.CTkLabel(hf, text="📌 Use Tareas para programar",
                         font=ctk.CTkFont("Segoe UI", 12),
                         text_color=C["texto_sec"]).pack(anchor="w")

        # Exportaciones
        ex = ctk.CTkFrame(ff, fg_color=C["card"],
                          border_width=1, border_color=C["borde"],
                          corner_radius=12)
        ex.grid(row=0, column=2, padx=(8, 0), sticky="nsew", pady=4)

        ctk.CTkLabel(ex, text="Exportar Calificaciones",
                     font=ctk.CTkFont("Segoe UI", 14, "bold"),
                     text_color=C["texto"]).pack(anchor="w", padx=16, pady=(14, 8))

        exf = ctk.CTkFrame(ex, fg_color="transparent")
        exf.pack(fill="both", expand=True, padx=16, pady=(0, 14))

        self.cb_trim_export = ctk.CTkOptionMenu(exf, values=["Trimestre 1", "Trimestre 2", "Trimestre 3", "Todos los trimestres"])
        self.cb_trim_export.pack(fill="x", pady=5)
        self.cb_trim_export.set("Todos los trimestres")

        btn_f_ex = ctk.CTkFrame(exf, fg_color="transparent")
        btn_f_ex.pack(fill="x", pady=5)

        ctk.CTkButton(btn_f_ex, text="📄 PDF", width=40, fg_color="#E11D48", command=lambda: self.exportar_calificaciones("pdf")).pack(side="left", padx=2, expand=True)
        ctk.CTkButton(btn_f_ex, text="📝 Word", width=40, fg_color="#2563EB", command=lambda: self.exportar_calificaciones("docx")).pack(side="left", padx=2, expand=True)
        ctk.CTkButton(btn_f_ex, text="📊 Excel", width=40, fg_color="#059669", command=lambda: self.exportar_calificaciones("xlsx")).pack(side="left", padx=2, expand=True)



    def exportar_calificaciones(self, formato):
        trimestre = self.cb_trim_export.get()
        import tempfile
        import os
        from tkinter import messagebox
        import pandas as pd

        try:
            grados = self.engine.obtener_grados_activos()
            if not grados: return messagebox.showerror("Error", "No hay grados activos")

            # Recopilar datos
            todas_notas = []
            for grado in grados:
                materias = self.engine.obtener_materias_por_grado(grado)
                estudiantes = self.engine.obtener_estudiantes_completos(grado)
                if not materias or not estudiantes: continue

                if trimestre == "Todos los trimestres":
                    bulk_t1 = getattr(self.engine, 'obtener_promedios_reales_bulk', lambda g,m,t: {})(grado, materias, "Trimestre 1")
                    bulk_t2 = getattr(self.engine, 'obtener_promedios_reales_bulk', lambda g,m,t: {})(grado, materias, "Trimestre 2")
                    bulk_t3 = getattr(self.engine, 'obtener_promedios_reales_bulk', lambda g,m,t: {})(grado, materias, "Trimestre 3")
                else:
                    bulk_trim = getattr(self.engine, 'obtener_promedios_reales_bulk', lambda g,m,t: {})(grado, materias, trimestre)

                for est in estudiantes:
                    nombre = est['nombre']
                    for mat in materias:
                        if trimestre == "Todos los trimestres":
                            t1 = bulk_t1.get(mat, {}).get(nombre, 0)
                            t2 = bulk_t2.get(mat, {}).get(nombre, 0)
                            t3 = bulk_t3.get(mat, {}).get(nombre, 0)
                            todas_notas.append({"Grado": grado, "Materia": mat, "Estudiante": nombre, "T1": t1, "T2": t2, "T3": t3})
                        else:
                            val = bulk_trim.get(mat, {}).get(nombre, 0)
                            todas_notas.append({"Grado": grado, "Materia": mat, "Estudiante": nombre, trimestre: val})

            if not todas_notas:
                return messagebox.showinfo("Aviso", "No se encontraron datos para exportar.")

            df = pd.DataFrame(todas_notas)

            if formato == "xlsx":
                with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                    df.to_excel(tmp.name, index=False)
                    res_path = tmp.name
                messagebox.showinfo("Exportar Excel", f"Datos exportados a:\n{res_path}")
            elif formato == "docx":
                from docx import Document
                doc = Document()
                try:
                    from utils.footer_utils import add_header_with_logo, get_school_logo_path
                    logo_esc = get_school_logo_path()
                    if logo_esc:
                        add_header_with_logo(doc, logo_esc)
                except Exception: pass

                doc.add_heading(f"Calificaciones - {trimestre}", 0)
                table = doc.add_table(rows=1, cols=len(df.columns))
                for i, col in enumerate(df.columns): table.rows[0].cells[i].text = col
                for idx, row in df.iterrows():
                    cells = table.add_row().cells
                    for i, val in enumerate(row): cells[i].text = str(val)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    doc.save(tmp.name)
                    res_path = tmp.name
                messagebox.showinfo("Exportar Word", f"Datos exportados a:\n{res_path}")
            elif formato == "pdf":
                from docx import Document
                doc = Document()
                try:
                    from utils.footer_utils import add_header_with_logo, get_school_logo_path
                    logo_esc = get_school_logo_path()
                    if logo_esc:
                        add_header_with_logo(doc, logo_esc)
                except Exception: pass

                doc.add_heading(f"Calificaciones - {trimestre}", 0)
                table = doc.add_table(rows=1, cols=len(df.columns))
                for i, col in enumerate(df.columns): table.rows[0].cells[i].text = col
                for idx, row in df.iterrows():
                    cells = table.add_row().cells
                    for i, val in enumerate(row): cells[i].text = str(val)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    doc.save(tmp.name)
                    tmp_docx = tmp.name
                try:
                    import comtypes.client
                    pdf_path = tmp_docx.replace(".docx", ".pdf")
                    word = comtypes.client.CreateObject('Word.Application')
                    doc_obj = word.Documents.Open(tmp_docx)
                    doc_obj.SaveAs(pdf_path, FileFormat=17)
                    doc_obj.Close()
                    word.Quit()
                    messagebox.showinfo("Exportar PDF", f"PDF generado en:\n{pdf_path}")
                except Exception:
                    messagebox.showinfo("Aviso PDF", f"Documento generado en DOCX:\n{tmp_docx}\n\nNo se pudo convertir automáticamente a PDF (requiere MS Word en Windows).")

        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error al exportar: {str(e)}")

    def _obtener_materia_actual(self):
        """Obtiene la materia actual basada en el día y hora actuales."""
        try:
            # Obtener día de la semana (0=lunes, 6=domingo)
            dia_semana = datetime.datetime.now().weekday()
            
            # Mapear a nombres del horario (solo días de semana)
            dias_horario = ["lunes", "martes", "miercoles", "jueves", "viernes"]
            
            if dia_semana >= 5:  # Sábado o domingo
                return None, None
            
            dia_actual = dias_horario[dia_semana]
            
            # Obtener hora actual
            hora_actual = datetime.datetime.now().time()
            
            # Obtener horario del engine
            horario = self.engine.obtener_horario()
            
            # Buscar el período actual
            for periodo in horario:
                if not periodo["horas"] or not periodo[dia_actual]:
                    continue
                
                # Parsear el rango de horas (ej: "7:00-7:45")
                try:
                    horas_str = periodo["horas"].strip()
                    if "-" in horas_str:
                        inicio_str, fin_str = horas_str.split("-")
                        inicio = datetime.datetime.strptime(inicio_str.strip(), "%H:%M").time()
                        fin = datetime.datetime.strptime(fin_str.strip(), "%H:%M").time()
                        
                        if inicio <= hora_actual <= fin:
                            return periodo[dia_actual], periodo["horas"]
                except (ValueError, AttributeError):
                    continue
            
            return None, None
            
        except Exception:
            return None, None

    def _asistente_docente(self, parent):
        frame = ctk.CTkFrame(parent, fg_color=C["card"], border_width=1, border_color=C["borde"], corner_radius=12)
        frame.pack(fill="x", padx=24, pady=(10, 15))
        frame.columnconfigure(0, weight=6, uniform="asist")
        frame.columnconfigure(1, weight=4, uniform="asist")

        # Columna Izquierda: Alertas y Alumnos en Riesgo
        col_izq = ctk.CTkFrame(frame, fg_color="transparent")
        col_izq.grid(row=0, column=0, padx=16, pady=16, sticky="nsew")

        ctk.CTkLabel(col_izq, text="🔔 Estado y Alertas de Alumnos",
                     font=ctk.CTkFont("Segoe UI", 14, "bold"),
                     text_color=C["cian"]).pack(anchor="w", pady=(0, 6))

        # Analizar promedios para buscar alumnos con notas < 3.0
        alertas = []
        try:
            grados = self.engine.obtener_grados_activos()
            for g in grados:
                proms = self.engine.obtener_promedios_reales(g, None, "Trimestre 1")
                if not proms:
                    proms = self.engine.obtener_promedios_reales(g, None, "Anual")
                for nom, prom in proms.items():
                    if prom is not None and prom < 3.0:
                        alertas.append(f"⚠️ {nom} (Promedio: {prom:.2f}) — {g}")
        except Exception:
            pass

        if not alertas:
            lbl_ex = ctk.CTkLabel(col_izq, text="✅ Todo en orden. Ningún estudiante se encuentra en riesgo académico actualmente.",
                                  font=ctk.CTkFont("Segoe UI", 12),
                                  text_color=C["verde"], anchor="w", justify="left")
            lbl_ex.pack(fill="x", pady=10)
        else:
            scroll_al = ctk.CTkScrollableFrame(col_izq, fg_color="#0A1628", height=90, corner_radius=8)
            scroll_al.pack(fill="both", expand=True, pady=4)
            for alert in alertas:
                ctk.CTkLabel(scroll_al, text=alert,
                             font=ctk.CTkFont("Segoe UI", 12),
                             text_color=C["rojo"], anchor="w").pack(fill="x", padx=8, pady=2)

        # Columna Derecha: Accesos Rápidos
        col_der = ctk.CTkFrame(frame, fg_color="transparent")
        col_der.grid(row=0, column=1, padx=16, pady=16, sticky="nsew")

        ctk.CTkLabel(col_der, text="⚡ Acciones Rápidas (Un Clic)",
                     font=ctk.CTkFont("Segoe UI", 14, "bold"),
                     text_color=C["cian"]).pack(anchor="w", pady=(0, 8))

        btn_grid = ctk.CTkFrame(col_der, fg_color="transparent")
        btn_grid.pack(fill="both", expand=True)
        btn_grid.columnconfigure((0, 1), weight=1, uniform="btns")
        btn_grid.rowconfigure((0, 1), weight=1)

        # Navegaciones de botones
        def ir_notas():
            if self.app: self.app.mostrar_notas()
        def ir_asistencia():
            if self.app: self.app.mostrar_asistencia()
        def ir_observaciones():
            if self.app: self.app.mostrar_observaciones()

        btn_n = ctk.CTkButton(btn_grid, text="📝 Notas", font=ctk.CTkFont("Segoe UI", 11, "bold"),
                              fg_color=C["input"], hover_color=C["hover"], height=32, command=ir_notas)
        btn_n.grid(row=0, column=0, padx=4, pady=4, sticky="ew")

        btn_a = ctk.CTkButton(btn_grid, text="📅 Asistencia", font=ctk.CTkFont("Segoe UI", 11, "bold"),
                              fg_color=C["input"], hover_color=C["hover"], height=32, command=ir_asistencia)
        btn_a.grid(row=0, column=1, padx=4, pady=4, sticky="ew")

        btn_o = ctk.CTkButton(btn_grid, text="🔍 Observar", font=ctk.CTkFont("Segoe UI", 11, "bold"),
                              fg_color=C["input"], hover_color=C["hover"], height=32, command=ir_observaciones)
        btn_o.grid(row=1, column=0, padx=4, pady=4, sticky="ew")

        btn_r = ctk.CTkButton(btn_grid, text="💾 Respaldo", font=ctk.CTkFont("Segoe UI", 11, "bold"),
                              fg_color="#10B981", hover_color="#059669", height=32, command=self._ejecutar_respaldo_manual)
        btn_r.grid(row=1, column=1, padx=4, pady=4, sticky="ew")

    def _ejecutar_respaldo_manual(self):
        import shutil
        import datetime
        import os
        from tkinter import messagebox
        try:
            ruta_original = self.engine.ruta
            if not os.path.exists(ruta_original):
                messagebox.showerror("Error", "No se encontró el archivo de datos original para respaldar.")
                return
                
            dir_base = os.path.dirname(os.path.abspath(ruta_original))
            dir_respaldos = os.path.join(dir_base, "Respaldos_Locales")
            if not os.path.exists(dir_respaldos):
                os.makedirs(dir_respaldos)
                
            ahora = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            nombre_base = os.path.basename(ruta_original).replace(".xlsx", "")
            nombre_respaldo = f"{nombre_base}_respaldo_{ahora}.xlsx"
            ruta_destino = os.path.join(dir_respaldos, nombre_respaldo)
            
            shutil.copy2(ruta_original, ruta_destino)
            
            # Limpiar antiguos
            respaldos = sorted(
                [os.path.join(dir_respaldos, f) for f in os.listdir(dir_respaldos) if f.endswith(".xlsx")],
                key=os.path.getmtime
            )
            while len(respaldos) > 15:
                viejo = respaldos.pop(0)
                try: os.remove(viejo)
                except: pass
                
            # Buscar root window para invocar mostrar_toast si existe
            root = self.winfo_toplevel()
            if hasattr(root, "mostrar_toast"):
                root.mostrar_toast("✓ ¡Respaldo local creado con éxito!", color="#10B981")
            else:
                messagebox.showinfo("Éxito", f"Respaldo creado con éxito en:\n{ruta_destino}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo crear el respaldo: {e}")


# ─── MOCK PARA PRUEBA STANDALONE ─────────────────────────────────────────────
if __name__ == "__main__":
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("blue")

    class MockEngine:
        def get_dashboard_stats(self):
            return {"total": 92, "riesgo": 7,
                    "honor": "SANTOS FIDEL", "honor_prom": "4.9",
                    "asistencia": "98%"}
        def obtener_grados_activos(self):
            return ["Sin datos"]
        def obtener_estudiantes_completos(self, g):
            names = ["Maria Gonzalez", "Pedro Pinto", "Ana Santos"]
            return [{"id": i+1, "nombre": n, "cedula": ""} for i, n in enumerate(names)]
        def obtener_datos_generales(self):
            return {"docente_nombre": "Prof. Elmer Tugri",
                    "correo": "elmer.tugri7@meduca.edu.pa"}
        def obtener_horario(self):
            return [
                {"horas": "7:00-7:45", "lunes": "Matemáticas", "martes": "Español", "miercoles": "Ciencias", "jueves": "Historia", "viernes": "Inglés"},
                {"horas": "7:45-8:30", "lunes": "Español", "martes": "Matemáticas", "miercoles": "Inglés", "jueves": "Ciencias", "viernes": "Historia"},
                {"horas": "8:30-9:15", "lunes": "Ciencias", "martes": "Historia", "miercoles": "Matemáticas", "jueves": "Español", "viernes": "Ciencias"},
                {"horas": "9:15-10:00", "lunes": "Historia", "martes": "Inglés", "miercoles": "Español", "jueves": "Matemáticas", "viernes": "Inglés"},
                {"horas": "10:15-11:00", "lunes": "Inglés", "martes": "Ciencias", "miercoles": "Historia", "jueves": "Inglés", "viernes": "Matemáticas"},
                {"horas": "11:00-11:45", "lunes": "Educación Física", "martes": "Arte", "miercoles": "Música", "jueves": "Computación", "viernes": "Educación Cívica"},
                {"horas": "11:45-12:30", "lunes": "Computación", "martes": "Educación Física", "miercoles": "Arte", "jueves": "Música", "viernes": "Computación"},
                {"horas": "12:30-13:15", "lunes": "Religión", "martes": "Computación", "miercoles": "Educación Física", "jueves": "Arte", "viernes": "Música"}
            ]

    root = ctk.CTk()
    root.title("RegistroDoc Pro v3.0")
    # ── VENTANA FIJA 1280×720 ──────────────────────────────────────────
    root.geometry("1280x720")
    root.minsize(1280, 720)
    root.maxsize(1280, 720)
    root.resizable(False, False)

    DashboardFrame(root, MockEngine()).pack(fill="both", expand=True)
    root.mainloop()
