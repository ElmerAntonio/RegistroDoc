"""
RegistroDoc Pro — Generador de Actas de Reuniones
Formatos para:
  1. Reunión de Docentes
  2. Reunión con Padres de Familia / Citación
100% offline — genera documentos Word (.docx)
"""
import os
import datetime
from config import BASE_DIR

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


def set_col_widths(table, widths_inches):
    """Fija el ancho de cada columna. widths_inches = lista de anchos en pulgadas."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    total = sum(int(w * 1440) for w in widths_inches)
    tblW.set(qn('w:w'), str(total))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(widths_inches[i] * 1440)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)


def agregar_numero_pagina(section):
    """Agrega 'Página X de Y' al pie de página derecho."""
    footer = section.footer
    # Limpiar footer existente si tiene contenido
    for p in footer.paragraphs:
        p.clear()
    
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = p.add_run("Página ")
    run.font.size = Pt(8)
    run.font.name = "Arial"
    
    # Campo PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run2 = p.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    
    run3 = p.add_run(" de ")
    run3.font.size = Pt(8)
    run3.font.name = "Arial"
    
    # Campo NUMPAGES
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    run4 = p.add_run()
    run4._r.append(fldChar3)
    run4._r.append(instrText2)
    run4._r.append(fldChar4)


def agregar_linea_color(doc, color_hex="1E3A8A"):
    """Agrega una línea horizontal de color institucional."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _estilo_celda(cell, texto, negrita=False, tam=10, centrar=False, color_fondo=None):
    """Aplica formato estándar a una celda de tabla."""
    cell.text = ""
    p = cell.paragraphs[0]
    if centrar:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color_fondo:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_fondo)
        cell._tc.get_or_add_tcPr().append(shd)
    run = p.add_run(texto)
    run.font.size = Pt(tam)
    run.font.name = "Arial"
    run.bold = negrita


def generar_acta_docentes(datos):
    """
    Genera un acta de reunión de docentes en formato Word.
    datos: dict con claves:
        - escuela, director, fecha, hora_inicio, hora_fin
        - lugar, asistentes (list of str), temas (list of str)
        - acuerdos (list of str), observaciones (str)
    """
    if not DOCX_OK:
        return None

    from rdsecurity import cargar_config_segura
    from utils.footer_utils import add_header_with_logo, add_footer_with_logo, get_school_logo_path, abrir_documento

    cfg = cargar_config_segura({})
    escuela_nombre = cfg.get("escuela_nombre", datos.get("escuela", "Centro Educativo"))
    escuela_region = cfg.get("escuela_region", "Comarca Ngäbe Buglé")
    ano_lectivo = cfg.get("ano_lectivo", datos.get("ano", "2026"))
    docente_nombre = cfg.get("docente_nombre", "Docente")

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Logo en encabezado
    logo_path = get_school_logo_path()
    if logo_path:
        add_header_with_logo(doc, logo_path)

    # Pie de página con frase y logo
    add_footer_with_logo(
        doc,
        logo_path or "",
        "RegistroDoc Pro",
        fecha=datos.get("fecha", "")
    )

    # Banda de color superior
    p_banda = doc.add_paragraph()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), '1E3A8A')
    p_banda._p.get_or_add_pPr().append(shading)
    r = p_banda.add_run(f"  MINISTERIO DE EDUCACIÓN — {escuela_nombre.upper()}")
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.name = "Arial"

    # Título del acta
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("ACTA DE REUNIÓN DE DOCENTES")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Año Lectivo {ano_lectivo}")
    run.font.size = Pt(11)
    run.font.name = "Arial"

    # Línea separadora azul
    agregar_linea_color(doc, "1E3A8A")
    doc.add_paragraph()

    # Datos generales
    tabla_info = doc.add_table(rows=4, cols=4)
    tabla_info.style = 'Table Grid'
    set_col_widths(tabla_info, [1.2, 2.0, 1.2, 2.0])

    _estilo_celda(tabla_info.cell(0, 0), "Fecha:", negrita=True)
    _estilo_celda(tabla_info.cell(0, 1), datos.get("fecha", ""))
    _estilo_celda(tabla_info.cell(0, 2), "Lugar:", negrita=True)
    _estilo_celda(tabla_info.cell(0, 3), datos.get("lugar", "Sala de reuniones"))

    _estilo_celda(tabla_info.cell(1, 0), "Hora Inicio:", negrita=True)
    _estilo_celda(tabla_info.cell(1, 1), datos.get("hora_inicio", ""))
    _estilo_celda(tabla_info.cell(1, 2), "Hora Fin:", negrita=True)
    _estilo_celda(tabla_info.cell(1, 3), datos.get("hora_fin", ""))

    _estilo_celda(tabla_info.cell(2, 0), "Director(a):", negrita=True)
    _estilo_celda(tabla_info.cell(2, 1), datos.get("director", ""))
    _estilo_celda(tabla_info.cell(2, 2), "Convocante:", negrita=True)
    _estilo_celda(tabla_info.cell(2, 3), datos.get("convocante", "Dirección"))

    _estilo_celda(tabla_info.cell(3, 0), "Acta N°:", negrita=True)
    _estilo_celda(tabla_info.cell(3, 1), datos.get("numero_acta", "001"))
    _estilo_celda(tabla_info.cell(3, 2), "Año Lectivo:", negrita=True)
    _estilo_celda(tabla_info.cell(3, 3), ano_lectivo)

    doc.add_paragraph("")

    # Asistentes
    p = doc.add_paragraph()
    run = p.add_run("ASISTENTES")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)

    asistentes = datos.get("asistentes", [])
    if asistentes:
        tabla_asist = doc.add_table(rows=len(asistentes) + 1, cols=4)
        tabla_asist.style = 'Table Grid'
        set_col_widths(tabla_asist, [0.4, 2.8, 1.5, 1.7])

        encabezados = ["N°", "Nombre del Docente", "Cargo", "Firma"]
        for i, enc in enumerate(encabezados):
            _estilo_celda(tabla_asist.cell(0, i), enc, negrita=True, centrar=True, color_fondo="1E3A8A")
            for p_run in tabla_asist.cell(0, i).paragraphs[0].runs:
                p_run.font.color.rgb = RGBColor(255, 255, 255)

        for idx, asist in enumerate(asistentes):
            color_fila = "F9FAFB" if idx % 2 == 0 else "FFFFFF"
            _estilo_celda(tabla_asist.cell(idx + 1, 0), str(idx + 1), centrar=True, color_fondo=color_fila)
            _estilo_celda(tabla_asist.cell(idx + 1, 1), asist, color_fondo=color_fila)
            _estilo_celda(tabla_asist.cell(idx + 1, 2), "Docente", color_fondo=color_fila)
            _estilo_celda(tabla_asist.cell(idx + 1, 3), "", color_fondo=color_fila)

    doc.add_paragraph("")

    # Agenda / Temas tratados
    p = doc.add_paragraph()
    run = p.add_run("AGENDA / TEMAS TRATADOS")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)

    temas = datos.get("temas", [])
    for i, tema in enumerate(temas, 1):
        p_t = doc.add_paragraph(f"{i}. {tema}", style="List Number")
        for r_t in p_t.runs:
            r_t.font.name = "Arial"
            r_t.font.size = Pt(10)

    doc.add_paragraph("")

    # Acuerdos
    p = doc.add_paragraph()
    run = p.add_run("ACUERDOS Y COMPROMISOS")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)

    acuerdos = datos.get("acuerdos", [])
    for i, acuerdo in enumerate(acuerdos, 1):
        p_a = doc.add_paragraph(f"{i}. {acuerdo}", style="List Number")
        for r_a in p_a.runs:
            r_a.font.name = "Arial"
            r_a.font.size = Pt(10)

    doc.add_paragraph("")

    # Observaciones
    p = doc.add_paragraph()
    run = p.add_run("OBSERVACIONES")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)

    p_obs = doc.add_paragraph(datos.get("observaciones", "Sin observaciones adicionales."))
    p_obs.runs[0].font.name = "Arial"
    p_obs.runs[0].font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Línea separadora
    agregar_linea_color(doc, "1E3A8A")

    tabla_firmas = doc.add_table(rows=3, cols=3)
    set_col_widths(tabla_firmas, [2.2, 2.2, 2.2])

    _estilo_celda(tabla_firmas.cell(0, 0), "_" * 28, centrar=True)
    _estilo_celda(tabla_firmas.cell(0, 1), "_" * 28, centrar=True)
    _estilo_celda(tabla_firmas.cell(0, 2), "_" * 28, centrar=True)

    _estilo_celda(tabla_firmas.cell(1, 0), datos.get("director", "Director(a)"), centrar=True, negrita=True, tam=10)
    _estilo_celda(tabla_firmas.cell(1, 1), datos.get("convocante", "Convocante"), centrar=True, negrita=True, tam=10)
    _estilo_celda(tabla_firmas.cell(1, 2), "Secretario(a)", centrar=True, negrita=True, tam=10)

    _estilo_celda(tabla_firmas.cell(2, 0), "Director(a)", centrar=True, tam=9)
    _estilo_celda(tabla_firmas.cell(2, 1), "Docente Convocante", centrar=True, tam=9)
    _estilo_celda(tabla_firmas.cell(2, 2), "Secretario de Acta", centrar=True, tam=9)

    # Agregar número de página
    for section in doc.sections:
        agregar_numero_pagina(section)

    # Guardar
    from rdsecurity import cargar_config_segura
    cfg = cargar_config_segura({})
    ruta_base = cfg.get("ruta_exportacion")
    if not ruta_base:
        ruta_base = os.path.join(os.path.expanduser("~"), "Documents", "RegistroDoc")
    carpeta = os.path.join(ruta_base, "Reuniones")
    os.makedirs(carpeta, exist_ok=True)

    fecha_str = datos.get("fecha", datetime.datetime.now().strftime("%d-%m-%Y"))
    nombre = f"Acta_Reunion_Docentes_{fecha_str}.docx"
    ruta = os.path.join(carpeta, nombre)
    doc.save(ruta)
    
    abrir_documento(ruta)
    return ruta


def generar_acta_padres(datos):
    """
    Genera un acta de reunión con padres de familia.
    datos: dict con claves:
        - escuela, docente, fecha, hora
        - alumno, grado, acudiente, parentesco, telefono
        - motivo, descripcion, acuerdos_padre, acuerdos_alumno
    """
    if not DOCX_OK:
        return None

    from rdsecurity import cargar_config_segura
    from utils.footer_utils import add_header_with_logo, add_footer_with_logo, get_school_logo_path, abrir_documento

    cfg = cargar_config_segura({})
    escuela_nombre = cfg.get("escuela_nombre", datos.get("escuela", "Centro Educativo"))
    escuela_region = cfg.get("escuela_region", "Comarca Ngäbe Buglé")
    docente_nombre = cfg.get("docente_nombre", datos.get("docente", "Docente"))
    docente_cedula = cfg.get("docente_cedula", "")

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Logo en encabezado
    logo_path = get_school_logo_path()
    if logo_path:
        add_header_with_logo(doc, logo_path)

    # Pie de página con frase y logo
    add_footer_with_logo(
        doc,
        logo_path or "",
        "RegistroDoc Pro",
        fecha=datos.get("fecha", "")
    )

    # Banda de color superior
    p_banda = doc.add_paragraph()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), '1E3A8A')
    p_banda._p.get_or_add_pPr().append(shading)
    r = p_banda.add_run(f"  MINISTERIO DE EDUCACIÓN — {escuela_nombre.upper()}")
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.name = "Arial"

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("ACTA DE REUNIÓN CON PADRE/MADRE DE FAMILIA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Dirección Regional — {escuela_region}")
    run.font.size = Pt(10)
    run.font.name = "Arial"

    # Línea separadora azul
    agregar_linea_color(doc, "1E3A8A")
    doc.add_paragraph()

    # Datos del alumno
    tabla = doc.add_table(rows=5, cols=4)
    tabla.style = 'Table Grid'
    set_col_widths(tabla, [1.1, 1.9, 1.1, 1.9])

    _estilo_celda(tabla.cell(0, 0), "Fecha:", negrita=True)
    _estilo_celda(tabla.cell(0, 1), datos.get("fecha", ""))
    _estilo_celda(tabla.cell(0, 2), "Hora:", negrita=True)
    _estilo_celda(tabla.cell(0, 3), datos.get("hora", ""))

    _estilo_celda(tabla.cell(1, 0), "Alumno(a):", negrita=True)
    _estilo_celda(tabla.cell(1, 1), datos.get("alumno", ""))
    _estilo_celda(tabla.cell(1, 2), "Grado:", negrita=True)
    _estilo_celda(tabla.cell(1, 3), datos.get("grado", ""))

    _estilo_celda(tabla.cell(2, 0), "Acudiente:", negrita=True)
    _estilo_celda(tabla.cell(2, 1), datos.get("acudiente", ""))
    _estilo_celda(tabla.cell(2, 2), "Parentesco:", negrita=True)
    _estilo_celda(tabla.cell(2, 3), datos.get("parentesco", ""))

    _estilo_celda(tabla.cell(3, 0), "Teléfono:", negrita=True)
    _estilo_celda(tabla.cell(3, 1), datos.get("telefono", ""))
    _estilo_celda(tabla.cell(3, 2), "Docente:", negrita=True)
    _estilo_celda(tabla.cell(3, 3), docente_nombre)

    _estilo_celda(tabla.cell(4, 0), "Motivo:", negrita=True)
    motivo = datos.get("motivo", "Citación general")
    tabla.cell(4, 1).merge(tabla.cell(4, 3))
    _estilo_celda(tabla.cell(4, 1), motivo)

    doc.add_paragraph("")

    # Descripción de la situación
    p = doc.add_paragraph()
    run = p.add_run("DESCRIPCIÓN DE LA SITUACIÓN")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)
    
    p_desc = doc.add_paragraph(datos.get("descripcion", ""))
    p_desc.runs[0].font.name = "Arial"
    p_desc.runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    # Acuerdos con el padre
    p = doc.add_paragraph()
    run = p.add_run("ACUERDOS CON EL PADRE/MADRE")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)
    
    p_acpad = doc.add_paragraph(datos.get("acuerdos_padre", ""))
    p_acpad.runs[0].font.name = "Arial"
    p_acpad.runs[0].font.size = Pt(10)

    # Compromisos del alumno
    p = doc.add_paragraph()
    run = p.add_run("COMPROMISOS DEL ALUMNO(A)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)
    
    p_acalu = doc.add_paragraph(datos.get("acuerdos_alumno", ""))
    p_acalu.runs[0].font.name = "Arial"
    p_acalu.runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    # Seguimiento y Próxima Cita
    p = doc.add_paragraph()
    run = p.add_run("SEGUIMIENTO Y PRÓXIMA CITA")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(30, 58, 138)

    tabla_seg = doc.add_table(rows=2, cols=2)
    tabla_seg.style = 'Table Grid'
    set_col_widths(tabla_seg, [3.3, 3.1])

    _estilo_celda(tabla_seg.cell(0, 0), "Próxima reunión programada para:", negrita=True, tam=10)
    _estilo_celda(tabla_seg.cell(0, 1), "__________________________", tam=10)
    _estilo_celda(tabla_seg.cell(1, 0), "Observaciones adicionales:", negrita=True, tam=10)
    _estilo_celda(tabla_seg.cell(1, 1), datos.get("observaciones_seguimiento", ""), tam=10)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Firmas (3 columnas)
    tabla_firmas = doc.add_table(rows=3, cols=3)
    set_col_widths(tabla_firmas, [2.2, 2.2, 2.2])

    _estilo_celda(tabla_firmas.cell(0, 0), "_" * 26, centrar=True)
    _estilo_celda(tabla_firmas.cell(0, 1), "_" * 26, centrar=True)
    _estilo_celda(tabla_firmas.cell(0, 2), "_" * 26, centrar=True)

    _estilo_celda(tabla_firmas.cell(1, 0), docente_nombre, centrar=True, negrita=True, tam=10)
    _estilo_celda(tabla_firmas.cell(1, 1), datos.get("acudiente", "Acudiente"), centrar=True, negrita=True, tam=10)
    _estilo_celda(tabla_firmas.cell(1, 2), datos.get("alumno", "Alumno(a)"), centrar=True, negrita=True, tam=10)

    _estilo_celda(tabla_firmas.cell(2, 0), f"Cédula: {docente_cedula or '_________'}", centrar=True, tam=9)
    _estilo_celda(tabla_firmas.cell(2, 1), f"Tel: {datos.get('telefono', '_________')}", centrar=True, tam=9)
    _estilo_celda(tabla_firmas.cell(2, 2), f"Grado: {datos.get('grado', '')}", centrar=True, tam=9)

    # Agregar número de página
    for section in doc.sections:
        agregar_numero_pagina(section)

    # Guardar
    from rdsecurity import cargar_config_segura
    cfg = cargar_config_segura({})
    ruta_base = cfg.get("ruta_exportacion")
    if not ruta_base:
        ruta_base = os.path.join(os.path.expanduser("~"), "Documents", "RegistroDoc")
    carpeta = os.path.join(ruta_base, "Reuniones")
    os.makedirs(carpeta, exist_ok=True)

    alumno = datos.get("alumno", "Alumno").replace(" ", "_")
    fecha_str = datos.get("fecha", datetime.datetime.now().strftime("%d-%m-%Y"))
    nombre = f"Acta_Padres_{alumno}_{fecha_str}.docx"
    ruta = os.path.join(carpeta, nombre)
    doc.save(ruta)
    
    abrir_documento(ruta)
    return ruta
