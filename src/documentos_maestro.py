"""
RegistroDoc Pro — Módulo de Generación de Documentos Maestro Word (.docx)
100% offline y adaptado a los estándares MEDUCA.
"""
import os
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from config import BASE_DIR
from utils.footer_utils import add_footer_with_logo, add_header_with_logo, get_school_logo_path

# --- UTILIDADES XML PARA DIAPOSITIVAS Y ESTILOS ---

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece los márgenes internos (padding) de una celda en dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    """Establece el color de fondo de una celda."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_width(cell, width_in_inches):
    """Establece un ancho fijo estricto en la celda."""
    cell.width = Inches(width_in_inches)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_in_inches * 1440))) # 1440 dxa por pulgada
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def add_page_number_to_footer(run):
    """Agrega el campo dinámico de número de página PAGE."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_total_pages_to_footer(run):
    """Agrega el campo dinámico de total de páginas NUMPAGES."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def setup_footer_page_numbers(doc):
    """Configura el pie de página con numeración 'Página X de Y'."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.text = ""
        run = p.add_run("Página ")
        run.font.size = Pt(8.5)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        run_num = p.add_run()
        run_num.font.size = Pt(8.5)
        run_num.font.name = "Arial"
        run_num.font.color.rgb = RGBColor(100, 100, 100)
        add_page_number_to_footer(run_num)
        
        run_of = p.add_run(" de ")
        run_of.font.size = Pt(8.5)
        run_of.font.name = "Arial"
        run_of.font.color.rgb = RGBColor(100, 100, 100)
        
        run_tot = p.add_run()
        run_tot.font.size = Pt(8.5)
        run_tot.font.name = "Arial"
        run_tot.font.color.rgb = RGBColor(100, 100, 100)
        add_total_pages_to_footer(run_tot)

def _crear_parrafo_celda(cell, texto, negrita=False, tam=10, centrar=False, color_rgb=None):
    """Ayudante para agregar texto formateado a una celda."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrar else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(texto)
    run.font.size = Pt(tam)
    run.font.name = "Arial"
    run.bold = negrita
    if color_rgb:
        run.font.color.rgb = color_rgb
    return p


# --- GENERADORES DE DOCUMENTOS ---

def generar_nota_acudiente(datos):
    """
    Genera una Nota Rápida para Acudientes en formato de media página.
    datos: dict con claves:
        - escuela_nombre, alumno, grado, docente_nombre, fecha, motivo, mensaje
    """
    doc = Document()
    
    # Configuración de márgenes estrechos para media página
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.page_width = Inches(8.5)
        section.page_height = Inches(5.5) # Media página Carta horizontal
        
    logo_path = get_school_logo_path()
    if logo_path:
        add_header_with_logo(doc, logo_path, text=datos.get("escuela_nombre", "Ministerio de Educación"))
        
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titulo.add_run("COMUNICADO AL ACUDIENTE")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(30, 58, 138) # Azul MEDUCA
    
    doc.add_paragraph("")
    
    # Tabla de metadatos
    tabla = doc.add_table(rows=3, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    anchos = [1.5, 5.0]
    
    meta_filas = [
        ("Fecha:", datos.get("fecha", datetime.datetime.now().strftime("%d-%m-%Y"))),
        ("Estudiante:", f"{datos.get('alumno', '')} ({datos.get('grado', '')})"),
        ("Motivo:", datos.get("motivo", "Comunicado Oficial")),
    ]
    
    for i, (label, val) in enumerate(meta_filas):
        c0 = tabla.cell(i, 0)
        c1 = tabla.cell(i, 1)
        set_cell_width(c0, anchos[0])
        set_cell_width(c1, anchos[1])
        set_cell_margins(c0, 40, 40, 60, 60)
        set_cell_margins(c1, 40, 40, 60, 60)
        
        _crear_parrafo_celda(c0, label, negrita=True, tam=9.5)
        _crear_parrafo_celda(c1, val, tam=9.5)
        
    doc.add_paragraph("")
    
    # Mensaje principal
    p_msg = doc.add_paragraph()
    run_msg = p_msg.add_run(datos.get("mensaje", ""))
    run_msg.font.size = Pt(10)
    run_msg.font.name = "Arial"
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Firmas
    tabla_firmas = doc.add_table(rows=2, cols=2)
    tabla_firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    c_doc_firma = tabla_firmas.cell(0, 0)
    c_acu_firma = tabla_firmas.cell(0, 1)
    set_cell_width(c_doc_firma, 3.25)
    set_cell_width(c_acu_firma, 3.25)
    
    _crear_parrafo_celda(c_doc_firma, "_________________________________", centrar=True, tam=9.5)
    _crear_parrafo_celda(c_acu_firma, "_________________________________", centrar=True, tam=9.5)
    
    c_doc_lbl = tabla_firmas.cell(1, 0)
    c_acu_lbl = tabla_firmas.cell(1, 1)
    set_cell_width(c_doc_lbl, 3.25)
    set_cell_width(c_acu_lbl, 3.25)
    
    from rdsecurity import cargar_config_segura
    cfg = cargar_config_segura({})
    docente = datos.get('docente_nombre') or cfg.get("docente_nombre", "")
    cedula = cfg.get("docente_cedula", "")

    _crear_parrafo_celda(c_doc_lbl, f"Firma del Docente\n{docente}\nCédula: {cedula}", centrar=True, tam=9)
    _crear_parrafo_celda(c_acu_lbl, "Firma del Acudiente\nCédula: _______________________", centrar=True, tam=9)

    # Carpeta
    ruta_base = cfg.get("ruta_exportacion")
    if not ruta_base:
        ruta_base = os.path.join(os.path.expanduser("~"), "Documents", "RegistroDoc")
    carpeta = os.path.join(ruta_base, "Expedientes_Estudiantes")
    os.makedirs(carpeta, exist_ok=True)
        
    nombre_seguro = datos.get("alumno", "Comunicado").replace(" ", "_")
    fecha_str = datos.get("fecha", datetime.datetime.now().strftime("%d-%m-%Y"))
    ruta = os.path.join(carpeta, f"Nota_Acudiente_{nombre_seguro}_{fecha_str}.docx")
    doc.save(ruta)
    return ruta


def generar_informe_calificaciones(datos):
    """
    Genera un informe individual de calificaciones en formato Word formal (boleta/ficha).
    datos: dict con las siguientes claves:
        - alumno, cedula, grado, docente_nombre, escuela_nombre, ano_lectivo, fecha, estado_final
        - trimestres: dict con claves "T1", "T2", "T3", y valores:
             {"materias": [{"nombre": "Español", "nota": 4.5, "estado": "APROBADO"}], "promedio": 4.5}
    """
    doc = Document()
    
    from rdsecurity import cargar_config_segura
    cfg = cargar_config_segura({})
    docente = datos.get("docente_nombre") or cfg.get("docente_nombre", "")
    cedula_docente = cfg.get("docente_cedula", "")
    docente_str = f"{docente} (Céd. {cedula_docente})" if cedula_docente else docente

    # Márgenes de 1 pulgada
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        
    logo_path = get_school_logo_path()
    if logo_path:
        add_header_with_logo(doc, logo_path, text=datos.get("escuela_nombre", "Centro Educativo"))
        
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titulo.add_run("INFORME DE RENDIMIENTO ACADÉMICO")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(30, 58, 138)
    
    doc.add_paragraph("")
    
    # Tabla de Datos Generales (2x4)
    tabla_datos = doc.add_table(rows=4, cols=4)
    tabla_datos.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    col_widths = [1.2, 2.05, 1.2, 2.05] # Total: 6.5 pulgadas de ancho disponible
    
    datos_campos = [
        [("Estudiante:", datos.get("alumno", "")), ("Cédula:", datos.get("cedula", "N/A"))],
        [("Grado/Grupo:", datos.get("grado", "")), ("Año Lectivo:", datos.get("ano_lectivo", ""))],
        [("Docente:", docente_str), ("Fecha Emisión:", datos.get("fecha", ""))],
        [("Centro Educativo:", datos.get("escuela_nombre", "")), ("Estado Final:", datos.get("estado_final", "EN PROCESO"))]
    ]
    
    for r_idx, fila in enumerate(datos_campos):
        col_pair_idx = 0
        for label, val in fila:
            c0 = tabla_datos.cell(r_idx, col_pair_idx * 2)
            c1 = tabla_datos.cell(r_idx, col_pair_idx * 2 + 1)
            set_cell_width(c0, col_widths[col_pair_idx * 2])
            set_cell_width(c1, col_widths[col_pair_idx * 2 + 1])
            set_cell_margins(c0, 50, 50, 80, 80)
            set_cell_margins(c1, 50, 50, 80, 80)
            
            _crear_parrafo_celda(c0, label, negrita=True, tam=9.5)
            # Colorear estado final
            if label == "Estado Final:":
                color_fin = RGBColor(16, 185, 129) if "APROBADO" in val.upper() else RGBColor(239, 68, 68)
                _crear_parrafo_celda(c1, val, negrita=True, tam=9.5, color_rgb=color_fin)
            else:
                _crear_parrafo_celda(c1, val, tam=9.5)
            col_pair_idx += 1
            
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Recopilar materias únicas
    materias_unidades = {}
    for tk in ["T1", "T2", "T3"]:
        t_data = datos.get("trimestres", {}).get(tk, {})
        if t_data:
            for mat in t_data.get("materias", []):
                materias_unidades[mat["nombre"]] = True
                
    mats_lista = sorted(list(materias_unidades.keys()))
    
    if mats_lista:
        # Tabla de calificaciones consolidada
        # Columnas: Materia | Trimestre 1 | Trimestre 2 | Trimestre 3 | Promedio Final
        tabla_notas = doc.add_table(rows=len(mats_lista) + 2, cols=5)
        tabla_notas.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        tab_widths = [2.2, 0.9, 0.9, 0.9, 1.6] # Total: 6.5
        
        # Encabezado principal
        headers = ["Asignatura / Materia", "Trimestre 1", "Trimestre 2", "Trimestre 3", "Promedio Final"]
        for col_idx, text in enumerate(headers):
            cell = tabla_notas.cell(0, col_idx)
            set_cell_width(cell, tab_widths[col_idx])
            set_cell_background(cell, "1E3A88") # Azul MEDUCA oscuro
            set_cell_margins(cell, 100, 100, 100, 100)
            _crear_parrafo_celda(cell, text, negrita=True, tam=10, centrar=(col_idx > 0), color_rgb=RGBColor(255, 255, 255))
            
        # Filas de materias
        for row_idx, mat_name in enumerate(mats_lista, 1):
            cell_name = tabla_notas.cell(row_idx, 0)
            set_cell_width(cell_name, tab_widths[0])
            set_cell_margins(cell_name, 60, 60, 100, 100)
            _crear_parrafo_celda(cell_name, mat_name, negrita=True, tam=9.5)
            
            suma_notas = 0.0
            conteo_notas = 0
            
            for col_idx, tk in enumerate(["T1", "T2", "T3"], 1):
                cell_nota = tabla_notas.cell(row_idx, col_idx)
                set_cell_width(cell_nota, tab_widths[col_idx])
                set_cell_margins(cell_nota, 60, 60, 60, 60)
                
                # Buscar nota de esta materia en este trimestre
                t_mats = datos.get("trimestres", {}).get(tk, {}).get("materias", [])
                match_mat = next((m for m in t_mats if m["nombre"] == mat_name), None)
                
                if match_mat:
                    nota_val = match_mat["nota"]
                    nota_str = f"{nota_val:.1f}"
                    color_nota = RGBColor(16, 185, 129) if nota_val >= 3.0 else RGBColor(239, 68, 68)
                    _crear_parrafo_celda(cell_nota, nota_str, tam=9.5, centrar=True, color_rgb=color_nota)
                    suma_notas += nota_val
                    conteo_notas += 1
                else:
                    _crear_parrafo_celda(cell_nota, "-", tam=9.5, centrar=True)
                    
            # Promedio Final
            cell_prom = tabla_notas.cell(row_idx, 4)
            set_cell_width(cell_prom, tab_widths[4])
            set_cell_margins(cell_prom, 60, 60, 60, 60)
            if conteo_notas > 0:
                prom_val = round(suma_notas / conteo_notas, 1)
                color_nota = RGBColor(16, 185, 129) if prom_val >= 3.0 else RGBColor(239, 68, 68)
                _crear_parrafo_celda(cell_prom, f"{prom_val:.1f}", negrita=True, tam=9.5, centrar=True, color_rgb=color_nota)
            else:
                _crear_parrafo_celda(cell_prom, "-", tam=9.5, centrar=True)
                
        # Fila de promedio general del salón
        prom_row_idx = len(mats_lista) + 1
        cell_lbl = tabla_notas.cell(prom_row_idx, 0)
        set_cell_width(cell_lbl, tab_widths[0])
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        _crear_parrafo_celda(cell_lbl, "PROMEDIO GENERAL", negrita=True, tam=10)
        set_cell_background(cell_lbl, "E2E8F0")
        
        for col_idx, tk in enumerate(["T1", "T2", "T3"], 1):
            cell_t_prom = tabla_notas.cell(prom_row_idx, col_idx)
            set_cell_width(cell_t_prom, tab_widths[col_idx])
            set_cell_margins(cell_t_prom, 80, 80, 60, 60)
            set_cell_background(cell_t_prom, "E2E8F0")
            
            prom_trim = datos.get("trimestres", {}).get(tk, {}).get("promedio", 0.0)
            if prom_trim > 0.0:
                color_nota = RGBColor(16, 185, 129) if prom_trim >= 3.0 else RGBColor(239, 68, 68)
                _crear_parrafo_celda(cell_t_prom, f"{prom_trim:.1f}", negrita=True, tam=9.5, centrar=True, color_rgb=color_nota)
            else:
                _crear_parrafo_celda(cell_t_prom, "-", tam=9.5, centrar=True)
                
        # Celda promedio final general
        cell_f_prom = tabla_notas.cell(prom_row_idx, 4)
        set_cell_width(cell_f_prom, tab_widths[4])
        set_cell_margins(cell_f_prom, 80, 80, 60, 60)
        set_cell_background(cell_f_prom, "CBD5E1")
        
        proms_validos = [datos.get("trimestres", {}).get(tk, {}).get("promedio", 0.0) for tk in ["T1", "T2", "T3"] if datos.get("trimestres", {}).get(tk, {}).get("promedio", 0.0) > 0.0]
        if proms_validos:
            prom_final_gen = round(sum(proms_validos) / len(proms_validos), 1)
            color_nota = RGBColor(16, 185, 129) if prom_final_gen >= 3.0 else RGBColor(239, 68, 68)
            _crear_parrafo_celda(cell_f_prom, f"{prom_final_gen:.1f}", negrita=True, tam=10, centrar=True, color_rgb=color_nota)
        else:
            _crear_parrafo_celda(cell_f_prom, "-", tam=9.5, centrar=True)
            
    else:
        doc.add_paragraph("No hay calificaciones registradas para este estudiante.")
        
    doc.add_paragraph("")
    
    # Observación General / Sugerencias de Rendimiento (Tarea 23)
    obs_val = datos.get("observacion_general", "").strip()
    if obs_val:
        p_obs_title = doc.add_paragraph()
        run_obs_t = p_obs_title.add_run("Observaciones y Recomendaciones de Rendimiento:")
        run_obs_t.bold = True
        run_obs_t.font.size = Pt(11)
        run_obs_t.font.name = "Arial"
        run_obs_t.font.color.rgb = RGBColor(30, 58, 138)
        
        p_obs_text = doc.add_paragraph()
        run_obs_val = p_obs_text.add_run(obs_val)
        run_obs_val.font.size = Pt(10)
        run_obs_val.font.name = "Arial"
        
        doc.add_paragraph("")
    else:
        doc.add_paragraph("")
        doc.add_paragraph("")
    
    # Firmas
    tabla_firmas = doc.add_table(rows=2, cols=2)
    tabla_firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    c_doc_firma = tabla_firmas.cell(0, 0)
    c_dir_firma = tabla_firmas.cell(0, 1)
    set_cell_width(c_doc_firma, 3.25)
    set_cell_width(c_dir_firma, 3.25)
    
    _crear_parrafo_celda(c_doc_firma, "_________________________________", centrar=True, tam=10)
    _crear_parrafo_celda(c_dir_firma, "_________________________________", centrar=True, tam=10)
    
    c_doc_lbl = tabla_firmas.cell(1, 0)
    c_dir_lbl = tabla_firmas.cell(1, 1)
    set_cell_width(c_doc_lbl, 3.25)
    set_cell_width(c_dir_lbl, 3.25)
    
    _crear_parrafo_celda(c_doc_lbl, f"Firma del Docente\n{docente}\nCédula: {cedula_docente}", centrar=True, tam=9.5)
    _crear_parrafo_celda(c_dir_lbl, "Firma de la Dirección\nCentro Educativo", centrar=True, tam=9.5)

    # Footer y numeración
    try:
        add_footer_with_logo(doc, logo_path, "RegistroDoc Pro v3.0", "Instruye al niño en su camino, y aun cuando fuere viejo no se apartará de él. - Proverbios 22:6")
    except Exception:
        pass
        
    setup_footer_page_numbers(doc)

    # Carpeta
    ruta_base = cfg.get("ruta_exportacion")
    if not ruta_base:
        ruta_base = os.path.join(os.path.expanduser("~"), "Documents", "RegistroDoc")
    carpeta = os.path.join(ruta_base, "Expedientes_Estudiantes")
    os.makedirs(carpeta, exist_ok=True)
        
    nombre_seguro = datos.get("alumno", "Estudiante").replace(" ", "_")
    ruta = os.path.join(carpeta, f"Informe_Calificaciones_{nombre_seguro}.docx")
    doc.save(ruta)
    return ruta


def generar_lista_clase(datos):
    """
    Genera una lista de clase en cuadrícula (plantilla de firmas/asistencia).
    datos: dict con claves:
        - grado, docente_nombre, escuela_nombre, ano_lectivo, estudiantes (lista de dicts con 'id', 'nombre', 'cedula')
    """
    from rdsecurity import cargar_config_segura
    cfg = cargar_config_segura({})
    docente = datos.get("docente_nombre") or cfg.get("docente_nombre", "")
    cedula_docente = cfg.get("docente_cedula", "")

    doc = Document()
    
    # Márgenes de 0.6 pulgadas
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        
    logo_path = get_school_logo_path()
    if logo_path:
        add_header_with_logo(doc, logo_path, text=datos.get("escuela_nombre", "Centro Educativo"))
        
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titulo.add_run(f"LISTA DE CLASE — GRADO/GRUPO: {datos.get('grado', '')}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(30, 58, 138)
    
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = p_info.add_run(f"Docente: {docente} (Cédula: {cedula_docente})   |   Año Lectivo: {datos.get('ano_lectivo', '')}")
    run_info.font.size = Pt(9.5)
    run_info.font.name = "Arial"
    
    doc.add_paragraph("")
    
    estudiantes = datos.get("estudiantes", [])
    if not estudiantes:
        doc.add_paragraph("No hay estudiantes registrados en este grupo.")
        return None
        
    # Anchos de columna: N°(0.4), Estudiante(2.2), Cédula(1.2), 10 columnas vacías de 0.35 para notas/asistencia
    # Total: 0.4 + 2.2 + 1.2 + (0.35 * 10) = 7.3 pulgadas.
    tabla = doc.add_table(rows=len(estudiantes) + 1, cols=13)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    col_widths = [0.4, 2.2, 1.2] + [0.35] * 10
    
    # Encabezado
    headers = ["N°", "Nombre del Estudiante", "Cédula"] + [f"{i}" for i in range(1, 11)]
    for col_idx, text in enumerate(headers):
        cell = tabla.cell(0, col_idx)
        set_cell_width(cell, col_widths[col_idx])
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 60, 60, 40, 40)
        _crear_parrafo_celda(cell, text, negrita=True, tam=9, centrar=(col_idx != 1), color_rgb=RGBColor(255, 255, 255))
        
    # Filas de estudiantes
    for row_idx, est in enumerate(estudiantes, 1):
        # N°
        c_num = tabla.cell(row_idx, 0)
        set_cell_width(c_num, col_widths[0])
        set_cell_margins(c_num, 40, 40, 40, 40)
        _crear_parrafo_celda(c_num, str(row_idx), tam=8.5, centrar=True)
        
        # Nombre
        c_nom = tabla.cell(row_idx, 1)
        set_cell_width(c_nom, col_widths[1])
        set_cell_margins(c_nom, 40, 40, 60, 60)
        _crear_parrafo_celda(c_nom, est.get("nombre", ""), tam=8.5)
        
        # Cédula
        c_ced = tabla.cell(row_idx, 2)
        set_cell_width(c_ced, col_widths[2])
        set_cell_margins(c_ced, 40, 40, 40, 40)
        _crear_parrafo_celda(c_ced, est.get("cedula", "N/A"), tam=8.5, centrar=True)
        
        # Celdas vacías cuadrícula
        for c_idx in range(3, 13):
            c_empty = tabla.cell(row_idx, c_idx)
            set_cell_width(c_empty, col_widths[c_idx])
            set_cell_margins(c_empty, 40, 40, 40, 40)
            _crear_parrafo_celda(c_empty, "", tam=8.5)
            
            # Zebra striping muy sutil para facilitar la lectura
            if row_idx % 2 == 0:
                set_cell_background(c_empty, "F8FAFC")
        
        if row_idx % 2 == 0:
            set_cell_background(c_num, "F8FAFC")
            set_cell_background(c_nom, "F8FAFC")
            set_cell_background(c_ced, "F8FAFC")
            
    doc.add_paragraph("")
    
    # Footer y numeración
    try:
        add_footer_with_logo(doc, logo_path, "RegistroDoc Pro v3.0", "Instruye al niño en su camino, y aun cuando fuere viejo no se apartará de él. - Proverbios 22:6")
    except Exception:
        pass
        
    setup_footer_page_numbers(doc)

    # Carpeta
    ruta_base = cfg.get("ruta_exportacion")
    if not ruta_base:
        ruta_base = os.path.join(os.path.expanduser("~"), "Documents", "RegistroDoc")
    carpeta = os.path.join(ruta_base, "Documentos")
    os.makedirs(carpeta, exist_ok=True)
        
    grado_str = datos.get("grado", "Grado").replace("°", "")
    ruta = os.path.join(carpeta, f"Lista_Clase_{grado_str}.docx")
    doc.save(ruta)
    return ruta


def generar_citacion_txt(datos):
    """
    Genera una citación estructurada en texto plano (.txt) para Notepad (Tarea 20).
    """
    from rdsecurity import cargar_config_segura
    cfg = cargar_config_segura({})
    escuela = datos.get("escuela_nombre", "Ministerio de Educación")
    fecha = datos.get("fecha", datetime.datetime.now().strftime("%d-%m-%Y"))
    docente = datos.get("docente_nombre") or cfg.get("docente_nombre", "Docente")
    cedula_docente = cfg.get("docente_cedula", "")
    ano = datos.get("ano_lectivo", "2026")
    alumno = datos.get("alumno", "Estudiante")
    grado = datos.get("grado", "")
    cedula = datos.get("cedula", "N/A")
    motivo = datos.get("motivo", "Comunicado Oficial")
    descripcion = datos.get("descripcion", "")

    txt = f"""========================================================================
             REPÚBLICA DE PANAMÁ - MINISTERIO DE EDUCACIÓN
                      CITACIÓN OFICIAL DE PADRES
========================================================================

Fecha: {fecha}
Centro Educativo: {escuela}
Docente: {docente} (Cédula: {cedula_docente})
Año Lectivo: {ano}

Estimado(a) Acudiente de: {alumno}
Grado: {grado}
Cédula del Estudiante: {cedula}

Por medio de la presente, se le cita formalmente a una reunión individual
con el docente de su acudido(a) para tratar asuntos de carácter oficial.

Detalle de la Citación:
------------------------------------------------------------------------
Motivo: {motivo}
Descripción/Observación: 
{descripcion}
------------------------------------------------------------------------

Lugar: Instalaciones del Plantel Educativo
Fecha de la cita: ___/___/______   Hora: ______:______

Su presencia es de vital importancia para garantizar el seguimiento y
apoyo adecuado en el proceso formativo de su acudido(a).

Agradeciendo su puntual asistencia.

Atentamente,



__________________________________       __________________________________
     Firma del Docente ({docente})            Firma del Acudiente
     Cédula: {cedula_docente}                  Cédula: _____________________

========================================================================
                 "La educación es el camino al éxito"
========================================================================
"""
    ruta_base = cfg.get("ruta_exportacion")
    if not ruta_base:
        ruta_base = os.path.join(os.path.expanduser("~"), "Documents", "RegistroDoc")
    carpeta = os.path.join(ruta_base, "Expedientes_Estudiantes", "Citaciones")
    os.makedirs(carpeta, exist_ok=True)
        
    nombre_seguro = alumno.replace(" ", "_")
    ruta_txt = os.path.join(carpeta, f"Citacion_{nombre_seguro}_{fecha.replace('-', '_')}.txt")
    with open(ruta_txt, "w", encoding="utf-8") as f:
        f.write(txt)
    return ruta_txt


# ══════════════════════════════════════════════════════════════════════════════
#  EXPEDIENTE DE RETIRO
# ══════════════════════════════════════════════════════════════════════════════

def generar_expediente_retiro(datos: dict) -> str:
    """Genera un documento Word con el expediente completo del alumno retirado.

    datos dict keys:
        nombre, cedula, grado, motivo, acudiente,
        docente_nombre, escuela_nombre, ano_lectivo,
        asistencia (dict con pct_asistencia, total_dias, dias_asistidos),
        notas_t1, notas_t2, notas_t3 (dicts materia->promedio)
    Retorna la ruta al .docx generado.
    """
    doc  = Document()
    hoy  = datetime.date.today().strftime("%d/%m/%Y")
    nombre   = datos.get("nombre", "")
    cedula   = datos.get("cedula", "")
    grado    = datos.get("grado", "")
    motivo   = datos.get("motivo", "")
    acudiente= datos.get("acudiente", "")
    docente  = datos.get("docente_nombre", "")
    escuela  = datos.get("escuela_nombre", "")
    ano      = datos.get("ano_lectivo", str(datetime.date.today().year))

    # ── Márgenes ──────────────────────────────────────────────────────────
    from docx.oxml.ns import qn
    sec = doc.sections[0]
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    # ── Encabezado institucional ───────────────────────────────────────────
    try:
        add_header_with_logo(doc, escuela, ano)
    except Exception:
        pass

    # ── Título ────────────────────────────────────────────────────────────
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = p_tit.add_run("EXPEDIENTE DE RETIRO DE ESTUDIANTE")
    run_tit.bold      = True
    run_tit.font.size = Pt(14)
    run_tit.font.color.rgb = RGBColor(0xC0, 0x38, 0x38)
    doc.add_paragraph()

    # ── Datos personales ──────────────────────────────────────────────────
    def _encabezado_seccion(texto):
        p = doc.add_paragraph()
        r = p.add_run(texto)
        r.bold      = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        p.paragraph_format.space_after = Pt(4)
        return p

    _encabezado_seccion("I. DATOS PERSONALES")

    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    pares = [
        ("Nombre completo:", nombre),
        ("Cédula / ID:",     cedula),
        ("Grado:",           grado),
        ("Nombre del acudiente:", acudiente),
    ]
    for i, (lbl, val) in enumerate(pares):
        row = tbl.rows[i]
        row.cells[0].text = lbl
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = val
    doc.add_paragraph()

    # ── Datos del retiro ──────────────────────────────────────────────────
    _encabezado_seccion("II. DATOS DEL RETIRO")
    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.style = "Table Grid"
    tbl2.rows[0].cells[0].text = "Fecha de retiro:"
    tbl2.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    tbl2.rows[0].cells[1].text = hoy
    tbl2.rows[1].cells[0].text = "Motivo:"
    tbl2.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    tbl2.rows[1].cells[1].text = motivo
    doc.add_paragraph()

    # ── Asistencia ────────────────────────────────────────────────────────
    _encabezado_seccion("III. RESUMEN DE ASISTENCIA")
    asis = datos.get("asistencia") or {}
    if isinstance(asis, dict) and asis:
        pct  = asis.get("pct_asistencia", asis.get("porcentaje", 0))
        tot  = asis.get("total_dias", asis.get("total", 0))
        asi  = asis.get("dias_asistidos", asis.get("asistidos", 0))
        p_a  = doc.add_paragraph(f"Días totales: {tot}   |   Días asistidos: {asi}   |   "
                                  f"Porcentaje de asistencia: {pct:.1f}%")
    else:
        p_a = doc.add_paragraph("No hay datos de asistencia registrados.")
    doc.add_paragraph()

    # ── Notas por trimestre ───────────────────────────────────────────────
    _encabezado_seccion("IV. RESUMEN ACADÉMICO POR TRIMESTRE")

    trimestres = [
        ("Trimestre 1", datos.get("notas_t1") or {}),
        ("Trimestre 2", datos.get("notas_t2") or {}),
        ("Trimestre 3", datos.get("notas_t3") or {}),
    ]

    for t_nombre, notas in trimestres:
        if not notas:
            continue
        p_tit_t = doc.add_paragraph()
        r_tit_t = p_tit_t.add_run(t_nombre)
        r_tit_t.bold      = True
        r_tit_t.font.size = Pt(10)

        materias = list(notas.keys())
        if materias:
            tbl_n = doc.add_table(rows=1 + len(materias), cols=2)
            tbl_n.style = "Table Grid"
            tbl_n.rows[0].cells[0].text = "Materia"
            tbl_n.rows[0].cells[1].text = "Promedio"
            for ci in range(2):
                tbl_n.rows[0].cells[ci].paragraphs[0].runs[0].bold = True
                set_cell_background(tbl_n.rows[0].cells[ci], "1E40AF")
                for r in tbl_n.rows[0].cells[ci].paragraphs[0].runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for i, mat in enumerate(materias):
                prom = notas[mat]
                tbl_n.rows[i + 1].cells[0].text = mat
                tbl_n.rows[i + 1].cells[1].text = f"{prom:.1f}" if isinstance(prom, float) else str(prom)
        doc.add_paragraph()

    # ── Firmas ────────────────────────────────────────────────────────────
    _encabezado_seccion("V. FIRMAS")
    doc.add_paragraph()
    from rdsecurity import cargar_config_segura
    cfg = cargar_config_segura({})
    docente = datos.get("docente_nombre") or cfg.get("docente_nombre", "")
    cedula_docente = cfg.get("docente_cedula", "")

    p_f = doc.add_paragraph()
    p_f.add_run(f"Docente a cargo: {docente} (Cédula: {cedula_docente})\n").bold = True
    p_f.add_run("Firma: _________________________________    ")
    p_f.add_run("Fecha: ").bold = True
    p_f.add_run("_____________")

    doc.add_paragraph()
    p_d = doc.add_paragraph()
    p_d.add_run("Director(a): ").bold = True
    p_d.add_run("_________________________________    ")
    p_d.add_run("Fecha: ").bold = True
    p_d.add_run("_____________")

    doc.add_paragraph()
    p_ac = doc.add_paragraph()
    p_ac.add_run("Acudiente: ").bold = True
    p_ac.add_run("_________________________________    ")
    p_ac.add_run("Fecha: ").bold = True
    p_ac.add_run("_____________")

    # ── Pie de página ─────────────────────────────────────────────────────
    try:
        add_footer_with_logo(doc, escuela, ano)
    except Exception:
        pass

    # ── Guardar ───────────────────────────────────────────────────────────
    ruta_base = cfg.get("ruta_exportacion")
    if not ruta_base:
        ruta_base = os.path.join(os.path.expanduser("~"), "Documents", "RegistroDoc")
    carpeta = os.path.join(ruta_base, "Expedientes_Estudiantes", "Retiros")
    os.makedirs(carpeta, exist_ok=True)
    nombre_seguro = nombre.replace(" ", "_").replace(",", "")
    ruta = os.path.join(carpeta, f"Expediente_Retiro_{nombre_seguro}_{hoy.replace('/', '-')}.docx")
    doc.save(ruta)
    return ruta
