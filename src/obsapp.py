import customtkinter as ctk
from tkinter import messagebox
import datetime
import os
import threading
import json
from config import BASE_DIR
from utils.spellchecker import SpellcheckTextbox
from theme import C
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False



class ObservacionesFrame(ctk.CTkFrame):
    def __init__(self, master, engine, **kwargs):
        super().__init__(master, fg_color="transparent", **kwargs)
        self.engine = engine
        self.estudiante_seleccionado = None
        self.grado_actual = ""
        self._debounce_job = None

        # Layout
        self.grid_columnconfigure(0, weight=4)  # Lista Izquierda
        self.grid_columnconfigure(1, weight=6)  # Panel Derecha
        self.grid_rowconfigure(0, weight=1)

        self.crear_panel_izquierdo()
        self.crear_panel_derecho()

        self.al_cambiar_grado(self.combo_grado.get())

    def crear_panel_izquierdo(self):
        frame_izq = ctk.CTkFrame(self, fg_color=C["card"], corner_radius=10)
        frame_izq.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        # Controles
        top = ctk.CTkFrame(frame_izq, fg_color="transparent")
        top.pack(fill="x", padx=15, pady=15)

        ctk.CTkLabel(
            top,
            text="Grado:",
            font=(
                "Segoe UI",
                14,
                "bold")).pack(
            side="left")
        grados = self.engine.obtener_grados_activos() or ["Sin datos"]

        self.combo_grado = ctk.CTkOptionMenu(
            top, values=grados, command=self.al_cambiar_grado, width=80)
        self.combo_grado.pack(side="left", padx=10)

        # Buscador en tiempo real
        self.entry_buscar = ctk.CTkEntry(
            top, placeholder_text="Buscar por nombre...", width=180)
        self.entry_buscar.pack(side="left", padx=5)
        self.entry_buscar.bind("<KeyRelease>", self.filtrar_lista)

        # Lista
        self.scroll_estudiantes = ctk.CTkScrollableFrame(
            frame_izq, fg_color="transparent")
        self.scroll_estudiantes.pack(fill="both", expand=True, padx=10, pady=5)
        self.lista_completa = []

    def crear_panel_derecho(self):
        self.frame_der = ctk.CTkFrame(
            self, fg_color=C["card_alt"], corner_radius=10)
        self.frame_der.grid(row=0, column=1, sticky="nsew")

        ctk.CTkLabel(
            self.frame_der,
            text="Expediente y Observaciones",
            font=(
                "Segoe UI",
                20,
                "bold"),
            text_color="#3B82F6").pack(
            pady=(
                20,
                10))

        # Etiqueta dinámica del alumno
        self.lbl_alumno_sel = ctk.CTkLabel(
            self.frame_der, text="Seleccione un estudiante en la lista 👈", font=(
                "Segoe UI", 16, "italic"), text_color="#94A3B8")
        self.lbl_alumno_sel.pack(pady=10)

        # Controles de observación
        panel_obs = ctk.CTkFrame(self.frame_der, fg_color=C["card"])
        panel_obs.pack(fill="both", expand=True, padx=30, pady=10)

        ctk.CTkLabel(
            panel_obs,
            text="Fecha del reporte (DD-MM-AAAA):",
            font=(
                "Segoe UI",
                12)).pack(
            anchor="w",
            padx=20,
            pady=(
                15,
                5))
        self.entry_fecha = ctk.CTkEntry(panel_obs)
        self.entry_fecha.insert(
            0, datetime.datetime.now().strftime("%d-%m-%Y"))
        self.entry_fecha.pack(fill="x", padx=20)

        ctk.CTkLabel(
            panel_obs,
            text="Categoría:",
            font=(
                "Segoe UI",
                12)).pack(
            anchor="w",
            padx=20,
            pady=(
                15,
                5))
        self.combo_categoria = ctk.CTkOptionMenu(
            panel_obs,
            values=[
                "Conducta",
                "Académico",
                "Citación a Acudiente",
                "Mención Honorífica",
                "Otro"])
        self.combo_categoria.pack(fill="x", padx=20)

        self.plantillas = self.cargar_plantillas()
        self.combo_plantilla = ctk.CTkOptionMenu(
            panel_obs,
            values=["Seleccionar plantilla..."] + list(self.plantillas.keys()),
            command=self.aplicar_plantilla)
        self.combo_plantilla.pack(fill="x", padx=20)

        ctk.CTkLabel(
            panel_obs,
            text="Redacte la observación:",
            font=(
                "Segoe UI",
                12)).pack(
            anchor="w",
            padx=20,
            pady=(
                15,
                5))
        self.texto_obs = SpellcheckTextbox(panel_obs, height=110)
        self.texto_obs.pack(fill="both", expand=True, padx=20, pady=(0, 10))
        self.texto_obs.bind("<KeyRelease>", self.al_escribir_observacion)

        btn_f = ctk.CTkFrame(panel_obs, fg_color="transparent")
        btn_f.pack(fill="x", padx=20, pady=(0, 15))

        self.btn_guardar_plantilla = ctk.CTkButton(
            btn_f, text="💾 Guardar plantilla", fg_color="#334155", hover_color="#475569",
            font=("Segoe UI", 11, "bold"), height=28, command=self.guardar_como_plantilla
        )
        self.btn_guardar_plantilla.pack(side="left", fill="x", expand=True, padx=(0, 5))

        self.btn_sugerir_ia = ctk.CTkButton(
            btn_f, text="🧠 Sugerir con IA", fg_color="#3B82F6", hover_color="#2563EB",
            font=("Segoe UI", 11, "bold"), height=28, command=self.sugerir_observacion_ia, state="disabled"
        )
        self.btn_sugerir_ia.pack(side="right", fill="x", expand=True, padx=(5, 0))

        self.btn_guardar = ctk.CTkButton(self.frame_der, text="📝 GUARDAR EN EXPEDIENTE OFICIAL", fg_color="#10B981", hover_color="#059669",
                                         font=("Segoe UI", 14, "bold"), height=45, command=self.guardar_observacion, state="disabled")
        self.btn_guardar.pack(pady=(20, 10), padx=30, fill="x")

        # Botones de Acción Extra
        extra_frame = ctk.CTkFrame(self.frame_der, fg_color="transparent")
        extra_frame.pack(pady=(0, 20), padx=30, fill="x")

        self.btn_nota_acudiente = ctk.CTkButton(
            extra_frame, text="✉️ Nota para Acudiente", fg_color="#6366F1", hover_color="#4F46E5",
            font=("Segoe UI", 12, "bold"), height=35, command=self.enviar_nota_acudiente, state="disabled"
        )
        self.btn_nota_acudiente.pack(side="left", fill="x", expand=True, padx=(0, 5))

        self.btn_abrir_expediente = ctk.CTkButton(
            extra_frame, text="📂 Abrir Expediente", fg_color="#8B5CF6", hover_color="#7C3AED",
            font=("Segoe UI", 12, "bold"), height=35, command=self.abrir_expediente_estudiante, state="disabled"
        )
        self.btn_abrir_expediente.pack(side="right", fill="x", expand=True, padx=(5, 0))

    def al_cambiar_grado(self, grado):
        self.grado_actual = grado
        self.lista_completa = self.engine.obtener_estudiantes_completos(grado)
        self.entry_buscar.delete(0, 'end')
        self.mostrar_lista(self.lista_completa)

        # Reset panel derecho
        self.estudiante_seleccionado = None
        self.lbl_alumno_sel.configure(
            text="Seleccione un estudiante en la lista 👈",
            text_color="#94A3B8",
            font=(
                "Segoe UI",
                16,
                "italic"))
        self.btn_guardar.configure(state="disabled")
        self.btn_nota_acudiente.configure(state="disabled")
        self.btn_abrir_expediente.configure(state="disabled")


    def filtrar_lista(self, event=None):
        # Debounce: cancelar búsqueda anterior y esperar 200ms
        if hasattr(self, "_filtro_after_id") and self._filtro_after_id:
            try:
                self.after_cancel(self._filtro_after_id)
            except Exception:
                pass
        self._filtro_after_id = self.after(200, self._ejecutar_filtro)

    def _ejecutar_filtro(self):
        self._filtro_after_id = None
        texto = self.entry_buscar.get().lower()
        filtrados = [
            est for est in self.lista_completa if texto in est["nombre"].lower()]
        self.mostrar_lista(filtrados)

    def mostrar_lista(self, lista):
        # Pool de botones reciclables
        if not hasattr(self, "_pool_btns"):
            self._pool_btns = []

        # Ocultar todos los botones del pool
        for btn in self._pool_btns:
            btn.pack_forget()

        # Destruir solo widgets que no estén en el pool (ej: labels de "no hay coincidencias")
        pool_set = set(self._pool_btns)
        for w in self.scroll_estudiantes.winfo_children():
            if w not in pool_set:
                w.destroy()

        if not lista:
            ctk.CTkLabel(
                self.scroll_estudiantes,
                text="No hay coincidencias.",
                text_color="#94A3B8").pack(
                pady=20)
            return

        for i, est in enumerate(lista):
            if i < len(self._pool_btns):
                btn = self._pool_btns[i]
                btn.configure(text=est['nombre'],
                              command=lambda e=est: self.seleccionar_estudiante(e))
            else:
                btn = ctk.CTkButton(self.scroll_estudiantes, text=est['nombre'], fg_color="transparent",
                                    anchor="w", text_color=C["texto"], hover_color=C["hover"],
                                    command=lambda e=est: self.seleccionar_estudiante(e))
                self._pool_btns.append(btn)
            btn.pack(fill="x", pady=2)

    def seleccionar_estudiante(self, estudiante):
        self.estudiante_seleccionado = estudiante
        self.lbl_alumno_sel.configure(
            text=f"👨‍🎓 Estudiante: {estudiante['nombre']}",
            text_color="white",
            font=("Segoe UI", 18, "bold"))
        self.btn_guardar.configure(state="normal")
        self.btn_nota_acudiente.configure(state="normal")
        self.btn_abrir_expediente.configure(state="normal")
        self.btn_sugerir_ia.configure(state="normal")
        self.texto_obs.delete("1.0", "end")

        # Cargar observación borrador desde SQLite si existe (Tarea 12 y 24)
        try:
            cursor = self.engine.db_conn.cursor()
            from utils.date_helpers import obtener_trimestre_actual
            trim_num = obtener_trimestre_actual()
            cursor.execute("SELECT comentario FROM observaciones WHERE estudiante_id = ? AND trimestre = ?;", (str(estudiante['id']), trim_num))
            row = cursor.fetchone()
            if row:
                comentario_desc = self.engine.db_manager.desencriptar_campo(row[0])
                self.texto_obs.insert("1.0", comentario_desc)
                if hasattr(self.texto_obs, "check_spelling"):
                    self.texto_obs.check_spelling()
        except Exception as e:
            print(f"[!] Error cargando borrador desde SQLite: {e}")

    def al_escribir_observacion(self, event=None):
        if getattr(self, "_debounce_job", None):
            self.after_cancel(self._debounce_job)
        self._debounce_job = self.after(2000, self.auto_guardar_borrador)

    def auto_guardar_borrador(self):
        if not self.estudiante_seleccionado:
            return
        observacion = self.texto_obs.get("1.0", "end").strip()
        try:
            cursor = self.engine.db_conn.cursor()
            est_id = str(self.estudiante_seleccionado['id'])
            from utils.date_helpers import obtener_trimestre_actual
            trim_num = obtener_trimestre_actual()
            
            # Cifrar comentario a nivel de columna (Tarea 12)
            comentario_cifrado = self.engine.db_manager.encriptar_campo(observacion)
            
            cursor.execute("SELECT id FROM observaciones WHERE estudiante_id = ? AND trimestre = ?;", (est_id, trim_num))
            row = cursor.fetchone()
            if row:
                cursor.execute("UPDATE observaciones SET comentario = ? WHERE id = ?;", (comentario_cifrado, row[0]))
            else:
                cursor.execute("INSERT INTO observaciones (estudiante_id, trimestre, comentario) VALUES (?, ?, ?);", (est_id, trim_num, comentario_cifrado))
            self.engine.db_conn.commit()
            self.engine.db_manager.guardar_cifrado()
            
            # Mostrar toast sutil indicando el guardado
            root = self.winfo_toplevel()
            if hasattr(root, "mostrar_toast"):
                root.mostrar_toast("✓ Borrador guardado automáticamente", color="#3B82F6")
        except Exception as e:
            print(f"[!] Error en auto_guardar_borrador: {e}")


    def aplicar_plantilla(self, opcion):
        if opcion == "Seleccionar plantilla...":
            return
            
        info = self.plantillas.get(opcion)
        if info:
            texto = info.get("texto", "")
            cat = info.get("categoria", "Otro")
            
            self.combo_categoria.set(cat)
            self.texto_obs.delete("1.0", "end")
            self.texto_obs.insert("1.0", texto)
            if hasattr(self.texto_obs, "check_spelling"):
                self.texto_obs.check_spelling()

    def guardar_observacion(self):
        if not self.estudiante_seleccionado:
            return

        fecha = self.entry_fecha.get().strip()
        categoria = self.combo_categoria.get()
        observacion = self.texto_obs.get("1.0", "end").strip()

        if not observacion:
            messagebox.showwarning(
                "Falta información",
                "Debe redactar la observación antes de guardar.")
            return

        # Sincronizar borrador a base de datos de inmediato (Tarea 24)
        self.auto_guardar_borrador()

        self.btn_guardar.configure(text="⏳ GUARDANDO...", state="disabled")
        self.btn_nota_acudiente.configure(state="disabled")
        self.btn_abrir_expediente.configure(state="disabled")

        carpeta = os.path.join(BASE_DIR, "..", "Expedientes_Estudiantes")
        if not os.path.exists(carpeta):
            os.makedirs(carpeta)

        nombre_est = self.estudiante_seleccionado['nombre']

        def tarea_fondo():
            try:
                self._actualizar_o_crear_word(
                    carpeta,
                    nombre_est,
                    self.grado_actual,
                    fecha,
                    categoria,
                    observacion)
                self.after(0, lambda: self.finalizar_guardado_obs(True, nombre_est, carpeta))
            except Exception as e:
                self.after(0, lambda: self.finalizar_guardado_obs(False, nombre_est, str(e)))

        threading.Thread(target=tarea_fondo, daemon=True).start()

    def finalizar_guardado_obs(self, exito, nombre_est, extra):
        self.btn_guardar.configure(text="📝 GUARDAR EN EXPEDIENTE OFICIAL", state="normal")
        self.btn_nota_acudiente.configure(state="normal")
        self.btn_abrir_expediente.configure(state="normal")

        if exito:
            root = self.winfo_toplevel()
            if hasattr(root, "mostrar_toast"):
                root.mostrar_toast(f"✓ Observación agregada al expediente de {nombre_est}", color="#10B981")
            self.texto_obs.delete("1.0", "end")

            # Preguntar al usuario si desea abrir el documento inmediatamente
            nombre_archivo = f"{nombre_est} - {self.grado_actual.replace('°', '')}.docx".replace("/", "-")
            ruta_archivo = os.path.join(extra, nombre_archivo)
            if messagebox.askyesno("Expediente Actualizado", f"Observación guardada.\n\n¿Desea abrir el expediente de {nombre_est} ahora mismo?"):
                from utils.footer_utils import abrir_documento
                abrir_documento(ruta_archivo)
        else:
            messagebox.showerror("Error", f"No se pudo guardar el expediente: {extra}")

    def enviar_nota_acudiente(self):
        if not self.estudiante_seleccionado:
            return
            
        nombre_est = self.estudiante_seleccionado['nombre']
        cedula_est = self.estudiante_seleccionado.get('cedula', '')
        grado = self.grado_actual
        
        modal = ctk.CTkToplevel(self)
        modal.title(f"✉️ Nota para Acudiente - {nombre_est}")
        modal.geometry("400x380")
        modal.resizable(False, False)
        modal.transient(self)
        modal.grab_set()
        
        from config import establecer_icono_ventana
        establecer_icono_ventana(modal)
        
        modal.update_idletasks()
        w = modal.winfo_width()
        h = modal.winfo_height()
        x = self.winfo_toplevel().winfo_x() + (self.winfo_toplevel().winfo_width() // 2) - (w // 2)
        y = self.winfo_toplevel().winfo_y() + (self.winfo_toplevel().winfo_height() // 2) - (h // 2)
        modal.geometry(f"+{x}+{y}")
        
        ctk.CTkLabel(modal, text="Asunto/Motivo:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=20, pady=(15, 0))
        combo_asunto = ctk.CTkOptionMenu(modal, values=[
            "Citación a Acudiente", 
            "Reporte de Conducta", 
            "Aviso de Bajo Rendimiento", 
            "Felicitaciones por Desempeño"
        ], width=360)
        combo_asunto.pack(padx=20, pady=5)
        
        ctk.CTkLabel(modal, text="Detalle de la nota:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=20, pady=(10, 0))
        txt_detalle = ctk.CTkTextbox(modal, height=120, width=360)
        txt_detalle.pack(padx=20, pady=5)
        
        def generar():
            motivo = combo_asunto.get()
            desc = txt_detalle.get("1.0", "end").strip()
            if not desc:
                messagebox.showwarning("Atención", "Escriba el detalle de la nota.")
                return
                
            from rdsecurity import cargar_config_segura
            from documentos_maestro import generar_nota_acudiente, generar_citacion_txt
            from utils.footer_utils import abrir_documento
            
            cfg = cargar_config_segura({})
            datos = {
                "alumno": nombre_est,
                "cedula": cedula_est,
                "grado": grado,
                "docente_nombre": cfg.get("docente_nombre", ""),
                "escuela_nombre": cfg.get("escuela_nombre", ""),
                "ano_lectivo": cfg.get("ano_lectivo", "2026"),
                "fecha": datetime.datetime.now().strftime("%d-%m-%Y"),
                "motivo": motivo,
                "descripcion": desc,
                "mensaje": desc
            }
            ruta = generar_nota_acudiente(datos)
            ruta_txt = generar_citacion_txt(datos)
            modal.destroy()
            
            if ruta or ruta_txt:
                abrir_txt = messagebox.askyesno(
                    "Abrir Citación", 
                    "La citación ha sido generada.\n\n¿Desea abrir la versión de Texto Plano (.txt) en el Bloc de Notas?\n(Haga clic en 'No' para intentar abrir la plantilla de Word .docx)"
                )
                if abrir_txt and os.path.exists(ruta_txt):
                    abrir_documento(ruta_txt)
                elif not abrir_txt and os.path.exists(ruta):
                    abrir_documento(ruta)
                
                root = self.winfo_toplevel()
                if hasattr(root, "mostrar_toast"):
                    root.mostrar_toast("✓ Citación generada exitosamente", color="#10B981")
                    
        btn_gen = ctk.CTkButton(modal, text="Generar e Imprimir", fg_color="#10B981", hover_color="#059669", command=generar)
        btn_gen.pack(pady=20)

    def abrir_expediente_estudiante(self):
        if not self.estudiante_seleccionado:
            return
            
        nombre_est = self.estudiante_seleccionado['nombre']
        nombre_archivo = f"{nombre_est} - {self.grado_actual.replace('°', '')}.docx".replace("/", "-")
        carpeta = os.path.join(BASE_DIR, "..", "Expedientes_Estudiantes")
        ruta_archivo = os.path.join(carpeta, nombre_archivo)
        
        from utils.footer_utils import abrir_documento
        if os.path.exists(ruta_archivo):
            abrir_documento(ruta_archivo)
        else:
            if messagebox.askyesno("Crear Expediente", f"El expediente de {nombre_est} aún no se ha creado.\n\n¿Desea crearlo ahora?"):
                try:
                    self._actualizar_o_crear_word(
                        carpeta,
                        nombre_est,
                        self.grado_actual,
                        datetime.datetime.now().strftime("%d-%m-%Y"),
                        "Apertura",
                        "Apertura de Expediente Oficial de Seguimiento"
                    )
                    abrir_documento(ruta_archivo)
                except Exception as e:
                    messagebox.showerror("Error", f"No se pudo crear el expediente: {e}")

    # ==========================================
    # LÓGICA DE WORD (Plantilla Oficial)
    # ==========================================
    def _actualizar_o_crear_word(
            self, carpeta, nombre_est, grado, fecha, tipo, motivo):
        if not DOCX_DISPONIBLE:
            raise Exception(
                "Librería python-docx no instalada. (pip install python-docx)")

        nombre_archivo = f"{nombre_est} - {
            grado.replace(
                '°', '')}.docx".replace(
            "/", "-")
        ruta_archivo = os.path.join(carpeta, nombre_archivo)

        def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
                node = OxmlElement(f'w:{m}')
                node.set(qn('w:w'), str(val))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
            
        def set_cell_width(cell, width_dxa):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(width_dxa))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            cell.width = width_dxa

        def get_categoria_color(cat):
            cat = cat.lower()
            if "conducta" in cat:
                return "E2F0D9" # Verde suave
            elif "acad" in cat:
                return "DDEBF7" # Azul suave
            elif "cita" in cat:
                return "FCE4D6" # Naranja suave
            elif "menci" in cat:
                return "FFF2CC" # Amarillo suave
            else:
                return "F2F2F2" # Gris suave

        col_widths = [1152, 1728, 2160, 4320] # N.º, Fecha, Tipo, Descripción

        if os.path.exists(ruta_archivo):
            doc = Document(ruta_archivo)
            if doc.tables:
                tabla = doc.tables[0]
                num_reg = str(len(tabla.rows))
                fila = tabla.add_row()
                
                fila.cells[0].text = num_reg
                fila.cells[1].text = fecha
                fila.cells[2].text = tipo
                fila.cells[3].text = motivo

                bg_color = get_categoria_color(tipo)
                for idx, cell in enumerate(fila.cells):
                    set_cell_width(cell, col_widths[idx])
                    set_cell_margins(cell, 100, 100, 150, 150)
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), bg_color)
                    cell._tc.get_or_add_tcPr().append(shading_elm)
            doc.save(ruta_archivo)
        else:
            doc = Document()
            try:
                from utils.footer_utils import add_header_with_logo, get_school_logo_path
                logo_esc = get_school_logo_path()
                if logo_esc:
                    add_header_with_logo(doc, logo_esc)
            except Exception:
                pass

            from rdsecurity import cargar_config_segura
            cfg = cargar_config_segura({})
            docente = cfg.get("docente_nombre", "Docente de Ejemplo")
            escuela = cfg.get("escuela_nombre", "ESCUELA DE EJEMPLO")
            ano = cfg.get("ano_lectivo", "2026")

            p_head = doc.add_paragraph()
            p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p_head.add_run("MINISTERIO DE EDUCACIÓN\n")
            run1.bold = True
            run1.font.size = Pt(14)

            run2 = p_head.add_run(f"{escuela.upper()}\n")
            run2.bold = True
            run2.font.size = Pt(12)

            run3 = p_head.add_run("DIRECCIÓN REGIONAL DE EDUCACIÓN")
            run3.font.size = Pt(11)

            doc.add_paragraph("\n")

            p_info = doc.add_paragraph()
            p_info.add_run("Docente: ").bold = True
            p_info.add_run(f"{docente}\t\t\t")
            p_info.add_run("Estudiante: ").bold = True
            p_info.add_run(f"{nombre_est}\n")

            p_info.add_run("Grado: ").bold = True
            p_info.add_run(f"{grado}\t\t\t\t")
            p_info.add_run("Año Lectivo: ").bold = True
            p_info.add_run(ano)

            doc.add_paragraph("\n")

            tabla = doc.add_table(rows=1, cols=4)
            tabla.style = 'Table Grid'

            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'Registro N.º'
            hdr_cells[1].text = 'Fecha'
            hdr_cells[2].text = 'Tipo de registro'
            hdr_cells[3].text = 'Descripción (Motivo u observación)'

            for idx, cell in enumerate(hdr_cells):
                set_cell_width(cell, col_widths[idx])
                set_cell_margins(cell, 100, 100, 150, 150)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '1F4E78') # MEDUCA azul oscuro
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)

            fila = tabla.add_row()
            fila.cells[0].text = "1"
            fila.cells[1].text = fecha
            fila.cells[2].text = tipo
            fila.cells[3].text = motivo

            bg_color = get_categoria_color(tipo)
            for idx, cell in enumerate(fila.cells):
                set_cell_width(cell, col_widths[idx])
                set_cell_margins(cell, 100, 100, 150, 150)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), bg_color)
                cell._tc.get_or_add_tcPr().append(shading_elm)

            doc.add_paragraph("\n\n\n\n\n")
            p_firma = doc.add_paragraph(
                "_________________________________\nFirma del Docente")
            p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

            section = doc.sections[0]
            footer = section.footer
            p_foot = footer.paragraphs[0]
            p_foot.text = "Documento Oficial de Seguimiento Estudiantil - RegistroDoc Pro"
            p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_foot.runs[0].font.size = Pt(8)
            p_foot.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            doc.save(ruta_archivo)

    def actualizar_vista(self):
        """Recarga la lista de grados y estudiantes."""
        opciones = self.engine.obtener_grados_activos() or ["Sin datos"]
        old_sel = self.combo_grado.get()
        self.combo_grado.configure(values=opciones)
        if old_sel in opciones:
            self.combo_grado.set(old_sel)
        else:
            self.combo_grado.set(opciones[0])
        self.al_cambiar_grado(self.combo_grado.get())

    def cargar_plantillas(self):
        ruta = os.path.join(BASE_DIR, "..", "data", "observaciones_plantillas.json")
        if os.path.exists(ruta):
            try:
                with open(ruta, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
                
        # Defaults
        defaults = {
            "Conducta: Excelente comportamiento": {
                "categoria": "Conducta",
                "texto": "El estudiante demuestra una conducta ejemplar, compañerismo y una actitud de respeto y liderazgo muy positiva en el aula."
            },
            "Conducta: Falta de respeto a compañeros": {
                "categoria": "Conducta",
                "texto": "El estudiante mostró conductas inapropiadas e irrespetuosas hacia sus compañeros durante la jornada de clases."
            },
            "Conducta: Uso de celular en clase": {
                "categoria": "Conducta",
                "texto": "Se observó al estudiante utilizando el teléfono celular durante la clase sin la debida autorización."
            },
            "Conducta: Interrumpe la clase": {
                "categoria": "Conducta",
                "texto": "El estudiante habla constantemente durante las explicaciones, distrayendo a sus compañeros e interrumpiendo el desarrollo normal de la clase."
            },
            "Conducta: Actitud desafiante / Indisciplina": {
                "categoria": "Conducta",
                "texto": "El estudiante muestra una actitud desafiante ante las indicaciones del docente y se niega a seguir las normas de convivencia."
            },
            "Conducta: Agresión física / Riña": {
                "categoria": "Conducta",
                "texto": "El estudiante se vio involucrado en una agresión física con un compañero dentro del centro educativo."
            },
            "Conducta: Vocabulario inadecuado": {
                "categoria": "Conducta",
                "texto": "El estudiante utiliza lenguaje vulgar o expresiones soeces para comunicarse en el aula de clases."
            },
            "Conducta: Daño a la propiedad escolar": {
                "categoria": "Conducta",
                "texto": "El estudiante causó daños materiales a la propiedad escolar de forma deliberada."
            },
            "Conducta: Falta de uniforme / Presentación": {
                "categoria": "Conducta",
                "texto": "El estudiante asiste a clases sin cumplir con las normas de uniforme y presentación personal requeridas por la institución."
            },
            "Académico: Excelente esfuerzo y participación": {
                "categoria": "Académico",
                "texto": "El estudiante demuestra un alto nivel de compromiso, entrega todas sus asignaciones a tiempo y participa de forma activa y destacada en clase."
            },
            "Académico: Incumplimiento de tareas": {
                "categoria": "Académico",
                "texto": "El estudiante no entregó las tareas o actividades programadas para la fecha establecida, lo cual afecta negativamente su rendimiento."
            },
            "Académico: Desinterés en clase": {
                "categoria": "Académico",
                "texto": "Se observa falta de interés y concentración por parte del estudiante en las explicaciones y actividades de clase."
            },
            "Académico: No trae materiales": {
                "categoria": "Académico",
                "texto": "El estudiante asiste a clases de manera recurrente sin sus libros, cuadernos o materiales necesarios para trabajar."
            },
            "Académico: Progreso académico notable": {
                "categoria": "Académico",
                "texto": "Se felicita al estudiante por mostrar una mejoría significativa en sus calificaciones y nivel de atención."
            },
            "Citación: Bajo rendimiento académico": {
                "categoria": "Citación a Acudiente",
                "texto": "Se cita formalmente al acudiente para conversar acerca del rendimiento académico deficiente y buscar estrategias de mejora conjuntas."
            },
            "Citación: Reiteradas fallas de conducta": {
                "categoria": "Citación a Acudiente",
                "texto": "Se solicita la presencia urgente del acudiente para abordar las reiteradas conductas inapropiadas reportadas del estudiante."
            },
            "Citación: Inasistencias injustificadas": {
                "categoria": "Citación a Acudiente",
                "texto": "Se convoca al acudiente debido a las constantes inasistencias sin justificar que presenta el estudiante."
            },
            "Mención: Excelente promedio y conducta": {
                "categoria": "Mención Honorífica",
                "texto": "Se otorga esta mención honorífica por su sobresaliente desempeño académico, esfuerzo constante y conducta impecable durante este período."
            }
        }
        
        try:
            with open(ruta, "w", encoding="utf-8") as f:
                json.dump(defaults, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
        return defaults

    def guardar_como_plantilla(self):
        observacion = self.texto_obs.get("1.0", "end").strip()
        if not observacion:
            messagebox.showwarning("Atención", "Redacte una observación antes de guardarla como plantilla.")
            return

        modal = ctk.CTkToplevel(self)
        modal.title("💾 Guardar como Plantilla")
        modal.geometry("380x210")
        modal.resizable(False, False)
        modal.transient(self)
        modal.grab_set()
        
        from config import establecer_icono_ventana
        establecer_icono_ventana(modal)
        
        modal.update_idletasks()
        w = modal.winfo_width()
        h = modal.winfo_height()
        x = self.winfo_toplevel().winfo_x() + (self.winfo_toplevel().winfo_width() // 2) - (w // 2)
        y = self.winfo_toplevel().winfo_y() + (self.winfo_toplevel().winfo_height() // 2) - (h // 2)
        modal.geometry(f"+{x}+{y}")
        
        ctk.CTkLabel(modal, text="Categoría de la plantilla:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=20, pady=(15, 0))
        cat_actual = self.combo_categoria.get()
        combo_cat = ctk.CTkOptionMenu(modal, values=["Conducta", "Académico", "Citación a Acudiente", "Mención Honorífica", "Otro"], width=340)
        combo_cat.set(cat_actual)
        combo_cat.pack(padx=20, pady=5)
        
        ctk.CTkLabel(modal, text="Nombre/Título de la plantilla:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=20, pady=(10, 0))
        entry_nombre = ctk.CTkEntry(modal, placeholder_text="Ej: Conducta: Falta grave", width=340)
        snippet = observacion[:25] + "..." if len(observacion) > 25 else observacion
        entry_nombre.insert(0, f"{cat_actual}: {snippet}")
        entry_nombre.pack(padx=20, pady=5)
        
        def guardar():
            nombre = entry_nombre.get().strip()
            cat = combo_cat.get()
            if not nombre:
                messagebox.showwarning("Atención", "Escriba un nombre para la plantilla.")
                return
            
            ruta = os.path.join(BASE_DIR, "..", "data", "observaciones_plantillas.json")
            self.plantillas[nombre] = {
                "categoria": cat,
                "texto": observacion
            }
            try:
                os.makedirs(os.path.dirname(ruta), exist_ok=True)
                with open(ruta, "w", encoding="utf-8") as f:
                    json.dump(self.plantillas, f, ensure_ascii=False, indent=2)
                
                # Update option menu values
                self.combo_plantilla.configure(values=["Seleccionar plantilla..."] + list(self.plantillas.keys()))
                self.combo_plantilla.set(nombre)
                
                modal.destroy()
                root = self.winfo_toplevel()
                if hasattr(root, "mostrar_toast"):
                    root.mostrar_toast("✓ Plantilla guardada exitosamente", color="#10B981")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar la plantilla: {e}")
                
        btn_g = ctk.CTkButton(modal, text="Guardar Plantilla", fg_color="#10B981", hover_color="#059669", command=guardar)
        btn_g.pack(pady=15)

    def sugerir_observacion_ia(self):
        if not self.estudiante_seleccionado:
            return
            
        est = self.estudiante_seleccionado
        id_est = est['id']
        grado = self.combo_grado.get()
        
        from utils.date_helpers import obtener_trimestre_actual
        trim_num = obtener_trimestre_actual()
        trim_str = f"Trimestre {trim_num}"
        
        # 1. Obtener promedio de calificaciones
        promedio = 3.0
        try:
            proms = self.engine.obtener_promedios_reales(grado, None, trim_str)
            if proms and est['nombre'] in proms:
                promedio = proms[est['nombre']]
        except Exception:
            pass
            
        # 2. Obtener estadísticas de asistencia
        ausencias = 0
        tardanzas = 0
        try:
            stats_asistencia = self.engine.obtener_estadisticas_asistencia(grado, trim_str, id_est)
            ausencias = stats_asistencia.get("ausencias", 0)
            tardanzas = stats_asistencia.get("tardanzas", 0)
        except Exception:
            pass
            
        # 3. Obtener hábitos (Tarea 18)
        habitos_negativos = []
        try:
            cursor = self.engine.db_conn.cursor()
            cursor.execute("SELECT criterio_codigo, nota FROM habitos WHERE estudiante_id = ? AND trimestre = ?;", (str(id_est), trim_num))
            for crit, nota in cursor.fetchall():
                if nota in ["R", "X"]:
                    habitos_negativos.append(crit)
        except Exception:
            pass
            
        # 4. Generar sugerencia cualitativa adaptada
        obs_text = ""
        cat_sug = "Académico"
        
        if promedio >= 4.5:
            obs_text += f"El estudiante presenta un rendimiento académico sobresaliente con un promedio general de {promedio:.1f}. Demuestra un alto nivel de asimilación de conceptos y compromiso."
            cat_sug = "Mención Honorífica"
        elif promedio >= 3.0:
            obs_text += f"El estudiante muestra un desempeño académico satisfactorio con un promedio de {promedio:.1f}. Se le insta a mantener el ritmo de estudio y participación."
            cat_sug = "Académico"
        else:
            obs_text += f"El estudiante presenta dificultades académicas críticas con un promedio de {promedio:.1f}. Requiere reforzamiento continuo y apoyo guiado para mejorar sus resultados."
            cat_sug = "Académico"
            
        # Asistencia
        if ausencias > 3:
            obs_text += f" Sin embargo, registra {ausencias} ausencias injustificadas, lo cual impacta negativamente su continuidad escolar."
        elif tardanzas > 4:
            obs_text += f" Se recomienda mejorar la puntualidad, ya que acumula {tardanzas} tardanzas en el periodo."
            
        # Hábitos
        if habitos_negativos:
            hab_str = ", ".join(habitos_negativos)
            obs_text += f" Adicionalmente, requiere atención y reforzamiento conductual en los siguientes criterios de hábitos y actitudes: {hab_str}."
            if "Responsabilidad" in habitos_negativos or "Puntualidad" in habitos_negativos:
                obs_text += " Se sugiere establecer pautas de disciplina en casa para fortalecer el compromiso."
            cat_sug = "Conducta"
            
        # Si no hay problemas
        if promedio >= 4.5 and ausencias == 0 and not habitos_negativos:
            obs_text += " ¡Excelente conducta y asistencia ejemplar!"
            
        # Insertar en el cuadro de texto
        self.combo_categoria.set(cat_sug)
        self.texto_obs.delete("1.0", "end")
        self.texto_obs.insert("1.0", obs_text)
        if hasattr(self.texto_obs, "check_spelling"):
            self.texto_obs.check_spelling()
            
        # Guardar borrador automáticamente
        self.auto_guardar_borrador()

