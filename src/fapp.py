import customtkinter as ctk
from tkinter import messagebox
import datetime
import threading
import os
from config import BASE_DIR

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

# DICCIONARIOS DE TRADUCCIÓN PANTALLA <-> EXCEL
UI_A_EXCEL = {"P": ".", "A": "-", "T": "T", "E": "E"}
EXCEL_A_UI = {".": "P", "-": "A", "T": "T", "E": "E", None: "P"}


class AsistenciaFrame(ctk.CTkFrame):
    def __init__(self, master, engine, **kwargs):
        super().__init__(master, fg_color="transparent", **kwargs)
        self.engine = engine
        self.entradas_asistencia = {}
        self.col_a_modificar = None

        self.grid_columnconfigure(0, weight=7)
        self.grid_columnconfigure(1, weight=3)
        self.grid_rowconfigure(0, weight=1)

        self.crear_panel_izquierdo()
        self.crear_panel_derecho()

        self.al_cambiar_grado(self.combo_grado.get())

    def crear_panel_izquierdo(self):
        frame_izq = ctk.CTkFrame(self, fg_color="#1A2638", corner_radius=10)
        frame_izq.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        top = ctk.CTkFrame(frame_izq, fg_color="transparent")
        top.pack(fill="x", padx=20, pady=15)

        ctk.CTkLabel(
            top,
            text="Grado:",
            font=(
                "Segoe UI",
                16,
                "bold")).pack(
            side="left")
        grados = self.engine.obtener_grados_activos() or ["Sin datos"]

        self.combo_grado = ctk.CTkOptionMenu(
            top, values=grados, command=self.al_cambiar_grado)
        self.combo_grado.pack(side="left", padx=10)

        header = ctk.CTkFrame(frame_izq, fg_color="#253650", corner_radius=5)
        header.pack(fill="x", padx=15, pady=(5, 0), ipady=5)
        ctk.CTkLabel(
            header,
            text="N°",
            width=30,
            font=(
                "Segoe UI",
                13,
                "bold")).pack(
            side="left",
            padx=5)
        ctk.CTkLabel(
            header,
            text="ESTUDIANTE",
            width=220,
            anchor="w",
            font=(
                "Segoe UI",
                13,
                "bold")).pack(
            side="left",
            padx=5)
        ctk.CTkLabel(
            header,
            text="ESTADO",
            width=120,
            font=(
                "Segoe UI",
                13,
                "bold")).pack(
            side="left",
            padx=10)
        ctk.CTkLabel(
            header,
            text="JUSTIFICACIÓN DE AUSENCIA/TARDANZA",
            anchor="w",
            font=(
                "Segoe UI",
                13,
                "bold")).pack(
            side="left",
            padx=10)

        self.scroll_estudiantes = ctk.CTkScrollableFrame(
            frame_izq, fg_color="transparent")
        self.scroll_estudiantes.pack(fill="both", expand=True, padx=15, pady=5)

    def crear_panel_derecho(self):
        frame_der = ctk.CTkFrame(self, fg_color="#1E2D42", corner_radius=10)
        frame_der.grid(row=0, column=1, sticky="nsew")

        ctk.CTkLabel(
            frame_der,
            text="Registro Diario",
            font=(
                "Segoe UI",
                18,
                "bold"),
            text_color="#3B82F6").pack(
            pady=(
                20,
                10))
        ctk.CTkLabel(
            frame_der,
            text="Leyenda: P(.) A(-) T(T)",
            text_color="#94A3B8",
            font=(
                "Segoe UI",
                11)).pack()

        self.tabs = ctk.CTkTabview(frame_der)
        self.tabs.pack(fill="both", expand=True, padx=10, pady=10)

        tab_nueva = self.tabs.add("Pasar Lista")
        tab_mod = self.tabs.add("Modificar")

        # ====== TAB: PASAR LISTA NUEVA ======
        ctk.CTkLabel(
            tab_nueva,
            text="Trimestre:",
            font=(
                "Segoe UI",
                12)).pack(
            anchor="w",
            padx=10,
            pady=(
                10,
                0))
        self.combo_trimestre = ctk.CTkOptionMenu(
            tab_nueva, values=["Trimestre 1", "Trimestre 2", "Trimestre 3"])
        self.combo_trimestre.pack(fill="x", padx=10, pady=5)

        ctk.CTkLabel(
            tab_nueva,
            text="Fecha (DD-MM):",
            font=(
                "Segoe UI",
                12)).pack(
            anchor="w",
            padx=10,
            pady=(
                15,
                0))
        self.entry_fecha = ctk.CTkEntry(tab_nueva)
        self.entry_fecha.insert(0, datetime.datetime.now().strftime("%d-%m"))
        self.entry_fecha.pack(fill="x", padx=10, pady=5)

        # ─── ACCIONES RÁPIDAS ───
        quick_frame = ctk.CTkFrame(tab_nueva, fg_color="transparent")
        quick_frame.pack(fill="x", padx=10, pady=(15, 5))

        ctk.CTkButton(
            quick_frame,
            text="✅ Todos Presentes",
            fg_color="#10B981",
            hover_color="#059669",
            font=("Segoe UI", 11, "bold"),
            height=30,
            command=lambda: self._marcar_todos("P")).pack(side="left", expand=True, fill="x", padx=(0, 4))

        ctk.CTkButton(
            quick_frame,
            text="❌ Todos Ausentes",
            fg_color="#EF4444",
            hover_color="#DC2626",
            font=("Segoe UI", 11, "bold"),
            height=30,
            command=lambda: self._marcar_todos("A")).pack(side="left", expand=True, fill="x", padx=(4, 0))

        ctk.CTkLabel(
            tab_nueva,
            text="💡 Marque todos presentes y corrija solo las excepciones",
            font=("Segoe UI", 10),
            text_color="#94A3B8").pack(padx=10, pady=(2, 8))

        self.btn_guardar_nueva = ctk.CTkButton(
            tab_nueva,
            text="💾 GUARDAR ASISTENCIA",
            fg_color="#10B981",
            hover_color="#059669",
            font=(
                "Segoe UI",
                14,
                "bold"),
            height=40,
            command=self.guardar_asistencia)
        self.btn_guardar_nueva.pack(pady=10, padx=10, fill="x")

        # ====== TAB: MODIFICAR ASISTENCIA ======
        ctk.CTkLabel(
            tab_mod,
            text="Trimestre a corregir:",
            font=(
                "Segoe UI",
                12)).pack(
            anchor="w",
            padx=10,
            pady=(
                10,
                0))
        self.combo_trimestre_mod = ctk.CTkOptionMenu(
            tab_mod,
            values=[
                "Trimestre 1",
                "Trimestre 2",
                "Trimestre 3"],
            command=self.cargar_fechas)
        self.combo_trimestre_mod.pack(fill="x", padx=10, pady=5)

        ctk.CTkLabel(
            tab_mod,
            text="Seleccione fecha existente:",
            font=(
                "Segoe UI",
                12,
                "bold")).pack(
            anchor="w",
            padx=10,
            pady=(
                15,
                0))
        self.combo_fechas_mod = ctk.CTkOptionMenu(
            tab_mod, values=["Buscando..."])
        self.combo_fechas_mod.pack(fill="x", padx=10, pady=5)

        ctk.CTkButton(
            tab_mod,
            text="🔍 CARGAR A LA LISTA",
            fg_color="#F59E0B",
            hover_color="#D97706",
            font=(
                "Segoe UI",
                12,
                "bold"),
            command=self.buscar_modificar).pack(
            pady=10,
            padx=10,
            fill="x")

        self.btn_actualizar = ctk.CTkButton(
            tab_mod,
            text="🔄 ACTUALIZAR EXCEL",
            fg_color="#3B82F6",
            hover_color="#2563EB",
            font=(
                "Segoe UI",
                14,
                "bold"),
            height=40,
            command=self.actualizar_asistencia)
        self.btn_actualizar.pack(pady=20, padx=10, fill="x")

    def al_cambiar_grado(self, grado):
        self.cargar_estudiantes(grado)
        self.cargar_fechas()

    def cargar_estudiantes(self, grado=None):
        if grado is None:
            grado = self.combo_grado.get()
        for w in self.scroll_estudiantes.winfo_children():
            w.destroy()
        self.entradas_asistencia.clear()
        self.col_a_modificar = None

        ests = self.engine.obtener_estudiantes_completos(grado)
        for est in ests:
            row = ctk.CTkFrame(self.scroll_estudiantes, fg_color="transparent")
            row.pack(fill="x", pady=2)

            ctk.CTkLabel(row, text=f"{est['id']}.", width=30).pack(side="left")
            ctk.CTkLabel(
                row,
                text=est['nombre'],
                width=220,
                anchor="w").pack(
                side="left")

            seg_btn = ctk.CTkSegmentedButton(
                row,
                values=[
                    "P",
                    "A",
                    "T",
                    "E"],
                width=160,
                selected_color="#3B82F6")
            seg_btn.set("P")
            seg_btn.pack(side="left", padx=10)

            # Bloqueado por defecto porque inicia en "P" (Presente)
            entry_exc = ctk.CTkEntry(
                row,
                placeholder_text="Solo si falta, tarde o excusa",
                fg_color="#0F1923",
                state="disabled")
            entry_exc.pack(side="left", fill="x", expand=True, padx=5)

            seg_btn.configure(
                command=lambda valor,
                entry=entry_exc: self.activar_excusa(
                    valor,
                    entry))

            # Bind keyboard navigation on the excuse entries
            entry_exc.bind("<Return>", lambda e, idx=est['id']: self.al_presionar_enter(idx))
            entry_exc.bind("<Down>", lambda e, idx=est['id']: self.al_presionar_abajo(idx))
            entry_exc.bind("<Up>", lambda e, idx=est['id']: self.al_presionar_arriba(idx))

            self.entradas_asistencia[est['id']] = {
                "nombre": est['nombre'], "btn": seg_btn, "exc": entry_exc}

    def activar_excusa(self, valor, entry_widget):
        """Habilita la casilla SOLO para ausencias, tardanzas y excusas."""
        if valor in ["A", "T", "E"]:
            entry_widget.configure(
                state="normal",
                placeholder_text="Escriba la justificación...")
        else:
            entry_widget.delete(0, "end")
            entry_widget.configure(
                state="disabled",
                placeholder_text="Solo si falta, tarde o excusa")

    def _marcar_todos(self, estado):
        """Marca todos los alumnos con el estado dado (P, A, T, E)."""
        for id_est, widgets in self.entradas_asistencia.items():
            widgets["btn"].set(estado)
            self.activar_excusa(estado, widgets["exc"])
        # Toast feedback
        root = self.winfo_toplevel()
        if estado == "P":
            msg = "✅ Todos marcados como PRESENTES — corrija las excepciones"
            color = "#10B981"
        else:
            msg = "❌ Todos marcados como AUSENTES — corrija las excepciones"
            color = "#EF4444"
        if hasattr(root, "mostrar_toast"):
            root.mostrar_toast(msg, color=color)

    def cargar_fechas(self, *args):
        grado = self.combo_grado.get()
        trimestre = self.combo_trimestre_mod.get()
        fechas = self.engine.obtener_fechas_asistencia(grado, trimestre)

        if fechas:
            self.combo_fechas_mod.configure(values=fechas)
            self.combo_fechas_mod.set(fechas[-1])
        else:
            vacio = ["Sin registro en este trimestre"]
            self.combo_fechas_mod.configure(values=vacio)
            self.combo_fechas_mod.set(vacio[0])

    def recopilar_datos(self):
        dic_asistencia = {}
        lista_excusas = []
        for id_est, widgets in self.entradas_asistencia.items():
            estado_ui = widgets["btn"].get()
            excusa = widgets["exc"].get().strip()

            estado_excel = UI_A_EXCEL.get(estado_ui, ".")
            dic_asistencia[id_est] = {"estado": estado_excel}

            # Guardamos excusas si es Ausencia, Tardanza o Excusa
            if estado_ui in ["A", "T", "E"]:
                if estado_ui == "A":
                    tipo_reg = "Ausencia"
                elif estado_ui == "T":
                    tipo_reg = "Tardanza"
                else:
                    tipo_reg = "Excusa Justificada"
                lista_excusas.append({
                    "nombre": widgets["nombre"],
                    "estado": tipo_reg,
                    "motivo": excusa if excusa else "Falta sin justificar"
                })
        return dic_asistencia, lista_excusas

    def _obtener_justificacion_word(self, nombre_est, grado, fecha):
        carpeta = os.path.join(BASE_DIR, "..", "Expedientes_Estudiantes")
        nombre_archivo = f"{nombre_est} - {grado.replace('°', '')}.docx".replace("/", "-")
        ruta_archivo = os.path.join(carpeta, nombre_archivo)
        if os.path.exists(ruta_archivo) and DOCX_DISPONIBLE:
            try:
                doc = Document(ruta_archivo)
                if doc.tables:
                    tabla = doc.tables[0]
                    for row in reversed(tabla.rows[1:]):
                        if len(row.cells) >= 4:
                            f = row.cells[1].text.strip()
                            t = row.cells[2].text.strip()
                            desc = row.cells[3].text.strip()
                            if f == fecha and t in ["Ausencia", "Tardanza", "Excusa Justificada"]:
                                return desc
            except Exception as e:
                print(f"Error cargando justificación de Word para {nombre_est}: {e}")
        return ""

    def guardar_asistencia(self):
        grado = self.combo_grado.get()
        trimestre = self.combo_trimestre.get()
        fecha = self.entry_fecha.get().strip()

        if not fecha:
            messagebox.showwarning("Atención", "Debe colocar una fecha.")
            return

        dic_asistencia, lista_excusas = self.recopilar_datos()

        # Validar si hay faltas/tardanzas/excusas sin justificación escrita
        for exc in lista_excusas:
            if exc["motivo"] == "Falta sin justificar":
                msg_alerta = (
                    f"El estudiante {exc['nombre']} está marcado como '{exc['estado']}' pero no se ha escrito un motivo o justificación.\n\n"
                    "Recuerde preguntar al grupo o acudiente si se sabe la razón (ej. cita, enfermedad, transporte).\n\n"
                    "¿Desea guardarlo como 'Falta sin justificar'?"
                )
                if not messagebox.askyesno("Justificación Faltante", msg_alerta):
                    return

        self.btn_guardar_nueva.configure(
            text="⏳ GUARDANDO...",
            fg_color="#F59E0B",
            state="disabled")
        self.update()

        def tarea():
            exito, msj = self.engine.guardar_asistencia(
                grado, trimestre, fecha, dic_asistencia)
            self.after(0, lambda: self.finalizar_guardado(
                exito, msj, lista_excusas, grado, fecha))

        threading.Thread(target=tarea, daemon=True).start()

    def finalizar_guardado(self, exito, msj, lista_excusas, grado, fecha):
        self.btn_guardar_nueva.configure(
            text="💾 GUARDAR ASISTENCIA",
            fg_color="#10B981",
            state="normal")

        if not exito:
            messagebox.showerror("Error de Excel", msj)
            return

        self.cargar_fechas()
        threading.Thread(
            target=self.procesar_expedientes_word,
            args=(grado, fecha, lista_excusas, "Guardado exitoso."),
            daemon=True
        ).start()
        self.cargar_estudiantes(grado)

    def buscar_modificar(self):
        grado = self.combo_grado.get()
        trimestre = self.combo_trimestre_mod.get()
        fecha = self.combo_fechas_mod.get()

        if "Sin registro" in fecha or "Buscando" in fecha:
            return

        resultado = self.engine.buscar_asistencia_existente(
            grado, trimestre, fecha)
        if not resultado:
            return

        self.col_a_modificar = resultado["columna"]
        datos_excel = resultado["asistencia"]

        for id_est, widgets in self.entradas_asistencia.items():
            if id_est in datos_excel:
                estado_excel = datos_excel[id_est]
                estado_ui = EXCEL_A_UI.get(estado_excel, "P")
                widgets["btn"].set(estado_ui)
                self.activar_excusa(estado_ui, widgets["exc"])
                widgets["exc"].delete(0, 'end')
                if estado_ui in ["A", "T", "E"]:
                    motivo_guardado = self._obtener_justificacion_word(widgets["nombre"], grado, fecha)
                    if motivo_guardado:
                        widgets["exc"].insert(0, motivo_guardado)

        messagebox.showinfo(
            "Modo Edición",
            f"Asistencia del {fecha} cargada.\n\nEdítela y presione ACTUALIZAR EXCEL.")

    def actualizar_asistencia(self):
        if not self.col_a_modificar:
            return

        grado = self.combo_grado.get()
        trimestre = self.combo_trimestre_mod.get()
        fecha = self.combo_fechas_mod.get()
        dic_asistencia, lista_excusas = self.recopilar_datos()

        # Validar si hay faltas/tardanzas/excusas sin justificación escrita en el modo edición
        for exc in lista_excusas:
            if exc["motivo"] == "Falta sin justificar":
                msg_alerta = (
                    f"El estudiante {exc['nombre']} está marcado como '{exc['estado']}' pero no se ha escrito un motivo o justificación.\n\n"
                    "Recuerde preguntar al grupo o acudiente si se sabe la razón (ej. cita, enfermedad, transporte).\n\n"
                    "¿Desea guardarlo como 'Falta sin justificar'?"
                )
                if not messagebox.askyesno("Justificación Faltante", msg_alerta):
                    return

        self.btn_actualizar.configure(
            text="⏳ ACTUALIZANDO...",
            fg_color="#F59E0B",
            state="disabled")
        self.update()

        def tarea():
            exito = self.engine.actualizar_asistencia(
                grado, trimestre, self.col_a_modificar, dic_asistencia)
            self.after(
                0, lambda: self.finalizar_actualizacion(
                    exito, grado, fecha, lista_excusas))

        threading.Thread(target=tarea, daemon=True).start()

    def finalizar_actualizacion(self, exito, grado, fecha, lista_excusas):
        self.btn_actualizar.configure(
            text="🔄 ACTUALIZAR EXCEL",
            fg_color="#3B82F6",
            state="normal")
        if exito:
            self.col_a_modificar = None
            threading.Thread(
                target=self.procesar_expedientes_word,
                args=(grado, fecha, lista_excusas, "Asistencia actualizada correctamente."),
                daemon=True
            ).start()
            self.cargar_estudiantes(grado)
        else:
            messagebox.showerror("Error", "No se pudo actualizar.")

    # =============================================================
    # MOTOR GENERADOR DE EXPEDIENTES EN WORD
    # =============================================================
    def procesar_expedientes_word(
            self,
            grado,
            fecha,
            lista_excusas,
            mensaje_base):
        if not lista_excusas:
            self.after(0, lambda: self._mostrar_toast_safe(f"✓ {mensaje_base} (Asistencia perfecta)", "#10B981"))
            return

        if not DOCX_DISPONIBLE:
            self.after(0, lambda: self._mostrar_toast_safe(f"✓ {mensaje_base} (Falta python-docx)", "#F59E0B"))
            return

        carpeta_expedientes = os.path.join(
            BASE_DIR, "..", "Expedientes_Estudiantes")
        if not os.path.exists(carpeta_expedientes):
            os.makedirs(carpeta_expedientes)

        for exc in lista_excusas:
            self._actualizar_o_crear_word(
                carpeta_expedientes,
                exc['nombre'],
                grado,
                fecha,
                exc['estado'],
                exc['motivo'])

        self.after(0, lambda: self._mostrar_toast_safe(f"✓ {mensaje_base} (Expedientes actualizados)", "#10B981"))


    def _actualizar_o_crear_word(
            self,
            carpeta,
            nombre_est,
            grado,
            fecha,
            tipo,
            motivo):
        nombre_archivo = f"{nombre_est} - {
            grado.replace(
                '°', '')}.docx".replace(
            "/", "-")
        ruta_archivo = os.path.join(carpeta, nombre_archivo)

        if os.path.exists(ruta_archivo):
            try:
                doc = Document(ruta_archivo)
                if doc.tables:
                    tabla = doc.tables[0]
                    
                    # Buscar si ya existe una fila de asistencia para esta fecha
                    fila_existente = None
                    for row in tabla.rows[1:]:
                        if len(row.cells) >= 4:
                            f = row.cells[1].text.strip()
                            t = row.cells[2].text.strip()
                            if f == fecha and t in ["Ausencia", "Tardanza", "Excusa Justificada"]:
                                fila_existente = row
                                break
                    
                    if fila_existente:
                        fila_existente.cells[2].text = tipo
                        fila_existente.cells[3].text = motivo
                    else:
                        num_reg = str(len(tabla.rows))
                        fila = tabla.add_row()
                        fila.cells[0].text = num_reg
                        fila.cells[1].text = fecha
                        fila.cells[2].text = tipo
                        fila.cells[3].text = motivo
                doc.save(ruta_archivo)
            except Exception as e:
                print(f"No se pudo actualizar el Word de {nombre_est}: {e}")
        else:
            doc = Document()
            try:
                from utils.footer_utils import add_header_with_logo, get_school_logo_path
                logo_esc = get_school_logo_path()
                if logo_esc:
                    add_header_with_logo(doc, logo_esc)
            except Exception:
                pass

            p_head = doc.add_paragraph()
            p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p_head.add_run("MINISTERIO DE EDUCACIÓN\n")
            run1.bold = True
            run1.font.size = Pt(14)

            from rdsecurity import cargar_config_segura
            cfg = cargar_config_segura({})
            docente = cfg.get("docente_nombre", "Elmer Tugri")
            escuela = cfg.get("escuela_nombre", "ESCUELA CERRO CACICÓN")
            ano = cfg.get("ano_lectivo", "2026")

            run2 = p_head.add_run(f"{escuela.upper()}\n")
            run2.bold = True
            run2.font.size = Pt(12)

            run3 = p_head.add_run("DIRECCIÓN REGIONAL DE EDUCACIÓN")
            run3.font.size = Pt(11)

            doc.add_paragraph("\n")

            p_info = doc.add_paragraph()
            p_info.add_run("Docente: ").bold = True
            p_info.add_run(f"{docente}\t\t\t")
            p_info.add_run("Estudiante: ").bold = True
            p_info.add_run(f"{nombre_est}\n")

            p_info.add_run("Grado: ").bold = True
            p_info.add_run(f"{grado}\t\t\t\t")
            p_info.add_run("Año Lectivo: ").bold = True
            p_info.add_run(ano)


            doc.add_paragraph("\n")

            tabla = doc.add_table(rows=1, cols=4)
            tabla.style = 'Table Grid'

            for cell in tabla.rows[0].cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D9E2F3')
                cell._tc.get_or_add_tcPr().append(shading_elm)

            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'Registro N.º'
            hdr_cells[1].text = 'Fecha'
            hdr_cells[2].text = 'Tipo de registro'
            hdr_cells[3].text = 'Descripción (Motivo u observación)'

            for celda in hdr_cells:
                for p in celda.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

            fila = tabla.add_row()
            fila.cells[0].text = "1"
            fila.cells[1].text = fecha
            fila.cells[2].text = tipo
            fila.cells[3].text = motivo

            doc.add_paragraph("\n\n\n\n\n")
            p_firma = doc.add_paragraph(
                "_________________________________\nFirma del Docente")
            p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

            section = doc.sections[0]
            footer = section.footer
            p_foot = footer.paragraphs[0]
            p_foot.text = "Documento Oficial de Seguimiento Estudiantil - RegistroDoc Pro"
            p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_foot.runs[0].font.size = Pt(8)
            p_foot.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            try:
                doc.save(ruta_archivo)
            except Exception as e:
                print(f"Error guardando Word: {e}")

    def _mostrar_toast_safe(self, mensaje, color):
        root = self.winfo_toplevel()
        if hasattr(root, "mostrar_toast"):
            root.mostrar_toast(mensaje, color=color)

    def al_presionar_enter(self, id_est):
        self.al_presionar_abajo(id_est)
        return "break"

    def al_presionar_abajo(self, id_est):
        ids = sorted(list(self.entradas_asistencia.keys()))
        try:
            curr_idx = ids.index(id_est)
            for i in range(curr_idx + 1, len(ids)):
                next_id = ids[i]
                widgets = self.entradas_asistencia[next_id]
                entry = widgets["exc"]
                if entry.cget("state") != "disabled":
                    entry.focus_set()
                    entry.select_range(0, 'end')
                    break
        except ValueError:
            pass

    def al_presionar_arriba(self, id_est):
        ids = sorted(list(self.entradas_asistencia.keys()))
        try:
            curr_idx = ids.index(id_est)
            for i in range(curr_idx - 1, -1, -1):
                prev_id = ids[i]
                widgets = self.entradas_asistencia[prev_id]
                entry = widgets["exc"]
                if entry.cget("state") != "disabled":
                    entry.focus_set()
                    entry.select_range(0, 'end')
                    break
        except ValueError:
            pass

    def actualizar_vista(self):
        """Recarga la lista de grados y estudiantes."""
        opciones = self.engine.obtener_grados_activos() or ["Sin datos"]
        old_sel = self.combo_grado.get()
        self.combo_grado.configure(values=opciones)
        if old_sel in opciones:
            self.combo_grado.set(old_sel)
        else:
            self.combo_grado.set(opciones[0])
        self.al_cambiar_grado(self.combo_grado.get())

