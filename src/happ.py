import customtkinter as ctk
import os
from tkinter import messagebox

class ReportesFrame(ctk.CTkFrame):
    def __init__(self, master, engine, **kwargs):
        super().__init__(master, fg_color="transparent", **kwargs)
        self.engine = engine

        self.pack_propagate(False)

        ctk.CTkLabel(self, text="Reportes y Estadísticas", font=("Segoe UI", 24, "bold")).pack(anchor="w", pady=(0, 10))

# Botones de acción
        acciones_frame = ctk.CTkFrame(self, fg_color="transparent")
        acciones_frame.pack(fill="x", pady=(0, 5))
        ctk.CTkButton(acciones_frame, text="🖨️ Imprimir Reporte", fg_color="#3B82F6", command=self.imprimir_reporte).pack(side="left", padx=10)
        ctk.CTkButton(acciones_frame, text="⬇️ Exportar a PDF", fg_color="#10B981", command=self.descargar_pdf).pack(side="left", padx=10)
        ctk.CTkButton(acciones_frame, text="🖨️ Solo Gráficos", fg_color="#F59E0B", command=self.imprimir_graficos).pack(side="left", padx=10)
        ctk.CTkButton(acciones_frame, text="⬇️ Gráficos a Word", fg_color="#334155", command=self.descargar_graficos).pack(side="left", padx=10)
        ctk.CTkButton(acciones_frame, text="⬇️ Exportar a Word", fg_color="#4F46E5", command=self.descargar_reportes_word).pack(side="left", padx=10)
        ctk.CTkButton(acciones_frame, text="📄 Informe del Estudiante", fg_color="#10B981", hover_color="#059669", command=self.generar_informe_estudiante_modal).pack(side="left", padx=10)

        # Panel de Controles
        self.frame_controles = ctk.CTkFrame(self, fg_color="#1E2D42", corner_radius=8)
        self.frame_controles.pack(fill="x", pady=5, ipadx=10, ipady=10)

        ctk.CTkLabel(self.frame_controles, text="Seleccione Grado:", font=("Segoe UI", 14)).pack(side="left", padx=10)

        opciones_grado = self.engine.obtener_grados_activos()
        if not opciones_grado: opciones_grado = ["Sin datos"]
        self.combo_grado = ctk.CTkOptionMenu(self.frame_controles, values=opciones_grado, command=self.cargar_reportes)
        self.combo_grado.pack(side="left", padx=10)

        # Tabs for different reports
        self.tabs = ctk.CTkTabview(self)
        self.tabs.pack(fill="both", expand=True, pady=10)

        self.tab_docente = self.tabs.add("Reporte Docente")
        self.tab_aprobados = self.tabs.add("Aprobados / Reprobados")
        self.tab_direccion = self.tabs.add("Reporte Dirección")
        self.tab_habitos = self.tabs.add("Hábitos y Aptitudes")

        # Docente
        self.scroll_docente = ctk.CTkScrollableFrame(self.tab_docente, fg_color="transparent")
        self.scroll_docente.pack(fill="both", expand=True)

        # Aprobados
        self.scroll_aprobados = ctk.CTkScrollableFrame(self.tab_aprobados, fg_color="transparent")
        self.scroll_aprobados.pack(fill="both", expand=True)

        # Direccion
        self.scroll_direccion = ctk.CTkScrollableFrame(self.tab_direccion, fg_color="transparent")
        self.scroll_direccion.pack(fill="both", expand=True)

        # Habitos
        self.scroll_habitos = ctk.CTkScrollableFrame(self.tab_habitos, fg_color="transparent")
        self.scroll_habitos.pack(fill="both", expand=True)

        self.cargar_reportes(self.combo_grado.get())

    def cargar_reportes(self, grado):
        self._limpiar(self.scroll_docente)
        self._limpiar(self.scroll_aprobados)
        self._limpiar(self.scroll_direccion)
        self._limpiar(self.scroll_habitos)

        reportes = getattr(self.engine, 'obtener_datos_reportes', lambda x: {"docente": [], "aprobados": [], "direccion": []})(grado)

        if not reportes["docente"]:
            ctk.CTkLabel(self.scroll_docente, text="No hay datos de resumen para mostrar.", font=("Segoe UI", 16)).pack(pady=20)
            ctk.CTkLabel(self.scroll_aprobados, text="No hay datos de resumen para mostrar.", font=("Segoe UI", 16)).pack(pady=20)
            ctk.CTkLabel(self.scroll_direccion, text="No hay datos de resumen para mostrar.", font=("Segoe UI", 16)).pack(pady=20)
            ctk.CTkLabel(self.scroll_habitos, text="No hay datos de resumen para mostrar.", font=("Segoe UI", 16)).pack(pady=20)
            return

        # DOCENTE HEADER
        f_h_doc = ctk.CTkFrame(self.scroll_docente, fg_color="#253650")
        f_h_doc.pack(fill="x", pady=2)
        ctk.CTkLabel(f_h_doc, text="Estudiante", width=250, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)
        ctk.CTkLabel(f_h_doc, text="Cédula", width=150, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)
        ctk.CTkLabel(f_h_doc, text="Nota Final", width=100, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)

        for fila in reportes["docente"]:
            row_d = ctk.CTkFrame(self.scroll_docente, fg_color="#1A2638")
            row_d.pack(fill="x", pady=1)
            ctk.CTkLabel(row_d, text=fila[0], width=250, anchor="w").pack(side="left", padx=10)
            ctk.CTkLabel(row_d, text=fila[1] if fila[1] else "N/A", width=150, anchor="w").pack(side="left", padx=10)
            ctk.CTkLabel(row_d, text=str(fila[2]), width=100, anchor="w").pack(side="left", padx=10)

        # APROBADOS HEADER
        f_h_apr = ctk.CTkFrame(self.scroll_aprobados, fg_color="#253650")
        f_h_apr.pack(fill="x", pady=2)
        ctk.CTkLabel(f_h_apr, text="Estudiante", width=250, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)
        ctk.CTkLabel(f_h_apr, text="Estado", width=150, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)

        for fila in reportes["aprobados"]:
            row_a = ctk.CTkFrame(self.scroll_aprobados, fg_color="#1A2638")
            row_a.pack(fill="x", pady=1)
            ctk.CTkLabel(row_a, text=fila[0], width=250, anchor="w").pack(side="left", padx=10)
            color_est = "#10B981" if str(fila[1]).upper() == "APROBADO" else "#EF4444"
            ctk.CTkLabel(row_a, text=str(fila[1]), width=150, anchor="w", text_color=color_est, font=("Segoe UI", 12, "bold")).pack(side="left", padx=10)

        # DIRECCION HEADER
        f_h_dir = ctk.CTkFrame(self.scroll_direccion, fg_color="#253650")
        f_h_dir.pack(fill="x", pady=2)
        ctk.CTkLabel(f_h_dir, text="Estudiante", width=250, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)
        ctk.CTkLabel(f_h_dir, text="Cédula", width=150, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)
        ctk.CTkLabel(f_h_dir, text="Nota Final", width=100, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)
        ctk.CTkLabel(f_h_dir, text="Estado", width=150, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)

        for fila in reportes["direccion"]:
            row_dir = ctk.CTkFrame(self.scroll_direccion, fg_color="#1A2638")
            row_dir.pack(fill="x", pady=1)
            ctk.CTkLabel(row_dir, text=fila[0], width=250, anchor="w").pack(side="left", padx=10)
            ctk.CTkLabel(row_dir, text=fila[1] if fila[1] else "N/A", width=150, anchor="w").pack(side="left", padx=10)
            ctk.CTkLabel(row_dir, text=str(fila[2]), width=100, anchor="w").pack(side="left", padx=10)
            color_est = "#10B981" if str(fila[3]).upper() == "APROBADO" else "#EF4444"
            ctk.CTkLabel(row_dir, text=str(fila[3]), width=150, anchor="w", text_color=color_est, font=("Segoe UI", 12, "bold")).pack(side="left", padx=10)

        # HABITOS TAB
        self._limpiar(self.scroll_habitos)
        
        import json
        ruta_json = os.path.abspath(os.path.join(os.path.dirname(self.engine.ruta), "Expedientes_Estudiantes", "habitos_evaluaciones.json"))
        
        estudiantes = self.engine.obtener_estudiantes_completos(grado)
        
        # Header
        f_h_hab = ctk.CTkFrame(self.scroll_habitos, fg_color="#253650")
        f_h_hab.pack(fill="x", pady=2)
        ctk.CTkLabel(f_h_hab, text="Estudiante", width=250, anchor="w", font=("Segoe UI", 13, "bold")).pack(side="left", padx=10)
        ctk.CTkLabel(f_h_hab, text="Satisfactorios (S)", width=130, anchor="w", font=("Segoe UI", 13, "bold"), text_color="#10B981").pack(side="left", padx=10)
        ctk.CTkLabel(f_h_hab, text="Regulares (R)", width=130, anchor="w", font=("Segoe UI", 13, "bold"), text_color="#F59E0B").pack(side="left", padx=10)
        ctk.CTkLabel(f_h_hab, text="Insatisfactorios (X)", width=130, anchor="w", font=("Segoe UI", 13, "bold"), text_color="#EF4444").pack(side="left", padx=10)

        # Aggregate stats per student
        stats_por_est = {}
        if os.path.exists(ruta_json):
            try:
                with open(ruta_json, "r", encoding="utf-8") as f:
                    habitos_data = json.load(f)
                
                for key, val_entry in habitos_data.items():
                    k_grado = key.split("::")[0]
                    if grado.replace("°","") in k_grado.replace("°",""):
                        est_evals = val_entry.get("estudiantes", {})
                        for est_id, crit_vals in est_evals.items():
                            e_idx = int(est_id) - 1
                            if e_idx not in stats_por_est:
                                stats_por_est[e_idx] = {"S": 0, "R": 0, "X": 0}
                            for score in crit_vals.values():
                                if score in ["S", "R", "X"]:
                                    stats_por_est[e_idx][score] += 1
            except Exception as e:
                print(f"Error loading habits report stats: {e}")

        for idx, est in enumerate(estudiantes):
            row_h = ctk.CTkFrame(self.scroll_habitos, fg_color="#1A2638")
            row_h.pack(fill="x", pady=1)
            
            ctk.CTkLabel(row_h, text=est["nombre"], width=250, anchor="w").pack(side="left", padx=10)
            
            stats = stats_por_est.get(idx, {"S": 0, "R": 0, "X": 0})
            total = sum(stats.values())
            
            s_txt = f"{stats['S']} ({(stats['S']/total*100):.0f}%)" if total > 0 else "0 (0%)"
            r_txt = f"{stats['R']} ({(stats['R']/total*100):.0f}%)" if total > 0 else "0 (0%)"
            x_txt = f"{stats['X']} ({(stats['X']/total*100):.0f}%)" if total > 0 else "0 (0%)"
            
            ctk.CTkLabel(row_h, text=s_txt, width=130, anchor="w", text_color="#10B981").pack(side="left", padx=10)
            ctk.CTkLabel(row_h, text=r_txt, width=130, anchor="w", text_color="#F59E0B").pack(side="left", padx=10)
            ctk.CTkLabel(row_h, text=x_txt, width=130, anchor="w", text_color="#EF4444").pack(side="left", padx=10)

    def _limpiar(self, frame):
        for w in frame.winfo_children(): w.destroy()

    def generar_informe_estudiante_modal(self):
        from rdsecurity import cargar_config_segura
        from documentos_maestro import generar_informe_calificaciones
        from utils.footer_utils import abrir_documento
        import datetime
        
        grado = self.combo_grado.get()
        if grado == "Sin datos":
            messagebox.showwarning("Atención", "No hay un grado válido seleccionado.")
            return
            
        ests = self.engine.obtener_estudiantes_completos(grado)
        if not ests:
            messagebox.showwarning("Atención", "No hay estudiantes registrados en este grado.")
            return
            
        # Modal para elegir estudiante
        modal = ctk.CTkToplevel(self)
        modal.title("📄 Generar Informe Individual")
        modal.geometry("350x200")
        modal.resizable(False, False)
        modal.transient(self)
        modal.grab_set()
        
        from config import establecer_icono_ventana
        establecer_icono_ventana(modal)
        
        modal.update_idletasks()
        w = modal.winfo_width()
        h = modal.winfo_height()
        x = self.winfo_toplevel().winfo_x() + (self.winfo_toplevel().winfo_width() // 2) - (w // 2)
        y = self.winfo_toplevel().winfo_y() + (self.winfo_toplevel().winfo_height() // 2) - (h // 2)
        modal.geometry(f"+{x}+{y}")
        
        ctk.CTkLabel(modal, text="Seleccione el Estudiante:", font=("Segoe UI", 14, "bold")).pack(pady=15)
        
        nombres_ests = [e["nombre"] for e in ests]
        combo_est = ctk.CTkOptionMenu(modal, values=nombres_ests, width=280)
        combo_est.pack(pady=10)
        
        def generar():
            nombre_sel = combo_est.get()
            est = next(e for e in ests if e["nombre"] == nombre_sel)
            modal.destroy()
            
            # Recopilar materias con al menos una nota registrada en T1, T2 o T3
            materias_con_nota = set()
            notas_por_trimestre = {}
            for t_num in [1, 2, 3]:
                notas_t = self.engine.obtener_notas_estudiante(est["nombre"], grado, trimestre=t_num)
                notas_por_trimestre[t_num] = notas_t
                for mat_nombre, nota_val in notas_t.items():
                    try:
                        val = float(nota_val)
                        if val >= 1.0:
                            materias_con_nota.add(mat_nombre)
                    except (TypeError, ValueError):
                        pass

            # Fallback: si el alumno no tiene ninguna nota en todo el año, usar todas las del grado
            if not materias_con_nota:
                all_mats = self.engine.obtener_materias_por_grado(grado)
                materias_con_nota = set([m for m in all_mats if m not in ["Sin materias registradas", "Sin materias", "No hay materias", "General"]])

            cfg = cargar_config_segura({})
            trimestres = {}
            for t_key, t_num in [("T1", 1), ("T2", 2), ("T3", 3)]:
                notas_t = notas_por_trimestre[t_num]
                materias = []
                for mat_nombre in sorted(list(materias_con_nota)):
                    nota_val = notas_t.get(mat_nombre, None)
                    if nota_val is not None:
                        try:
                            val = float(nota_val)
                        except (TypeError, ValueError):
                            val = 0.0
                    else:
                        val = 0.0
                    
                    materias.append({
                        "nombre": mat_nombre,
                        "nota": val,
                        "estado": "APROBADO" if val >= 3.0 else "REPROBADO"
                    })
                prom = sum(m["nota"] for m in materias if m["nota"] > 0) / len([m for m in materias if m["nota"] > 0]) if any(m["nota"] > 0 for m in materias) else 0
                trimestres[t_key] = {"materias": materias, "promedio": round(prom, 1)}
            
            datos = {
                "alumno":          est["nombre"],
                "cedula":          est.get("cedula", ""),
                "grado":           grado,
                "docente_nombre":  cfg.get("docente_nombre", ""),
                "escuela_nombre":  cfg.get("escuela_nombre", ""),
                "ano_lectivo":     cfg.get("ano_lectivo", "2026"),
                "fecha":           datetime.datetime.now().strftime("%d-%m-%Y"),
                "trimestres":      trimestres,
                "observacion_general": "",
                "estado_final":    "EN PROCESO"
            }
            
            # Promedio global
            proms = [trimestres[tk]["promedio"] for tk in trimestres if trimestres[tk].get("promedio")]
            if proms:
                prom_global = sum(proms) / len(proms)
                datos["estado_final"] = "APROBADO" if prom_global >= 3.0 else "REPROBADO"
            
            ruta = generar_informe_calificaciones(datos)
            if ruta:
                abrir_documento(ruta)
                messagebox.showinfo("✓ Informe Generado", f"Informe de calificaciones generado exitosamente:\n{ruta}")
                
        btn_gen = ctk.CTkButton(modal, text="Generar Word", fg_color="#10B981", hover_color="#059669", command=generar)
        btn_gen.pack(pady=15)

    def imprimir_reporte(self):
        # Abre el Excel para imprimir desde la hoja de resumen del grado seleccionado
        try:
            from rdprint import abrir_para_imprimir
            grado = self.combo_grado.get()
            self.engine.formatear_reportes_excel(grado)
            # Find the correct RESUMEN sheet. The format varies (e.g. RESUMEN (7°))
            hoja = None
            wb = self.engine._wb_cache
            if not wb:
                import openpyxl
                wb = openpyxl.load_workbook(self.engine.ruta, data_only=True)

            grado_num = grado.replace("°", "")
            for s in wb.sheetnames:
                if "RESUMEN" in s.upper() and (self.engine.modalidad == "primaria" or grado_num in s):
                    hoja = s
                    break

            if not hoja:
                messagebox.showerror("Error", f"No se encontró la hoja de resumen para {grado}.")
                return

            exito, msj = abrir_para_imprimir(hoja)
            messagebox.showinfo("Imprimir Reporte", msj)
        except Exception as e:
            messagebox.showerror("Error", f"Error al intentar imprimir: {e}")

    def imprimir_graficos(self):
        try:
            from rdprint import imprimir_hoja_directo
            from tkinter import messagebox
            import openpyxl

            grado = self.combo_grado.get()
            hoja = None
            wb = self.engine._wb_cache
            if not wb:
                wb = openpyxl.load_workbook(self.engine.ruta, data_only=True)

            grado_num = grado.replace("°", "")
            for s in wb.sheetnames:
                if "RESUMEN" in s.upper() and (self.engine.modalidad == "primaria" or grado_num in s):
                    hoja = s
                    break

            if not hoja:
                messagebox.showerror("Error", f"No se encontró la hoja para imprimir {grado}.")
                return

            exito, msj = imprimir_hoja_directo(hoja)
            if exito:
                messagebox.showinfo("Imprimir Gráficos", msj)
            else:
                messagebox.showerror("Error", msj)
        except Exception as e:
            messagebox.showerror("Error", f"Error al intentar imprimir: {e}")


    def descargar_graficos(self):
        # Genera y guarda un DOCX solo con los gráficos del grado seleccionado
        self._generar_docx_graficos(abrir=False)

    def descargar_pdf(self):
        import tempfile
        try:
            from docx import Document
            from docx.shared import Pt
            from utils.footer_utils import add_footer_with_logo, add_header_with_logo, get_school_logo_path
        except ImportError:
            messagebox.showerror("Error", "Falta librería python-docx.")
            return

        grado = self.combo_grado.get()
        reportes = getattr(self.engine, 'obtener_datos_reportes', lambda x: {"docente": [], "aprobados": [], "direccion": []})(grado)
        doc = Document()
        try:
            from utils.footer_utils import add_header_with_logo, get_school_logo_path
            logo_esc = get_school_logo_path()
            if logo_esc:
                add_header_with_logo(doc, logo_esc)
        except Exception: pass

        doc.add_heading(f"Reporte Académico — Grado {grado}", 0)
        # Encabezado escuela
        datos = self.engine.obtener_datos_generales() if hasattr(self.engine, 'obtener_datos_generales') else {}
        escuela = datos.get("escuela_nombre", "")
        region = datos.get("escuela_region", "")
        ano = datos.get("ano_lectivo", "")
        doc.add_paragraph(f"Escuela: {escuela}")
        doc.add_paragraph(f"Región: {region}")
        doc.add_paragraph(f"Año Lectivo: {ano}")
        doc.add_paragraph("")

        # Docente
        doc.add_heading("Reporte Docente", level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Estudiante'
        hdr_cells[1].text = 'Cédula'
        hdr_cells[2].text = 'Nota Final'
        for fila in reportes["docente"]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fila[0])
            row_cells[1].text = str(fila[1])
            row_cells[2].text = str(fila[2])
        doc.add_paragraph("")

        # Aprobados
        doc.add_heading("Aprobados", level=1)
        table2 = doc.add_table(rows=1, cols=2)
        hdr_cells2 = table2.rows[0].cells
        hdr_cells2[0].text = 'Estudiante'
        hdr_cells2[1].text = 'Estado'
        for fila in reportes["aprobados"]:
            if str(fila[1]).upper() == "APROBADO":
                row_cells = table2.add_row().cells
                row_cells[0].text = str(fila[0])
                row_cells[1].text = str(fila[1])
        doc.add_paragraph("")

        # Reprobados
        doc.add_heading("Reprobados", level=1)
        table3 = doc.add_table(rows=1, cols=2)
        hdr_cells3 = table3.rows[0].cells
        hdr_cells3[0].text = 'Estudiante'
        hdr_cells3[1].text = 'Estado'
        for fila in reportes["aprobados"]:
            if str(fila[1]).upper() != "APROBADO":
                row_cells = table3.add_row().cells
                row_cells[0].text = str(fila[0])
                row_cells[1].text = str(fila[1])
        doc.add_paragraph("")

        # Dirección
        doc.add_heading("Reporte Dirección", level=1)
        table4 = doc.add_table(rows=1, cols=4)
        hdr_cells4 = table4.rows[0].cells
        hdr_cells4[0].text = 'Estudiante'
        hdr_cells4[1].text = 'Cédula'
        hdr_cells4[2].text = 'Nota Final'
        hdr_cells4[3].text = 'Estado'
        for fila in reportes["direccion"]:
            row_cells = table4.add_row().cells
            row_cells[0].text = str(fila[0])
            row_cells[1].text = str(fila[1])
            row_cells[2].text = str(fila[2])
            row_cells[3].text = str(fila[3])

        # Gráficos integrados según tipo de reporte
        try:
            from utils.docx_graphics import save_figure_as_image, add_image_to_doc
            import matplotlib.pyplot as plt
            promedios_por_est = getattr(self.engine, 'obtener_promedios_reales', lambda g,m,t: {})(grado, None, None)
            estudiantes = getattr(self.engine, 'obtener_estudiantes_completos', lambda g: [])(grado)
            if not promedios_por_est and estudiantes:
                for est in estudiantes:
                    promedios_por_est[est['nombre']] = 1.0

            # Pastel
            etiquetas = ['Aprobados (>=3.0)', 'En Riesgo (2.5-2.9)', 'Reprobados (<2.5)']
            valores = [sum(1 for v in promedios_por_est.values() if v >= 3.0),
                       sum(1 for v in promedios_por_est.values() if 2.5 <= v < 3.0),
                       sum(1 for v in promedios_por_est.values() if v < 2.5)]
            colores = ["#00FF88", "#FFD700", "#FF4444"]
            fig1, ax1 = plt.subplots(figsize=(5, 4), dpi=100)
            if sum(valores) == 0:
                ax1.text(0.5, 0.5, "Sin Datos", ha='center', va='center', color='black')
            else:
                wedges, texts, autotexts = ax1.pie(valores, labels=etiquetas, colors=colores, autopct='%1.1f%%',
                                                  startangle=90)
            ax1.set_title("Distribución General del Salón", pad=20, fontweight="bold")
            img1 = save_figure_as_image(fig1, "pastel_")
            add_image_to_doc(doc, img1, ancho=5)
            plt.close(fig1)
        except Exception as e:
            print("Error integrando gráficos en el reporte PDF:", e)

        # Pie de página institucional
        try:
            logo_path = os.path.join(os.path.dirname(__file__), '..', 'img', 'icono.png')
            add_footer_with_logo(doc, logo_path, "RegistroDoc Pro v3.0", "Proverbios 22:6 — Instruye al niño en su camino, y aun cuando fuere viejo no se apartará de él.")
        except Exception:
            pass

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_docx = tmp.name
        try:
            import comtypes.client
            pdf_path = tmp_docx.replace(".docx", ".pdf")
            word = comtypes.client.CreateObject('Word.Application')
            doc_obj = word.Documents.Open(tmp_docx)
            doc_obj.SaveAs(pdf_path, FileFormat=17)
            doc_obj.Close()
            word.Quit()
            messagebox.showinfo("Descarga PDF", f"PDF generado en:\n{pdf_path}")
        except Exception:
            messagebox.showinfo("Descarga DOCX", f"Documento generado en:\n{tmp_docx}\n\nNo se pudo convertir a PDF automáticamente (requiere MS Word en Windows). Puedes abrirlo y guardarlo como PDF desde Word.")

    def _generar_docx_graficos(self, abrir=False, imprimir=False):
        import tempfile
        import os
        try:
            from docx import Document
            from utils.footer_utils import add_footer_with_logo, add_header_with_logo, get_school_logo_path
            from utils.docx_graphics import save_figure_as_image, add_image_to_doc
        except ImportError:
            messagebox.showerror("Error", "Falta librería python-docx o utils de gráficos.")
            return

        grado = self.combo_grado.get()
        materia = getattr(self, 'combo_materia', None)
        materia = materia.get() if materia else None
        trimestre = getattr(self, 'combo_trimestre', None)
        trimestre = trimestre.get() if trimestre else None

        doc = Document()
        try:
            from utils.footer_utils import add_header_with_logo, get_school_logo_path
            logo_esc = get_school_logo_path()
            if logo_esc:
                add_header_with_logo(doc, logo_esc)
        except Exception: pass

        doc.add_heading(f"Gráficos de Rendimiento — Grado {grado}", 0)

        engine = self.engine
        estudiantes = []
        try:
            estudiantes = self.engine.obtener_estudiantes_completos(grado)
        except Exception:
            pass

        promedios_por_est = getattr(self.engine, 'obtener_promedios_reales', lambda g,m,t: {})(grado, materia, trimestre)
        if not promedios_por_est and estudiantes:
            for est in estudiantes:
                promedios_por_est[est['nombre']] = 1.0

        aprobados = sum(1 for v in promedios_por_est.values() if v >= 3.0)
        en_riesgo = sum(1 for v in promedios_por_est.values() if 2.5 <= v < 3.0)
        reprobados = sum(1 for v in promedios_por_est.values() if v < 2.5)

# --- Gráficos ---
        import matplotlib.pyplot as plt
        import numpy as np

        # Pastel
        etiquetas = ['Aprobados (>=3.0)', 'En Riesgo (2.5-2.9)', 'Reprobados (<2.5)']
        valores = [aprobados, en_riesgo, reprobados]
        colores = ["#00FF88", "#FFD700", "#FF4444"]
        fig1, ax1 = plt.subplots(figsize=(5, 4), dpi=100)
        if sum(valores) == 0:
            ax1.text(0.5, 0.5, "Sin Datos", ha='center', va='center')
        else:
            wedges, texts, autotexts = ax1.pie(valores, labels=etiquetas, colors=colores, autopct='%1.1f%%',
                                              startangle=90)
        ax1.set_title("Distribución General del Salón", pad=20, fontweight="bold")
        img1 = save_figure_as_image(fig1, "pastel_")
        add_image_to_doc(doc, img1, ancho=5)
        plt.close(fig1)

        # Barras
        nombres = list(promedios_por_est.keys())[:15]
        notas = [promedios_por_est[n] for n in nombres]
        nombres_cortos = [n.split(" ")[0] for n in nombres]
        colores_barras = ["#FF4444" if n < 3.0 else "#00FF88" for n in notas]
        fig2, ax2 = plt.subplots(figsize=(6, 4), dpi=100)
        ax2.bar(nombres_cortos, notas, color=colores_barras, alpha=0.8)
        ax2.axhline(y=3.0, color="#FFD700", linestyle='--', alpha=0.7)
        ax2.tick_params(axis='x', rotation=45)
        ax2.set_ylim(1.0, 5.2)
        ax2.set_title("Top 15 Estudiantes (Promedios)", pad=15, fontweight="bold")
        fig2.tight_layout()
        img2 = save_figure_as_image(fig2, "barras_")
        add_image_to_doc(doc, img2, ancho=5)
        plt.close(fig2)

        # Tendencia / Histograma (Original y nuevo)
        fig3, ax3 = plt.subplots(figsize=(10, 3.5), dpi=100)
        notas_all = list(promedios_por_est.values())
        if notas_all:
            counts, bins, patches = ax3.hist(notas_all, bins=10, range=(1.0, 5.0), alpha=0.6, edgecolor='white')
            promedio_gen = sum(notas_all) / len(notas_all)
            ax3.axvline(x=promedio_gen, color="red", linestyle='-', linewidth=2, label=f'Promedio General ({promedio_gen:.1f})')
            ax3.legend()
        else:
            ax3.text(0.5, 0.5, "Sin Datos", ha='center', va='center')
        ax3.set_xticks([1.0, 2.0, 3.0, 4.0, 5.0])
        ax3.set_xlabel("Calificación")
        ax3.set_ylabel("Cant. Estudiantes")
        ax3.set_title("Distribución y Campana de Calificaciones", pad=15, fontweight="bold")
        fig3.tight_layout()
        img3 = save_figure_as_image(fig3, "tendencia_")
        add_image_to_doc(doc, img3, ancho=7)
        plt.close(fig3)

        # Nuevos: Pareto
        try:
            materias = self.engine.obtener_materias_por_grado(grado)
            if materias and materias != ["Sin materias"]:
                reprobados_por_materia = {}
                bulk_anual = getattr(self.engine, 'obtener_promedios_reales_bulk', lambda g,m,t: {})(grado, materias, "Anual")
                bulk_t1 = getattr(self.engine, 'obtener_promedios_reales_bulk', lambda g,m,t: {})(grado, materias, "Trimestre 1")

                for mat in materias:
                    proms = bulk_anual.get(mat, {})
                    if not proms: proms = bulk_t1.get(mat, {})
                    reps = sum(1 for v in proms.values() if v < 3.0)
                    reprobados_por_materia[mat] = reps
                reprobados_ordenado = dict(sorted(reprobados_por_materia.items(), key=lambda item: item[1], reverse=True))
                materias_nombres = list(reprobados_ordenado.keys())[:10]
                materias_cortas = [m[:10] for m in materias_nombres]
                conteos = [reprobados_ordenado[m] for m in materias_nombres]

                fig4, ax4 = plt.subplots(figsize=(6, 4), dpi=100)
                ax4_2 = ax4.twinx()
                if sum(conteos) > 0:
                    ax4.bar(materias_cortas, conteos, color="#FF4444", alpha=0.8)
                    acumulado = []
                    suma = 0
                    total = sum(conteos)
                    for c in conteos:
                        suma += c
                        acumulado.append(suma / total * 100)
                    ax4_2.plot(materias_cortas, acumulado, color="#FFD700", marker="D", ms=5)
                    ax4_2.set_ylim(0, 110)
                ax4.set_title("Pareto (Reprobación)", pad=15, fontweight="bold")
                ax4.tick_params(axis='x', rotation=45)
                fig4.tight_layout()
                img4 = save_figure_as_image(fig4, "pareto_")
                add_image_to_doc(doc, img4, ancho=5)
                plt.close(fig4)
        except Exception as e: print("Error Pareto", e)

        # Scatter
        try:
            fig5, ax5 = plt.subplots(figsize=(6, 4), dpi=100)
            if notas_all:
                x = []
                for n, p in promedios_por_est.items():
                    asistencia = 100
                    try:
                        fechas = getattr(self.engine, 'obtener_fechas_asistencia', lambda g,t: [])(grado, "Trimestre 1")
                        if fechas:
                            presentes = 0
                            total = len(fechas)
                            for fecha in fechas:
                                asis_dia = getattr(self.engine, 'buscar_asistencia_existente', lambda g,t,f: {})(grado, "Trimestre 1", fecha)
                                if asis_dia.get(n, "P") in ["P", "T"]: presentes += 1
                            asistencia = (presentes / total) * 100 if total > 0 else 100
                    except: pass
                    x.append(asistencia)

                ax5.scatter(x, notas_all, color="#00DDEB", alpha=0.7)
                ax5.set_xlabel("% Asistencia (Estimado)")
                ax5.set_ylabel("Nota Promedio")
            ax5.set_title("Scatter: Notas vs Asistencia", pad=15, fontweight="bold")
            fig5.tight_layout()
            img5 = save_figure_as_image(fig5, "scatter_")
            add_image_to_doc(doc, img5, ancho=5)
            plt.close(fig5)
        except Exception as e: print("Error Scatter", e)

        # Boxplot
        try:
            fig6, ax6 = plt.subplots(figsize=(6, 4), dpi=100)
            if notas_all:
                bplot = ax6.boxplot(notas_all, patch_artist=True, vert=True)
                for patch in bplot['boxes']: patch.set_facecolor("#00FF88")
                ax6.set_xticklabels([materia[:10] if materia else grado])
            ax6.set_title("Box-plot Comparativo", pad=15, fontweight="bold")
            fig6.tight_layout()
            img6 = save_figure_as_image(fig6, "box_")
            add_image_to_doc(doc, img6, ancho=5)
            plt.close(fig6)
        except Exception as e: print("Error Boxplot", e)

        # Evolucion
        try:
            fig7, ax7 = plt.subplots(figsize=(6, 4), dpi=100)
            p_t1 = getattr(self.engine, 'obtener_promedios_reales', lambda g,m,t: {})(grado, materia, "Trimestre 1")
            p_t2 = getattr(self.engine, 'obtener_promedios_reales', lambda g,m,t: {})(grado, materia, "Trimestre 2")
            p_t3 = getattr(self.engine, 'obtener_promedios_reales', lambda g,m,t: {})(grado, materia, "Trimestre 3")
            y_evol = [sum(p_t1.values())/len(p_t1) if p_t1 else 0, sum(p_t2.values())/len(p_t2) if p_t2 else 0, sum(p_t3.values())/len(p_t3) if p_t3 else 0]
            if any(y_evol):
                ax7.plot([1,2,3], y_evol, marker='o', color="#FFD700", linewidth=2)
                ax7.set_xticks([1,2,3])
                ax7.set_xticklabels(["T1", "T2", "T3"])
                ax7.set_ylim(1.0, 5.2)
            ax7.set_title("Evolución Trimestral", pad=15, fontweight="bold")
            fig7.tight_layout()
            img7 = save_figure_as_image(fig7, "evol_")
            add_image_to_doc(doc, img7, ancho=5)
            plt.close(fig7)
        except Exception as e: print("Error Evol", e)

        # Heatmap
        try:
            fig8, ax8 = plt.subplots(figsize=(6, 4), dpi=100)
            materias = self.engine.obtener_materias_por_grado(grado)
            if materias and materias != ["Sin materias"] and estudiantes:
                m_c = [m[:10] for m in materias[:5]]
                n_c = [e['nombre'].split(" ")[0] for e in estudiantes[:10]]
                # Pre-fetch promedios for all materias using bulk fetch
                materias_interes = materias[:5]
                bulk_reales = getattr(self.engine, 'obtener_promedios_reales_bulk', lambda g,ms,t: {})(grado, materias_interes, "Anual")
                bulk_t1 = getattr(self.engine, 'obtener_promedios_reales_bulk', lambda g,ms,t: {})(grado, materias_interes, "Trimestre 1")

                materias_reales = []
                for m in materias_interes:
                    pr = bulk_reales.get(m, {})
                    if not pr:
                        pr = bulk_t1.get(m, {})
                    materias_reales.append(pr)

                data_h = []
                for e in estudiantes[:10]:
                    row_data = [pr.get(e['nombre'], 0.0) for pr in materias_reales]
                    data_h.append(row_data)
                data_np = np.array(data_h)
                im = ax8.imshow(data_np, cmap="RdYlGn", vmin=1.0, vmax=5.0, aspect='auto')
                ax8.set_xticks(np.arange(len(m_c)))
                ax8.set_yticks(np.arange(len(n_c)))
                ax8.set_xticklabels(m_c, rotation=45)
                ax8.set_yticklabels(n_c)
                fig8.colorbar(im, ax=ax8)
            ax8.set_title("Heatmap: Notas (Top 10/5)", pad=15, fontweight="bold")
            fig8.tight_layout()
            img8 = save_figure_as_image(fig8, "heat_")
            add_image_to_doc(doc, img8, ancho=5)
            plt.close(fig8)
        except Exception as e: print("Error Heatmap", e)

        # Pie de página
        try:
            logo_path = os.path.join(os.path.dirname(__file__), '..', 'img', 'icono.png')
            add_footer_with_logo(doc, logo_path, "RegistroDoc Pro v3.0", "Proverbios 22:6 — Instruye al niño en su camino, y aun cuando fuere viejo no se apartará de él.")
        except Exception:
            pass

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_docx = tmp.name


        if imprimir:
            import platform
            if platform.system() == "Windows":
                try:
                    import win32api
                    win32api.ShellExecute(0, "print", tmp_docx, None, ".", 0)
                    messagebox.showinfo("Imprimir", "Enviado a la impresora.")
                except Exception as e:
                    messagebox.showerror("Error de impresión", f"No se pudo imprimir: {e}")
            else:
                try:
                    import subprocess
                    subprocess.call(['lp', os.path.abspath(tmp_docx)])
                    messagebox.showinfo("Imprimir", "Enviado a la cola de impresión.")
                except Exception as e:
                    messagebox.showerror("Error de impresión", f"No se pudo imprimir: {e}")
            return

        if abrir:

            try:
                os.startfile(tmp_docx)
            except AttributeError:
                import subprocess
                subprocess.call(['xdg-open', os.path.abspath(tmp_docx)])
        else:
            messagebox.showinfo("Descarga DOCX", f"Documento generado en:\n{tmp_docx}\n\nPuedes abrirlo y guardarlo como PDF desde Word.")

    def descargar_reportes_word(self):
        import tempfile
        import os
        try:
            from docx import Document
            from utils.footer_utils import add_footer_with_logo, add_header_with_logo, get_school_logo_path
        except ImportError:
            messagebox.showerror("Error", "Falta librería python-docx.")
            return

        grado = self.combo_grado.get()
        reportes = getattr(self.engine, 'obtener_datos_reportes', lambda x: {"docente": [], "aprobados": [], "direccion": []})(grado)
        doc = Document()
        try:
            from utils.footer_utils import add_header_with_logo, get_school_logo_path
            logo_esc = get_school_logo_path()
            if logo_esc:
                add_header_with_logo(doc, logo_esc)
        except Exception: pass

        doc.add_heading(f"Reportes y Estadísticas — Grado {grado}", 0)

        # Docente
        doc.add_heading("Reporte Docente", level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Estudiante'
        hdr_cells[1].text = 'Cédula'
        hdr_cells[2].text = 'Nota Final'
        for fila in reportes["docente"]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fila[0])
            row_cells[1].text = str(fila[1])
            row_cells[2].text = str(fila[2])
        doc.add_paragraph("")

        # Aprobados
        doc.add_heading("Aprobados / Reprobados", level=1)
        table2 = doc.add_table(rows=1, cols=2)
        hdr_cells2 = table2.rows[0].cells
        hdr_cells2[0].text = 'Estudiante'
        hdr_cells2[1].text = 'Estado'
        for fila in reportes["aprobados"]:
            row_cells = table2.add_row().cells
            row_cells[0].text = str(fila[0])
            row_cells[1].text = str(fila[1])
        doc.add_paragraph("")

        # Dirección
        doc.add_heading("Reporte Dirección", level=1)
        table4 = doc.add_table(rows=1, cols=4)
        hdr_cells4 = table4.rows[0].cells
        hdr_cells4[0].text = 'Estudiante'
        hdr_cells4[1].text = 'Cédula'
        hdr_cells4[2].text = 'Nota Final'
        hdr_cells4[3].text = 'Estado'
        for fila in reportes["direccion"]:
            row_cells = table4.add_row().cells
            row_cells[0].text = str(fila[0])
            row_cells[1].text = str(fila[1])
            row_cells[2].text = str(fila[2])
            row_cells[3].text = str(fila[3])

        try:
            logo_path = os.path.join(os.path.dirname(__file__), '..', 'img', 'icono.png')
            add_footer_with_logo(doc, logo_path, "RegistroDoc Pro v3.0", "Proverbios 22:6")
        except Exception: pass

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_docx = tmp.name

        messagebox.showinfo("Descarga DOCX", f"Reportes generados en Word:\n{tmp_docx}")

class ReportesYGraficosFrame(ctk.CTkFrame):
    def __init__(self, master, engine, **kwargs):
        super().__init__(master, fg_color="transparent", **kwargs)
        self.engine = engine

        # Tab view superior
        self.tabs = ctk.CTkTabview(self)
        self.tabs.pack(fill="both", expand=True, padx=10, pady=10)

        # Agregar pestañas
        self.tab_reportes = self.tabs.add("📋 Reportes y Descargas")
        self.tab_graficos = self.tabs.add("📈 Gráficos de Rendimiento")

        # Cargar los sub-frames correspondientes
        from grapp import GraficosFrame
        
        self.frame_reportes = ReportesFrame(self.tab_reportes, self.engine)
        self.frame_reportes.pack(fill="both", expand=True)

        self.frame_graficos = GraficosFrame(self.tab_graficos, self.engine)
        self.frame_graficos.pack(fill="both", expand=True)

    def actualizar_vista(self):
        """Actualiza ambos sub-frames cuando la vista activa cambia."""
        if hasattr(self.frame_reportes, "actualizar_vista"):
            try:
                self.frame_reportes.actualizar_vista()
            except Exception:
                pass
        
        if hasattr(self.frame_reportes, "cargar_reportes"):
            try:
                self.frame_reportes.cargar_reportes(self.frame_reportes.combo_grado.get())
            except Exception:
                pass

        if hasattr(self.frame_graficos, "actualizar_vista"):
            try:
                self.frame_graficos.actualizar_vista()
            except Exception:
                pass
        
        if hasattr(self.frame_graficos, "actualizar_graficos"):
            try:
                self.frame_graficos.actualizar_graficos()
            except Exception:
                pass

