import os
import json
import datetime
import tkinter as tk
import threading
from tkinter import messagebox
import customtkinter as ctk
from config import BASE_DIR
from theme import C

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

# Criterios oficiales de Panamá (MEDUCA)
CRITERIOS_PRIMARIA = [
    ("Responsabilidad", "Asistencia a clases y entrega puntual de tareas."),
    ("Orden y Aseo", "Cuidado de útiles, libretas y aseo personal."),
    ("Organización del Trabajo", "Planificación y estructura al hacer deberes."),
    ("Autodominio y Confianza en sí mismo", "Control emocional, seguridad y autonomía."),
    ("Iniciativa", "Proactividad, participación y propuestas nuevas."),
    ("Cooperación", "Trabajo en equipo y compañerismo en el aula."),
    ("Respeto a la propiedad ajena", "Cuidado de bienes del aula y de compañeros.")
]

CRITERIOS_PREMEDIA_MEDIA = [
    ("Responsabilidad", "Cumplimiento y compromiso con deberes y tareas."),
    ("Puntualidad", "Llegada a tiempo a clase y entrega de trabajos."),
    ("Honradez", "Honestidad académica (no copiar, no mentir) y rectitud."),
    ("Conciencia Cívica", "Respeto a símbolos patrios, normas y valores cívicos."),
    ("Organización del Trabajo", "Método ordenado y planificación de actividades."),
    ("Autodominio y Confianza en sí mismo", "Autocontrol emocional frente a retos y autoestima."),
    ("Iniciativa", "Proactividad para aprender y emprender tareas."),
    ("Cooperación", "Empatía, compañerismo y trabajo en grupo."),
    ("Respeto a la propiedad ajena", "Cuidado de materiales ajenos y del centro."),
    ("Modales", "Cortesía, vocabulario correcto y buen trato."),
    ("Orden y Aseo", "Limpieza de trabajos y excelente presentación personal."),
    ("Empleo del tiempo libre", "Uso productivo y sano de recesos y tiempo libre.")
]

class HabitosFrame(ctk.CTkFrame):
    def __init__(self, master, engine, **kwargs):
        super().__init__(master, fg_color="transparent", **kwargs)
        self.engine = engine

        # Estado interno
        self.estudiantes = []
        self.evaluaciones_temporales = {}  # id_estudiante -> {criterio: valor}
        self.combo_widgets = {}
        self.criterios_activos = []

        # Ruta de datos JSON
        from rdsecurity import cargar_config_segura
        cfg = cargar_config_segura({})
        ruta_base = cfg.get("ruta_exportacion")
        if not ruta_base:
            ruta_base = os.path.join(os.path.expanduser("~"), "Documents", "RegistroDoc")
        self.ruta_json = os.path.join(ruta_base, "Expedientes_Estudiantes", "habitos_evaluaciones.json")

        self.crear_interfaz()
        self.al_cambiar_grado_trimestre()

    def crear_interfaz(self):
        # 1. Contenedor Superior (Filtros y Controles)
        self.frame_filtros = ctk.CTkFrame(self, fg_color=C["card"], corner_radius=10)
        self.frame_filtros.pack(fill="x", padx=10, pady=10)

        # Fila 0: Grado y Trimestre
        ctk.CTkLabel(self.frame_filtros, text="Grado:", font=("Outfit", 12, "bold")).grid(row=0, column=0, padx=10, pady=10, sticky="w")
        self.combo_grado = ctk.CTkComboBox(
            self.frame_filtros, 
            values=self.engine.obtener_grados_activos() or ["Sin grados"], 
            command=self.al_cambiar_grado_trimestre,
            width=140
        )
        self.combo_grado.grid(row=0, column=1, padx=5, pady=10)
        grados_activos = self.engine.obtener_grados_activos()
        self.combo_grado.set(grados_activos[0] if grados_activos else "Sin grados")

        ctk.CTkLabel(self.frame_filtros, text="Trimestre:", font=("Outfit", 12, "bold")).grid(row=0, column=2, padx=10, pady=10, sticky="w")
        self.combo_trimestre = ctk.CTkComboBox(
            self.frame_filtros, 
            values=["Trimestre 1", "Trimestre 2", "Trimestre 3"], 
            command=self.al_cambiar_grado_trimestre,
            width=140
        )
        self.combo_trimestre.grid(row=0, column=3, padx=5, pady=10)
        self.combo_trimestre.set("Trimestre 1")

        # Fila 1: Frecuencia y Período/Fecha
        ctk.CTkLabel(self.frame_filtros, text="Frecuencia:", font=("Outfit", 12, "bold")).grid(row=1, column=0, padx=10, pady=10, sticky="w")
        self.combo_frecuencia = ctk.CTkComboBox(
            self.frame_filtros,
            values=["Trimestral", "Diario", "Semanal", "Mensual"],
            command=self.al_cambiar_frecuencia,
            width=140
        )
        self.combo_frecuencia.grid(row=1, column=1, padx=5, pady=10)
        self.combo_frecuencia.set("Trimestral")

        ctk.CTkLabel(self.frame_filtros, text="Período/Fecha:", font=("Outfit", 12, "bold")).grid(row=1, column=2, padx=10, pady=10, sticky="w")
        self.combo_periodo = ctk.CTkComboBox(
            self.frame_filtros,
            values=["Trimestre Completo"],
            command=self.al_cambiar_grado_trimestre,
            width=140
        )
        self.combo_periodo.grid(row=1, column=3, padx=5, pady=10)
        self.combo_periodo.set("Trimestre Completo")
        self.combo_periodo.configure(state="disabled")

        # Botón Autorellenar con IA (ocupa ambas filas)
        self.btn_sugerir = ctk.CTkButton(
            self.frame_filtros,
            text="🧠 Autorellenar con IA",
            font=("Outfit", 12, "bold"),
            fg_color="#3B82F6",
            hover_color="#2563EB",
            command=self.autorellenar_con_ia,
            width=160
        )
        self.btn_sugerir.grid(row=0, column=4, rowspan=2, padx=15, pady=10)

        # Botón Guardar (ocupa ambas filas)
        self.btn_guardar_todo = ctk.CTkButton(
            self.frame_filtros,
            text="💾 GUARDAR EVALUACIONES",
            font=("Outfit", 12, "bold"),
            fg_color="#10B981",
            hover_color="#059669",
            command=self.guardar_evaluaciones,
            width=190
        )
        self.btn_guardar_todo.grid(row=0, column=5, rowspan=2, padx=5, pady=10)

        # 2. Contenedor de la Tabla Matricial (Scrollable)
        self.scroll_tabla = ctk.CTkScrollableFrame(self, fg_color=C["card_alt"], corner_radius=10)
        self.scroll_tabla.pack(fill="both", expand=True, padx=10, pady=5)

    def al_cambiar_frecuencia(self, *args):
        frec = self.combo_frecuencia.get()
        if frec == "Trimestral":
            self.combo_periodo.configure(state="normal")
            self.combo_periodo.configure(values=["Trimestre Completo"])
            self.combo_periodo.set("Trimestre Completo")
            self.combo_periodo.configure(state="disabled")
        elif frec == "Diario":
            self.combo_periodo.configure(state="normal")
            hoy = datetime.date.today().strftime("%d-%m-%Y")
            self.combo_periodo.configure(values=[hoy])
            self.combo_periodo.set(hoy)
        elif frec == "Semanal":
            self.combo_periodo.configure(state="normal")
            semanas = [f"Semana {i}" for i in range(1, 16)]
            self.combo_periodo.configure(values=semanas)
            self.combo_periodo.set("Semana 1")
        elif frec == "Mensual":
            self.combo_periodo.configure(state="normal")
            meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
            self.combo_periodo.configure(values=meses)
            meses_map = {
                1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio",
                7: "Julio", 8: "Agosto", 9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
            }
            mes_actual = meses_map.get(datetime.date.today().month, "Enero")
            self.combo_periodo.set(mes_actual)
            
        self.al_cambiar_grado_trimestre()

    def al_cambiar_grado_trimestre(self, *args):
        grado = self.combo_grado.get()
        if not grado or grado == "Sin grados":
            return
        self.cargar_tabla_matricial()

    def _grid_widget_if_needed(self, widget, row, column, **kwargs):
        current_row = getattr(widget, "_grid_row", None)
        current_col = getattr(widget, "_grid_col", None)
        if current_row == row and current_col == column:
            return
        widget.grid(row=row, column=column, **kwargs)
        widget._grid_row = row
        widget._grid_col = column

    def _forget_widget_if_gridded(self, widget):
        if getattr(widget, "_grid_row", None) is not None:
            widget.grid_forget()
            widget._grid_row = None
            widget._grid_col = None

    def _obtener_todos_pool_widgets(self):
        pool = set()
        if hasattr(self, "_pool_headers"):
            pool.update(self._pool_headers)
        if hasattr(self, "_pool_filas"):
            for lbl_nom, opt_menus in self._pool_filas:
                pool.add(lbl_nom)
                pool.update(opt_menus)
        return pool

    def _ocultar_todos_widgets(self):
        if hasattr(self, "_pool_headers"):
            for h in self._pool_headers:
                self._forget_widget_if_gridded(h)
        if hasattr(self, "_pool_filas"):
            for lbl_nom, opt_menus in self._pool_filas:
                self._forget_widget_if_gridded(lbl_nom)
                for opt in opt_menus:
                    self._forget_widget_if_gridded(opt)

    def _limpiar_grid(self):
        pool_w = self._obtener_todos_pool_widgets()
        for w in self.scroll_tabla.winfo_children():
            if w not in pool_w:
                try:
                    w.destroy()
                except Exception:
                    pass

    def _obtener_header_label(self, index, text, text_color, wraplength, justify, font):
        if not hasattr(self, "_pool_headers"):
            self._pool_headers = []
        if index < len(self._pool_headers):
            lbl = self._pool_headers[index]
            lbl.configure(text=text, text_color=text_color, wraplength=wraplength, justify=justify, font=font)
            return lbl
        lbl = ctk.CTkLabel(self.scroll_tabla, text=text, font=font, text_color=text_color, wraplength=wraplength, justify=justify)
        self._pool_headers.append(lbl)
        return lbl

    def _obtener_fila_reciclada(self, index):
        if not hasattr(self, "_pool_filas"):
            self._pool_filas = []
        if index < len(self._pool_filas):
            return self._pool_filas[index]
        
        lbl_nom = ctk.CTkLabel(self.scroll_tabla, text="", font=("Outfit", 11, "bold"), text_color="#F3F4F6", anchor="w")
        
        opt_menus = []
        for _ in range(12):
            opt_menu = ctk.CTkOptionMenu(
                self.scroll_tabla,
                values=["-", "S", "R", "X"],
                width=65,
                height=24,
                font=("Outfit", 10, "bold"),
                button_hover_color="#374151"
            )
            opt_menus.append(opt_menu)
            
        row_tuple = (lbl_nom, opt_menus)
        self._pool_filas.append(row_tuple)
        return row_tuple

    def cargar_tabla_matricial(self):
        self._limpiar_grid()

        grado = self.combo_grado.get()
        trimestre = self.combo_trimestre.get()
        if not grado or grado == "Sin grados":
            return

        self.estudiantes = self.engine.obtener_estudiantes_completos(grado)
        if not self.estudiantes or "Sin estudiantes" in self.estudiantes[0]['nombre']:
            self._ocultar_todos_widgets()
            lbl = ctk.CTkLabel(self.scroll_tabla, text="No hay estudiantes inscritos en este grado.", font=("Outfit", 14, "italic"), text_color="#A3A3A3")
            lbl.pack(pady=40)
            return

        # Destroy non-pool widgets that might be left in scroll_tabla
        pool_w = self._obtener_todos_pool_widgets()
        for w in self.scroll_tabla.winfo_children():
            if w not in pool_w:
                try:
                    w.destroy()
                except Exception:
                    pass

        es_primaria = any(g in grado for g in ["1°", "2°", "3°", "4°", "5°", "6°"])
        criterios_info = CRITERIOS_PRIMARIA if es_primaria else CRITERIOS_PREMEDIA_MEDIA
        self.criterios_activos = [crit[0] for crit in criterios_info]

        trimestre_num = int(trimestre.replace("Trimestre ", ""))
        frecuencia = self.combo_frecuencia.get()
        periodo = self.combo_periodo.get()
        try:
            cursor = self.engine.db_conn.cursor()
            est_ids = [str(e['id']) for e in self.estudiantes]
            placeholders = ",".join(["?"] * len(est_ids))
            
            cursor.execute(f"""
                SELECT estudiante_id, criterio_codigo, nota 
                FROM habitos 
                WHERE estudiante_id IN ({placeholders}) AND trimestre = ? AND frecuencia = ? AND periodo = ?;
            """, est_ids + [trimestre_num, frecuencia, periodo])
            rows = cursor.fetchall()
            
            self.evaluaciones_temporales = {}
            for est in self.estudiantes:
                self.evaluaciones_temporales[str(est['id'])] = {crit: "-" for crit in self.criterios_activos}
                
            for est_id, crit_code, nota in rows:
                if est_id in self.evaluaciones_temporales and crit_code in self.evaluaciones_temporales[est_id]:
                    self.evaluaciones_temporales[est_id][crit_code] = nota
        except Exception as e:
            print(f"[!] Error cargando hábitos de SQLite: {e}")
            self.evaluaciones_temporales = {}
            for est in self.estudiantes:
                self.evaluaciones_temporales[str(est['id'])] = {crit: "-" for crit in self.criterios_activos}

        self.scroll_tabla.grid_columnconfigure(0, minsize=220)
        for c_idx in range(1, len(self.criterios_activos) + 1):
            self.scroll_tabla.grid_columnconfigure(c_idx, minsize=90)

        lbl_est = self._obtener_header_label(0, "Estudiante", "#3B82F6", 0, "left", ("Outfit", 12, "bold"))
        self._grid_widget_if_needed(lbl_est, row=0, column=0, padx=10, pady=10, sticky="w")

        for c_idx, crit in enumerate(self.criterios_activos, 1):
            lbl_crit = self._obtener_header_label(c_idx, crit, "#22D3EE", 85, "center", ("Outfit", 10, "bold"))
            self._grid_widget_if_needed(lbl_crit, row=0, column=c_idx, padx=5, pady=10)

        for idx in range(len(self.criterios_activos) + 1, len(self._pool_headers)):
            self._forget_widget_if_gridded(self._pool_headers[idx])

        self.combo_widgets = {}
        for r_idx, est in enumerate(self.estudiantes, 1):
            est_id = str(est['id'])
            lbl_nom, opt_menus = self._obtener_fila_reciclada(r_idx - 1)
            
            if lbl_nom.cget("text") != est['nombre']:
                lbl_nom.configure(text=est['nombre'])
            self._grid_widget_if_needed(lbl_nom, row=r_idx, column=0, padx=10, pady=4, sticky="w")
            
            def get_color(val):
                return "#10B981" if val == "S" else ("#F59E0B" if val == "R" else ("#EF4444" if val == "X" else "#4B5563"))
            
            for c_idx, crit in enumerate(self.criterios_activos, 1):
                val_actual = self.evaluaciones_temporales[est_id].get(crit, "-")
                opt_menu = opt_menus[c_idx - 1]
                
                color = get_color(val_actual)
                
                current_color = getattr(opt_menu, "_current_color", None)
                current_key = getattr(opt_menu, "_current_command_key", None)
                if current_color != color or current_key != (est_id, crit):
                    opt_menu.configure(
                        fg_color=color,
                        button_color=color,
                        command=lambda val, eid=est_id, cn=crit: self.cambiar_valor_celda(eid, cn, val)
                    )
                    opt_menu._current_color = color
                    opt_menu._current_command_key = (est_id, crit)
                
                if getattr(opt_menu, "_current_val", None) != val_actual:
                    opt_menu.set(val_actual)
                    opt_menu._current_val = val_actual
                
                self._grid_widget_if_needed(opt_menu, row=r_idx, column=c_idx, padx=5, pady=4)
                self.combo_widgets[(est_id, crit)] = opt_menu
                
            for idx in range(len(self.criterios_activos), 12):
                self._forget_widget_if_gridded(opt_menus[idx])

        for r_idx in range(len(self.estudiantes), len(self._pool_filas)):
            lbl_nom, opt_menus = self._pool_filas[r_idx]
            self._forget_widget_if_gridded(lbl_nom)
            for opt in opt_menus:
                self._forget_widget_if_gridded(opt)

    def cambiar_valor_celda(self, est_id, criterio, valor):
        self.evaluaciones_temporales[est_id][criterio] = valor
        opt_menu = self.combo_widgets.get((est_id, criterio))
        if opt_menu:
            color = "#10B981" if valor == "S" else ("#F59E0B" if valor == "R" else ("#EF4444" if valor == "X" else "#4B5563"))
            opt_menu.configure(fg_color=color, button_color=color)
            opt_menu._current_color = color
            opt_menu._current_val = valor

    def autorellenar_con_ia(self):
        grado = self.combo_grado.get()
        trimestre = self.combo_trimestre.get()
        trimestre_num = int(trimestre.replace("Trimestre ", ""))
        frecuencia = self.combo_frecuencia.get()
        if not self.estudiantes:
            return

        confirmar = messagebox.askyesno(
            "Autorellenar con IA",
            "¿Desea analizar el rendimiento, asistencia y disciplina de todos los estudiantes para rellenar las sugerencias automáticamente?\n\nEsto sobrescribirá las selecciones actuales no guardadas."
        )
        if not confirmar:
            return

        self.btn_sugerir.configure(text="⏳ ANALIZANDO...", state="disabled")
        self.update()

        try:
            for est in self.estudiantes:
                id_est = est['id']
                est_id_str = str(id_est)
                
                # Intentar consolidar desde evaluaciones históricas diarias/semanales/mensuales del trimestre en SQLite
                cursor = self.engine.db_conn.cursor()
                cursor.execute("""
                    SELECT criterio_codigo, nota 
                    FROM habitos 
                    WHERE estudiante_id = ? AND trimestre = ? AND frecuencia != 'Trimestral' AND nota IN ('S', 'R', 'X');
                """, (est_id_str, trimestre_num))
                past_records = cursor.fetchall()
                
                by_crit = {}
                for crit_code, nota in past_records:
                    if crit_code not in by_crit:
                        by_crit[crit_code] = []
                    by_crit[crit_code].append(nota)
                
                stats_asistencia = self.engine.obtener_estadisticas_asistencia(grado, trimestre, id_est)
                stats_notas = self.engine.obtener_tareas_sin_nota(grado, trimestre, id_est)
                
                reportes_disciplina = []
                from rdsecurity import cargar_config_segura
                cfg = cargar_config_segura({})
                ruta_base = cfg.get("ruta_exportacion")
                if not ruta_base:
                    ruta_base = os.path.join(os.path.expanduser("~"), "Documents", "RegistroDoc")
                carpeta_exp = os.path.join(ruta_base, "Expedientes_Estudiantes")
                nombre_archivo = f"{est['nombre']} - {grado.replace('°', '')}.docx".replace("/", "-")
                ruta_word = os.path.join(carpeta_exp, nombre_archivo)
                if os.path.exists(ruta_word) and DOCX_DISPONIBLE:
                    try:
                        doc = Document(ruta_word)
                        if doc.tables:
                            tabla = doc.tables[0]
                            for row in tabla.rows[1:]:
                                if len(row.cells) >= 4:
                                    tipo = row.cells[2].text.strip()
                                    motivo = row.cells[3].text.strip()
                                    if any(k in tipo for k in ["Conducta", "Citación", "Reporte", "Disciplina"]):
                                        reportes_disciplina.append(motivo)
                    except Exception:
                        pass
                
                sugerencias = {}
                aus_sin_just = stats_asistencia.get("ausencias", 0)
                tardanzas = stats_asistencia.get("tardanzas", 0)
                tareas_vacias = stats_notas.get("total_vacias", 0)
                
                # Consolidar cada criterio
                for crit in self.criterios_activos:
                    if crit in by_crit and by_crit[crit]:
                        # Consolidador: S=3, R=2, X=1. Si promedio >= 2.5 -> S, >= 1.7 -> R, sino X.
                        scores = [3 if n == 'S' else (2 if n == 'R' else 1) for n in by_crit[crit]]
                        avg_score = sum(scores) / len(scores)
                        if avg_score >= 2.5:
                            sugerencias[crit] = "S"
                        elif avg_score >= 1.7:
                            sugerencias[crit] = "R"
                        else:
                            sugerencias[crit] = "X"
                    else:
                        # Fallback a reglas heurísticas estándar de IA
                        if crit == "Responsabilidad":
                            resp_val = "S"
                            if aus_sin_just >= 4 or tareas_vacias >= 3:
                                resp_val = "X"
                            elif aus_sin_just >= 2 or tareas_vacias >= 1:
                                resp_val = "R"
                            sugerencias["Responsabilidad"] = resp_val
                        elif crit == "Puntualidad":
                            punt_val = "S"
                            if tardanzas >= 6:
                                punt_val = "X"
                            elif tardanzas >= 3:
                                punt_val = "R"
                            sugerencias["Puntualidad"] = punt_val
                        elif crit == "Organización del Trabajo":
                            org_val = "S"
                            if tareas_vacias >= 3:
                                org_val = "X"
                            elif tareas_vacias >= 1:
                                org_val = "R"
                            sugerencias["Organización del Trabajo"] = org_val
                        elif crit in ["Autodominio y Confianza en sí mismo", "Cooperación", "Modales", "Honradez"]:
                            con_val = "S"
                            coop_val = "S"
                            mod_val = "S"
                            hon_val = "S"
                            if reportes_disciplina:
                                con_val = "R"
                                coop_val = "R"
                                mod_val = "R"
                                hon_val = "R"
                                if any("Grave" in r or "Suspensión" in r for r in reportes_disciplina):
                                    con_val = "X"
                                    coop_val = "X"
                                    mod_val = "X"
                            if crit == "Autodominio y Confianza en sí mismo":
                                sugerencias[crit] = con_val
                            elif crit == "Cooperación":
                                sugerencias[crit] = coop_val
                            elif crit == "Modales":
                                sugerencias[crit] = mod_val
                            elif crit == "Honradez":
                                sugerencias[crit] = hon_val
                        else:
                            sugerencias[crit] = "S"
                
                # Asignar sugerencias a la vista
                for crit, val in sugerencias.items():
                    self.evaluaciones_temporales[est_id_str][crit] = val
                    opt_menu = self.combo_widgets.get((est_id_str, crit))
                    if opt_menu:
                        opt_menu.set(val)
                        color = "#10B981" if val == "S" else ("#F59E0B" if val == "R" else ("#EF4444" if val == "X" else "#4B5563"))
                        opt_menu.configure(fg_color=color, button_color=color)

            root = self.winfo_toplevel()
            if hasattr(root, "mostrar_toast"):
                root.mostrar_toast("✓ Autocompletado IA finalizado", color="#10B981")
        except Exception as e:
            messagebox.showerror("Error", f"Error en el análisis de IA: {e}")
        finally:
            self.btn_sugerir.configure(text="🧠 Autorellenar con IA", state="normal")

    def guardar_evaluaciones(self):
        grado = self.combo_grado.get()
        trimestre = self.combo_trimestre.get()
        trimestre_num = int(trimestre.replace("Trimestre ", ""))
        frecuencia = self.combo_frecuencia.get()
        periodo = self.combo_periodo.get()

        if not self.evaluaciones_temporales:
            messagebox.showwarning("Atención", "No hay datos de evaluación para guardar.")
            return

        self.btn_guardar_todo.configure(text="⏳ GUARDANDO...", state="disabled")
        self.btn_sugerir.configure(state="disabled")
        self.update()

        def tarea_fondo():
            try:
                cursor = self.engine.db_conn.cursor()
                # 1. Guardar en SQLite (Base de datos principal)
                for est_id, crit_vals in self.evaluaciones_temporales.items():
                    for crit_code, nota in crit_vals.items():
                        cursor.execute("""
                            SELECT id FROM habitos 
                            WHERE estudiante_id = ? AND trimestre = ? AND criterio_codigo = ? AND frecuencia = ? AND periodo = ?;
                        """, (est_id, trimestre_num, crit_code, frecuencia, periodo))
                        row = cursor.fetchone()
                        if row:
                            cursor.execute("""
                                UPDATE habitos SET nota = ? WHERE id = ?;
                            """, (nota, row[0]))
                        else:
                            cursor.execute("""
                                INSERT INTO habitos (estudiante_id, trimestre, criterio_codigo, nota, frecuencia, periodo) 
                                VALUES (?, ?, ?, ?, ?, ?);
                            """, (est_id, trimestre_num, crit_code, nota, frecuencia, periodo))
                
                self.engine.db_conn.commit()
                self.engine.db_manager.guardar_cifrado()

                # 2. Guardar en JSON/Word (Expedientes locales) para compatibilidad
                datos_json = {}
                if os.path.exists(self.ruta_json):
                    try:
                        with open(self.ruta_json, "r", encoding="utf-8") as f:
                            datos_json = json.load(f)
                    except Exception:
                        pass
                
                clave_json = f"{grado}::{trimestre}::{frecuencia}::{periodo}"
                datos_json[clave_json] = {
                    "frecuencia": frecuencia,
                    "periodo": periodo,
                    "estudiantes": {str(k): v for k, v in self.evaluaciones_temporales.items()}
                }
                
                carpeta_exp = os.path.dirname(self.ruta_json)
                if not os.path.exists(carpeta_exp):
                    os.makedirs(carpeta_exp)
                
                with open(self.ruta_json, "w", encoding="utf-8") as f:
                    json.dump(datos_json, f, ensure_ascii=False, indent=4)

                # Actualizar Word de cada estudiante si está disponible
                if DOCX_DISPONIBLE:
                    for est in self.estudiantes:
                        id_est = est['id']
                        id_est_str = str(id_est)
                        if id_est_str in self.evaluaciones_temporales:
                            nombre_archivo = f"{est['nombre']} - {grado.replace('°', '')}.docx".replace("/", "-")
                            ruta_word = os.path.join(carpeta_exp, nombre_archivo)
                            
                            if os.path.exists(ruta_word):
                                try:
                                    doc = Document(ruta_word)
                                    if doc.tables:
                                        tabla = doc.tables[0]
                                        crit_vals = self.evaluaciones_temporales[id_est_str]
                                        eval_str = ", ".join([f"{k}: {v}" for k, v in crit_vals.items()])
                                        
                                        tipo_registro = f"Hábitos ({trimestre}) [{frecuencia} - {periodo}]"
                                        fecha_act = datetime.date.today().strftime("%d-%m-%Y")
                                        
                                        fila_existente = None
                                        for row in tabla.rows[1:]:
                                            if len(row.cells) >= 4:
                                                t = row.cells[2].text.strip()
                                                if t == tipo_registro:
                                                    fila_existente = row
                                                    break
                                        
                                        if fila_existente:
                                            fila_existente.cells[1].text = fecha_act
                                            fila_existente.cells[3].text = eval_str
                                        else:
                                            num_reg = str(len(tabla.rows))
                                            fila = tabla.add_row()
                                            fila.cells[0].text = num_reg
                                            fila.cells[1].text = fecha_act
                                            fila.cells[2].text = tipo_registro
                                            fila.cells[3].text = eval_str
                                        
                                        doc.save(ruta_word)
                                except Exception as e:
                                    print(f"Error escribiendo en Word de {est['nombre']}: {e}")

                self.after(0, lambda: self.finalizar_guardado(True, ""))
            except Exception as e:
                self.after(0, lambda: self.finalizar_guardado(False, str(e)))

        threading.Thread(target=tarea_fondo, daemon=True).start()

    def finalizar_guardado(self, exito, error_msg):
        self.btn_guardar_todo.configure(text="💾 GUARDAR EVALUACIONES", state="normal")
        self.btn_sugerir.configure(state="normal")
        
        if not exito:
            messagebox.showerror("Error", f"No se pudo guardar la evaluación: {error_msg}")
            return
            
        root = self.winfo_toplevel()
        if hasattr(root, "mostrar_toast"):
            root.mostrar_toast("✓ Evaluaciones guardadas con éxito", color="#10B981")
        else:
            messagebox.showinfo("Éxito", "Evaluaciones registradas y guardadas correctamente.")

    def actualizar_vista(self):
        opciones = self.engine.obtener_grados_activos() or ["Sin grados"]
        old_sel = self.combo_grado.get()
        self.combo_grado.configure(values=opciones)
        if old_sel in opciones:
            self.combo_grado.set(old_sel)
        else:
            self.combo_grado.set(opciones[0])
            
        try:
            from utils.date_helpers import obtener_trimestre_actual
            curr_trim = obtener_trimestre_actual()
            self.combo_trimestre.set(curr_trim)
        except Exception:
            self.combo_trimestre.set("Trimestre 1")
            
        self.combo_frecuencia.set("Trimestral")
        self.al_cambiar_frecuencia()
