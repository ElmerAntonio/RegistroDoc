from datetime import datetime
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Inches
from PIL import Image
from utils.frases_educacion import frase_del_dia

def add_footer_with_logo(doc, logo_path, nombre_programa, versiculo=None, fecha=None):
    section = doc.sections[-1]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_para.add_run()
    # Logo
    if os.path.exists(logo_path):
        run.add_picture(logo_path, width=Inches(0.7))
    # Nombre del programa y fecha
    fecha_str = fecha or datetime.now().strftime('%d/%m/%Y')
    run.add_text(f"   {nombre_programa} — {fecha_str}\n")
    # Frase o versículo educativo
    texto, ref = frase_del_dia() if versiculo is None else (versiculo, "")
    run.add_text(f"{texto} {f'({ref})' if ref else ''}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    return doc



def add_header_with_logo(doc, logo_path, text=""):
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = header_para.add_run()
    if os.path.exists(logo_path):
        run.add_picture(logo_path, width=Inches(1.0))
    if text:
        run.add_text(f"  {text}")
    return doc

def get_school_logo_path():
    # 1. Intentar obtener desde el archivo de configuración cifrado (prioridad usuario)
    try:
        from rdsecurity import cargar_config_segura
        data = cargar_config_segura({})
        config_path = data.get("logo_escuela_path", "")
        if config_path and os.path.exists(config_path):
            return config_path
    except Exception:
        pass

    # 2. Intentar obtener desde el paquete (PyInstaller fallback)
    import sys
    if hasattr(sys, "_MEIPASS"):
        internal_path = os.path.join(sys._MEIPASS, "assets", "logo.png")
    else:
        from config import BASE_DIR
        internal_path = os.path.normpath(os.path.join(BASE_DIR, "..", "assets", "logo.png"))

    if os.path.exists(internal_path):
        return internal_path

    return ""

def abrir_documento(ruta):
    """Abre el archivo recién creado en el visor predeterminado del sistema (multiplataforma)."""
    import subprocess
    try:
        if sys.platform == "win32":
            os.startfile(ruta)
        elif sys.platform == "darwin":
            subprocess.run(["open", ruta])
        else:
            subprocess.run(["xdg-open", ruta])
    except Exception:
        pass # No bloquear si no se puede abrir

